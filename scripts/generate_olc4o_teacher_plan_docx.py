import json
import re
from datetime import datetime, timezone
from html import unescape
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


WORKSPACE_ROOT = Path("D:/工作文件/SUNNYBROOK")
REPO_ROOT = WORKSPACE_ROOT / "ossd-course-portal"
COURSE_ROOT = WORKSPACE_ROOT / "courseware" / "OLC4O"
MANIFEST_PATH = COURSE_ROOT / "course-manifest.json"
REPORT_PATH = REPO_ROOT / "deployment" / "OLC4O-docx-plans-report.json"
GENERATED_AT = datetime.now(timezone.utc).replace(microsecond=0).isoformat()
TEMPLATE_ROOT = REPO_ROOT / "templates" / "teacher-plans"
LESSON_PLAN_TEMPLATE = TEMPLATE_ROOT / "mdm4u-style-lesson-plan-template.docx"
UNIT_PLAN_TEMPLATE = TEMPLATE_ROOT / "mdm4u-style-unit-plan-template.docx"


SUBJECT = "OLC4O Ontario Secondary School Literacy Course, Grade 12 Open"

STRANDS = {
    "A": {
        "title": "Building Reading Skills",
        "focus": "reading informational, narrative, and graphic texts; using strategies before, during, and after reading; explaining meaning with evidence.",
    },
    "B": {
        "title": "Building Writing Skills",
        "focus": "planning, drafting, revising, editing, and publishing clear writing for audience, purpose, and context.",
    },
    "C": {
        "title": "Understanding and Assessing Growth in Literacy",
        "focus": "monitoring literacy growth, using feedback, maintaining a portfolio, and reflecting on reading and writing progress.",
    },
}

UNIT_INTENT = {
    1: {
        "title": "Reading and Responding to Texts",
        "big_idea": "Successful literacy learners read informational, narrative, and graphic texts actively, using structure, text features, vocabulary clues, and strategies to construct meaning.",
        "teacher_intent": "Build a common reading process and vocabulary so students can move from literal comprehension to evidence-based responses.",
        "essential": [
            "How do different text forms guide the way readers find and organize meaning?",
            "Which strategies help students monitor comprehension before, during, and after reading?",
            "How can students support answers with accurate evidence from a text?",
        ],
    },
    2: {
        "title": "Writing for Audience, Purpose, and Context",
        "big_idea": "Effective writing is planned for a reader and purpose, then revised for clarity, organization, grammar, and voice.",
        "teacher_intent": "Move students from generating ideas to producing polished summaries, paragraphs, news reports, and reflections.",
        "essential": [
            "How do audience, purpose, and context shape writing choices?",
            "How can pre-writing and revision make student writing clearer?",
            "What conventions help readers trust and understand a writer's message?",
        ],
    },
    3: {
        "title": "Communication, Participation, and Literacy Processes",
        "big_idea": "Students strengthen literacy by communicating clearly, participating actively, and using reading and writing processes deliberately.",
        "teacher_intent": "Help students name the habits that support independent literacy work and prepare evidence for portfolio reflection.",
        "essential": [
            "What does active participation look like in a literacy classroom?",
            "How do readers and writers use feedback and reflection to improve?",
            "How can students explain their growth using specific evidence?",
        ],
    },
    4: {
        "title": "Portfolio, Culminating, and Final Examination",
        "big_idea": "The final portfolio and examination ask students to synthesize their reading, writing, reflection, and process evidence into clear demonstrations of literacy growth.",
        "teacher_intent": "Guide students through portfolio organization, final product planning, revision checkpoints, and exam readiness without replacing the original Moodle task files.",
        "essential": [
            "Which evidence best demonstrates growth in reading and writing?",
            "How can students revise a culminating product for clarity, organization, and audience?",
            "What routines help students prepare responsibly for the final examination?",
        ],
    },
}


def clean_text(value):
    text = str(value or "")
    text = re.sub(r"<script[\s\S]*?</script>", " ", text, flags=re.I)
    text = re.sub(r"<style[\s\S]*?</style>", " ", text, flags=re.I)
    text = re.sub(r"<br\s*/?>", "\n", text, flags=re.I)
    text = re.sub(r"</(p|div|li|tr|h[1-6])>", "\n", text, flags=re.I)
    text = re.sub(r"<[^>]+>", " ", text)
    text = unescape(text).replace("\u00a0", " ")
    return re.sub(r"\s+", " ", text).strip()


def read_local_text(relative_path):
    if not relative_path:
        return ""
    path = COURSE_ROOT / relative_path
    if not path.exists():
        return ""
    return clean_text(path.read_text(encoding="utf-8", errors="ignore"))


def short_text(value, limit=320):
    text = clean_text(value)
    if len(text) <= limit:
        return text
    clipped = text[:limit]
    cut = max(clipped.rfind(". "), clipped.rfind("; "), clipped.rfind(", "))
    if cut > 140:
        clipped = clipped[: cut + 1]
    return clipped.strip() + "..."


def unique(items):
    out = []
    seen = set()
    for item in items:
        value = clean_text(item)
        key = value.lower()
        if value and key not in seen:
            seen.add(key)
            out.append(value)
    return out


def find_section(lesson, pattern):
    regex = re.compile(pattern, re.I)
    for section in lesson.get("bookSections") or []:
        label = section.get("sectionLabel") or section.get("label") or ""
        if regex.search(label):
            return section
    return None


def section_text(lesson, pattern):
    section = find_section(lesson, pattern)
    if not section:
        return ""
    return read_local_text(section.get("path")) or clean_text(section.get("textPreview"))


def strip_repeated_heading(text, lesson):
    title = re.escape(str(lesson.get("title") or ""))
    patterns = [
        rf"^OLC4O Unit \d+ Lesson \d+ - [A-Za-z &,]+",
        rf"^Lesson\s+\d+\s*:\s*{title}\s+Lesson Expectations",
        rf"^Lesson\s+\d+\s*:\s*{title}",
    ]
    out = clean_text(text)
    changed = True
    while changed:
        changed = False
        for pattern in patterns:
            new = re.sub(pattern, "", out, flags=re.I).strip()
            if new != out:
                out = clean_text(new)
                changed = True
    return out


def expectation_codes(text):
    return unique(code.replace(" ", "") for code in re.findall(r"\b[A-C]\s*\d(?:\.\d+)?\b", text or ""))[:18]


def split_expectation_items(block, limit=10):
    block = clean_text(block)
    if not block:
        return []
    pieces = re.split(r"(?=\b[A-C]\s*\d(?:\.\d+)?\b)|(?<=\.)\s+(?=[A-Z][a-z]+)", block)
    return [short_text(piece, 300).rstrip(".") for piece in pieces if len(clean_text(piece)) > 12][:limit]


def extract_between(text, start_labels, end_labels):
    lower = text.lower()
    start = -1
    label_len = 0
    for label in start_labels:
        idx = lower.find(label.lower())
        if idx >= 0 and (start < 0 or idx < start):
            start = idx
            label_len = len(label)
    if start < 0:
        return ""
    start += label_len
    if start < len(text) and text[start] == ":":
        start += 1
    end = len(text)
    for label in end_labels:
        idx = lower.find(label.lower(), start)
        if idx >= 0:
            end = min(end, idx)
    return clean_text(text[start:end])


def expectation_for(lesson):
    raw = strip_repeated_heading(section_text(lesson, r"expectations"), lesson)
    overall = extract_between(
        raw,
        ["Overall Expectations", "Overall Expectation"],
        ["Specific Lesson Expectations", "Specific Expectations", "Learning Goals", "Success Criteria"],
    )
    specific = extract_between(
        raw,
        ["Specific Lesson Expectations", "Specific Expectations", "Specific Lesson Expectation"],
        ["Learning Goals", "Success Criteria"],
    )
    goals = extract_between(raw, ["Learning Goals", "Learning Goal"], ["Success Criteria"])
    success = extract_between(raw, ["Success Criteria", "Success Criterion"], [])
    if not overall and raw:
        overall = raw
    if not specific and raw:
        specific = raw
    return {
        "raw": raw,
        "codes": expectation_codes(raw),
        "overall": split_expectation_items(overall, 6),
        "specific": split_expectation_items(specific, 8),
        "goals": split_expectation_items(goals, 5),
        "success": split_expectation_items(success, 5),
    }


def all_lesson_attachments(lesson):
    records = []
    for section in lesson.get("bookSections") or []:
        records.extend(section.get("attachments") or [])
    records.extend(lesson.get("downloads") or [])
    records.extend(lesson.get("textExports") or [])
    return records


def resource_labels(items):
    labels = []
    for item in items:
        label = item.get("label") or item.get("title") or Path(str(item.get("path") or "")).name
        if label:
            labels.append(label)
    return unique(labels)


def lesson_materials(lesson):
    media = []
    files = []
    for item in lesson.get("ispring") or []:
        media.append(f"iSpring lesson presentation: {item.get('label') or lesson.get('id')}")
    for item in (lesson.get("handsOn") or []) + (lesson.get("h5p") or []):
        role = item.get("role") or "practice"
        media.append(f"H5P/{role} activity: {item.get('label') or 'localized activity'}")
    for item in all_lesson_attachments(lesson):
        label = item.get("label") or item.get("title") or Path(str(item.get("path") or "")).name
        haystack = " ".join(str(item.get(k, "")) for k in ("type", "category", "role", "path", "label"))
        if re.search(r"h5p|ispring|video|mp4", haystack, re.I):
            media.append(label)
        elif re.search(r"docx?|pdf|worksheet|rubric|assignment|organizer|reflection|tracking|exit", haystack, re.I):
            files.append(label)
    return unique(media), unique(files)


def lesson_activity_summary(lesson):
    pieces = []
    for pattern, label, limit in [
        (r"^Lesson$", "Lesson page", 220),
        (r"Hands On", "Hands On", 180),
        (r"Consoldation|Consolidation", "Consolidation", 180),
        (r"Homework", "Homework", 180),
    ]:
        text = section_text(lesson, pattern)
        if text:
            pieces.append(f"{label}: {short_text(text, limit)}")
    return "\n".join(pieces) or "Use the localized Moodle activity sequence and attached student resources."


def lesson_goals(unit, lesson, expectation):
    if expectation["goals"]:
        return expectation["goals"]
    title = lesson.get("title") or "the lesson topic"
    if unit["unit"] == 1:
        return [
            f"Identify the form and purpose of texts used in {title}.",
            "Use reading strategies to find main ideas, supporting details, and implied meaning.",
            "Respond to text-based questions using accurate evidence.",
        ]
    if unit["unit"] == 2:
        return [
            f"Plan and draft writing connected to {title}.",
            "Make writing choices that fit audience, purpose, and context.",
            "Revise for organization, clarity, grammar, and completeness.",
        ]
    return [
        f"Apply the literacy process emphasized in {title}.",
        "Use feedback, reflection, and participation routines to strengthen independent learning.",
        "Document evidence of reading and writing growth for the course portfolio.",
    ]


def success_criteria(unit, lesson, expectation):
    if expectation["success"]:
        return expectation["success"]
    if unit["unit"] == 1:
        return [
            "I can identify the text form and explain how its structure supports meaning.",
            "I can select relevant evidence and explain how it supports my response.",
            "I can use context, features, and strategies when a text is challenging.",
        ]
    if unit["unit"] == 2:
        return [
            "I can plan writing before drafting.",
            "I can organize ideas for a specific audience and purpose.",
            "I can revise grammar, sentence clarity, and word choice before submission.",
        ]
    return [
        "I can participate actively and communicate ideas clearly.",
        "I can describe the reading or writing process I used.",
        "I can choose portfolio evidence and explain what it shows about my growth.",
    ]


def prior_knowledge(unit, lesson):
    title = lesson.get("title") or "this lesson"
    if unit["unit"] == 1:
        return f"Students have experience reading everyday texts. Activate prior knowledge by asking them to name where they encounter {title.lower()} and what clues help them decide what matters in a text."
    if unit["unit"] == 2:
        return f"Students have written short responses in school and everyday settings. Begin by comparing how the same idea changes when the audience, purpose, or context changes in {title.lower()}."
    return f"Students have used literacy routines across Units 1 and 2. Begin by asking them to identify a reading or writing habit that helped them improve before connecting it to {title.lower()}."


def assessment_lists(unit, lesson):
    asl = [
        "Student self-check against success criteria",
        "Reading/writing reflection or learning log",
        "Exit slip on the next literacy step",
    ]
    afl = [
        "Teacher conference during Moodle task sequence",
        "Observation of strategy use and participation",
        "Descriptive feedback on draft or practice response",
    ]
    aol = []
    title = lesson.get("title") or ""
    if re.search(r"homework|summary|paragraph|news|reflective|writing|graphic|idioms|context|reading", title, re.I):
        aol.append("Completed homework or submitted task where identified in Moodle")
    if unit["unit"] == 3:
        aol.append("Portfolio evidence and process reflection where assigned")
    return "\n".join(asl), "\n".join(afl), "\n".join(aol or ["No formal mark unless the Moodle activity identifies the task as assessment of learning."])


def unit_lessons_for(unit):
    return unit.get("lessons") or []


def all_unit_codes(unit):
    codes = []
    for lesson in unit_lessons_for(unit):
        codes.extend(expectation_for(lesson)["codes"])
    return unique(codes)


def unit_learning_goals(unit):
    intent = UNIT_INTENT[int(unit["unit"])]
    goals = [intent["teacher_intent"]]
    if int(unit["unit"]) == 1:
        goals.extend([
            "Students will read informational, narrative, and graphic texts with increasing independence.",
            "Students will use structures, text features, context clues, and comprehension strategies to explain meaning.",
            "Students will collect evidence for reading growth and later portfolio reflection.",
        ])
    elif int(unit["unit"]) == 2:
        goals.extend([
            "Students will write for different purposes, including summaries, paragraphs, news reports, and reflections.",
            "Students will plan, draft, revise, and edit with attention to audience, purpose, context, and conventions.",
            "Students will use feedback to improve clarity and organization before submission.",
        ])
    elif int(unit["unit"]) == 3:
        goals.extend([
            "Students will strengthen communication and active participation habits.",
            "Students will explain their reading and writing process using examples from the course.",
            "Students will prepare portfolio evidence that documents literacy growth.",
        ])
    else:
        goals.extend([
            "Students will organize reading logs, graphic organizers, reflections, and final product evidence.",
            "Students will complete the portfolio project and final brochure or examination requirements.",
            "Students will revise final work using rubrics, checklists, and teacher feedback.",
        ])
    return goals


def unit_summary(unit):
    intent = UNIT_INTENT[int(unit["unit"])]
    lesson_titles = [lesson.get("title") for lesson in unit_lessons_for(unit)]
    lesson_line = "; ".join(lesson_titles) if lesson_titles else "reading log, portfolio, culminating project, final brochure, and final examination tasks"
    return f"{intent['big_idea']} Lessons/resources include: {lesson_line}."


def unit_outline(unit):
    if int(unit["unit"]) == 4:
        items = [
            "Reading Log/Tracking Sheet: confirm students have ongoing evidence from independent and assigned reading.",
            "Portfolio Project: review organizer, learning journal reflection, rubrics, and final evidence requirements.",
            "Unit 1-3 Graphic Organizer and Reflection Dropboxes: check completion and provide short conferences.",
            "Final Brochure Submission: support final drafting, revision, and rubric-aligned self-assessment.",
            "Final Exam and Submission Dropbox: review expectations, timing, permitted supports, and submission process.",
        ]
        return "\n".join(f"{index + 1}. {item}" for index, item in enumerate(items))
    rows = []
    for lesson in unit_lessons_for(unit):
        expectation = expectation_for(lesson)
        goal = lesson_goals(unit, lesson, expectation)[0]
        rows.append(f"Lesson {lesson['lesson']}: {lesson['title']} - {goal}")
    evaluations = resource_labels((unit.get("unitResources") or {}).get("evaluations") or [])
    if evaluations:
        rows.append("Unit evaluations: " + "; ".join(evaluations))
    return "\n".join(rows)


def unit_technology(unit):
    if int(unit["unit"]) == 4:
        return "Localized Moodle final/culminating section, portfolio files, rubric files, final exam PDF, and submission dropboxes."
    h5p_count = sum(len(lesson.get("handsOn") or []) for lesson in unit_lessons_for(unit))
    ispring_count = sum(len(lesson.get("ispring") or []) for lesson in unit_lessons_for(unit))
    return f"Localized Moodle lesson pages; {ispring_count} iSpring lesson presentation(s); {h5p_count} H5P/practice or consolidation activity record(s); LMS dropboxes for homework/evaluation submission."


def unit_printed(unit):
    labels = []
    if int(unit["unit"]) == 4:
        labels.extend([
            "Reading-Tracking-Sheet.docx",
            "OLC4O-Learning-Graphic-Organizer.pdf",
            "OLC4O-Learning-Journal-Reflection.docx",
            "Learning-Journal-Rubric.pdf",
            "Brochure-Rubric.docx",
            "OLC4O-Portfolio-Project.docx",
            "OLC4O-Exam-Updated.pdf",
        ])
    else:
        for lesson in unit_lessons_for(unit):
            _, files = lesson_materials(lesson)
            labels.extend(files)
    return "\n".join(unique(labels)[:18] or ["No separate printed file identified; use localized Moodle pages and teacher-created note catcher if needed."])


def unique_cells(row):
    seen = set()
    cells = []
    for cell in row.cells:
        marker = id(cell._tc)
        if marker not in seen:
            seen.add(marker)
            cells.append(cell)
    return cells


def set_text(cell, text, bold=False, size=9.5):
    cell.text = ""
    for line_index, line in enumerate(str(text or "").split("\n")):
        paragraph = cell.paragraphs[0] if line_index == 0 else cell.add_paragraph()
        run = paragraph.add_run(line)
        run.bold = bold
        run.font.name = "Calibri"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
        run.font.size = Pt(size)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP


def set_row_text(row, text, bold=False, size=9.5):
    for cell in unique_cells(row):
        set_text(cell, text, bold=bold, size=size)


def template_document(kind):
    return Document(LESSON_PLAN_TEMPLATE if kind == "lesson" else UNIT_PLAN_TEMPLATE)


def style_document(doc):
    section = doc.sections[0]
    section.top_margin = Inches(0.9)
    section.bottom_margin = Inches(0.9)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    normal.font.size = Pt(10)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.08
    for name, size, color in [("Heading 1", 16, "002B57"), ("Heading 2", 12, "002B57"), ("Heading 3", 10.5, "002B57")]:
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)


def lesson_plan_doc(unit, lesson):
    expectation = expectation_for(lesson)
    media, files = lesson_materials(lesson)
    doc = template_document("lesson")
    style_document(doc)

    overview = doc.tables[0]
    rows = overview.rows
    set_text(rows[0].cells[0], "Lesson Plan", bold=True)
    set_text(rows[0].cells[2], f"Subject: {SUBJECT}", bold=True)
    set_row_text(rows[1], f"Lesson Name: {lesson['title']}")
    set_row_text(rows[2], f"Unit of Study: {UNIT_INTENT[int(unit['unit'])]['title']}")
    set_row_text(rows[3], "PRIOR KNOWLEDGE What do my students already know? What key questions or strategies will activate prior knowledge? What connections can I help students make?\n" + prior_knowledge(unit, lesson))

    expectation_text = "\n".join(
        [
            "CURRICULUM EXPECTATIONS",
            "OVERALL",
            *(expectation["overall"] or ["Use the localized Lesson Expectations page for exact wording."]),
            "SPECIFIC",
            *(expectation["specific"] or ["Use the localized Lesson Expectations page for exact wording."]),
        ]
    )
    set_row_text(rows[4], expectation_text)
    set_row_text(rows[5], "\n".join(["LEARNING GOALS What do I want students to know and be able to do?", *lesson_goals(unit, lesson, expectation)]))
    set_row_text(rows[6], "\n".join(["SUCCESS CRITERIA(S) How will students know they attained the goal? What tools can they use to self-assess?", *success_criteria(unit, lesson, expectation)]))

    asl, afl, aol = assessment_lists(unit, lesson)
    set_text(rows[7].cells[0], "Assessment as Learning", bold=True)
    set_text(rows[7].cells[1], "Assessment for Learning", bold=True)
    set_text(rows[7].cells[2], "Assessment of Learning", bold=True)
    set_text(rows[8].cells[0], asl)
    set_text(rows[8].cells[1], afl)
    set_text(rows[8].cells[2], aol)
    set_row_text(rows[9], "What will I do?\nConfer\nObserve\nGive descriptive feedback\nGrade only where the Moodle activity identifies the task as assessment of learning")
    set_row_text(
        rows[10],
        "Accommodations: How will you change the lesson to meet the needs of individual students?\n"
        "Chunk texts into short reading targets; pre-teach vocabulary; provide read-aloud or audio support where available; model one response before independent work; offer graphic organizers and sentence frames; allow oral rehearsal before written submission; extend through a higher-complexity text or revision challenge.",
    )
    set_row_text(
        rows[11],
        "Materials and Resources: What do I need to prepare in advance? Equipment? Student resources? Teacher resources?\n"
        + "\n".join((media + files)[:14] or ["Localized Moodle lesson page and attached homework/evaluation resources."])
        + f"\nSource status: locally authored from OLC4O Moodle content on {GENERATED_AT}.",
    )

    lesson_table = doc.tables[1]
    set_text(lesson_table.rows[0].cells[0], "DELIVERING THE LESSON", bold=True)
    phases = [
        (
            1,
            "Timing\n5-8\nminutes",
            "Minds On!\nOpen with a short reading or writing diagnostic prompt. Name the learning goal and success criteria, then connect the task to OSSLT-style literacy habits.",
            "Materials/Resources\nLesson Expectations page\nQuick-write or KWL prompt",
        ),
        (
            2,
            "Timing\n20-30\nminutes",
            "Action!\nUse the Lesson page and iSpring presentation where present to model the skill. Pause for vocabulary, strategy, and evidence checks before students continue independently.",
            "Materials/Resources\nLocalized lesson page\n" + ("\n".join(media[:5]) if media else "Course activity page"),
        ),
        (
            3,
            "Timing\n15-25\nminutes",
            "Action!\nGuide students through Hands On/H5P/practice work. Require students to explain what strategy they used and what evidence or revision choice supports their answer.",
            "Materials/Resources\nHands On section\nPractice/H5P resources\nTeacher conference notes",
        ),
        (
            4,
            "Timing\n10-15\nminutes",
            "Consolidation!\nUse the consolidation or homework instructions to check understanding. Confirm the correct dropbox and have students record one next step for reading, writing, or portfolio growth.",
            "Materials/Resources\n" + ("\n".join(files[:7]) if files else "Homework/consolidation instructions"),
        ),
    ]
    for row_index, timing, action, materials in phases:
        row = lesson_table.rows[row_index]
        set_text(row.cells[0], timing)
        set_text(row.cells[1], "W\n[ ]")
        set_text(row.cells[2], "S\n[ ]")
        set_text(row.cells[3], "I\n[ ]")
        set_text(row.cells[4], action + "\n\n" + short_text(lesson_activity_summary(lesson), 430))
        set_text(row.cells[5], materials)

    set_text(
        doc.tables[2].rows[0].cells[0],
        "Notes:\nGenerated teacher planning aid reconstructed from localized OLC4O Moodle content. Keep answer-key resources separate in the teacher-facing area. Use these plans for lesson preparation, pacing, differentiation, and observation notes; do not treat them as original Moodle teacher packet documents.",
    )
    return doc


def unit_plan_doc(unit):
    doc = template_document("unit")
    style_document(doc)
    unit_number = int(unit["unit"])
    intent = UNIT_INTENT[unit_number]
    if doc.paragraphs:
        doc.paragraphs[0].text = f"UNIT {unit_number}"
        doc.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    table = doc.tables[0]
    rows = table.rows
    for row_index, label in {0: "Unit Author", 5: "Unit Overview", 14: "Unit Foundation", 19: "Assessment Plan", 22: "Unit Details", 25: "Materials and Resources"}.items():
        set_text(rows[row_index].cells[0], label, bold=True)
    for row_index, label in [(1, "Name:"), (2, "School District:"), (3, "School Name:"), (4, "School City, Province:")]:
        set_row_text(rows[row_index], label)

    codes = all_unit_codes(unit) if unit_number != 4 else []
    strand_titles = []
    for code in codes:
        strand = STRANDS.get(code[0])
        if strand:
            strand_titles.append(f"{code[0]}. {strand['title']} - {strand['focus']}")
    if unit_number == 4:
        strand_titles = [f"C. {STRANDS['C']['title']} - {STRANDS['C']['focus']}"]

    values = [
        (6, "Unit Title Name"),
        (7, intent["title"]),
        (8, "Unit Summary"),
        (9, unit_summary(unit)),
        (10, "Year Level"),
        (11, "Grade 12, Open Ontario Secondary School Literacy Course"),
        (12, "Approximate Time Needed"),
        (13, f"{max(len(unit_lessons_for(unit)) * 3, 4)} instructional hours, plus homework and evaluation time" if unit_number != 4 else "Final portfolio/final examination preparation period, adjusted to teacher pacing and submission windows"),
        (15, "Targeted Curriculum Expectations"),
        (
            16,
            "\n".join(
                [
                    "Expectation evidence is drawn from the localized OLC4O Lesson Expectations pages.",
                    f"Codes found in this unit: {', '.join(codes) if codes else 'see culminating and portfolio requirements.'}",
                    *unique(strand_titles),
                ]
            ),
        ),
        (17, "Learning Goals"),
        (18, "\n".join(unit_learning_goals(unit))),
    ]
    for row_index, value in values:
        set_row_text(rows[row_index], value)

    set_text(rows[20].cells[0], "Assessment as Learning (ASL)", bold=True)
    set_text(rows[20].cells[2], "Assessment for Learning (AFL)", bold=True)
    set_text(rows[20].cells[3], "Assessment of Learning (AoL)", bold=True)
    set_text(rows[21].cells[0], "[x] KWL/reflection routines\n[x] Reading or writing log\n[x] Portfolio self-assessment\n[x] Exit card next-step notes")
    set_text(rows[21].cells[2], "[x] Hands On/H5P checks\n[x] Teacher conferences\n[x] Draft review\n[x] Homework or practice review")
    evaluations = resource_labels((unit.get("unitResources") or {}).get("evaluations") or [])
    if unit_number == 4:
        evaluations = [
            "Reading Log/Tracking Sheet",
            "Portfolio Project",
            "Unit 1-3 Graphic Organizer and Reflection submissions",
            "Final Brochure Submission",
            "Final Exam",
        ]
    set_text(rows[21].cells[3], "[x] " + "\n[x] ".join(evaluations[:10]) if evaluations else "[x] Unit assignment where identified in Moodle")
    set_row_text(rows[23], "Lesson and Assessment Outlines", bold=True)
    set_row_text(rows[24], unit_outline(unit))
    set_text(rows[26].cells[0], "Technology", bold=True)
    set_text(rows[26].cells[1], unit_technology(unit))
    set_text(rows[27].cells[0], "Printed", bold=True)
    set_text(rows[27].cells[1], unit_printed(unit))
    set_text(rows[28].cells[0], "Other Resources", bold=True)
    set_text(rows[28].cells[1], "Teacher conference notes\nStudent reading and writing samples\nPortfolio checklist\nSuccess criteria and rubric language\nOSSLT-style reading and writing prompts selected by the teacher")
    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    return doc


def save_doc(doc, relative_path):
    path = COURSE_ROOT / relative_path
    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(path)
    return path.stat().st_size


def file_record(label, relative_path, category, role):
    path = COURSE_ROOT / relative_path
    rel_posix = relative_path.replace("\\", "/")
    return {
        "label": label,
        "type": "docx",
        "category": category,
        "role": role,
        "path": rel_posix,
        "bytes": path.stat().st_size,
        "source": "reconstructed_from_moodle_content",
        "sourceStatus": "locally_authored_from_course_materials",
        "teacherUse": "Teacher planning aid reconstructed from localized OLC4O course content; not an original Moodle teacher packet.",
        "templateReference": "MDM4U teacher plan format and local mdm4u-style plan templates.",
        "previewPath": f"previews-html/{rel_posix}.html",
    }


def culminating_unit(manifest):
    return {
        "unit": 4,
        "title": UNIT_INTENT[4]["title"],
        "lessons": [],
        "unitResources": {
            "evaluations": [
                item
                for item in manifest.get("evaluations") or []
                if re.search(r"culminating|portfolio|brochure|final|reading log|graphic organizer", item.get("label") or "", re.I)
            ]
        },
    }


def main():
    manifest = json.loads(MANIFEST_PATH.read_text(encoding="utf-8"))
    unit_count = 0
    lesson_count = 0
    written = []

    for unit in manifest.get("units") or []:
        unit_number = int(unit["unit"])
        unit["title"] = UNIT_INTENT[unit_number]["title"]
        unit_rel = f"plans/generated/unit-plans/U{unit_number:02d}-unit-plan.docx"
        save_doc(unit_plan_doc(unit), unit_rel)
        unit["unitPlan"] = file_record(f"Unit Plan - {UNIT_INTENT[unit_number]['title']}", unit_rel, "unit_plan", "unit_plan")
        unit_count += 1
        written.append(unit_rel)

        for lesson in unit.get("lessons") or []:
            lesson_rel = f"plans/generated/lesson-plans/{lesson['id']}-lesson-plan.docx"
            save_doc(lesson_plan_doc(unit, lesson), lesson_rel)
            lesson["lessonPlan"] = file_record(
                f"Lesson Plan - Unit {unit_number} Lesson {lesson['lesson']}",
                lesson_rel,
                "lesson_plan",
                "lesson_plan",
            )
            lesson["resourceCounts"] = {**(lesson.get("resourceCounts") or {}), "lessonPlan": 1}
            lesson_count += 1
            written.append(lesson_rel)

    u4 = culminating_unit(manifest)
    u4_rel = "plans/generated/unit-plans/U04-unit-plan.docx"
    save_doc(unit_plan_doc(u4), u4_rel)
    u4_record = file_record("Unit Plan - Portfolio, Culminating, and Final Examination", u4_rel, "unit_plan", "unit_plan")
    written.append(u4_rel)
    unit_count += 1

    for section in manifest.get("courseSections") or []:
        if section.get("role") == "final_examination_culminating":
            section["unitPlan"] = u4_record
            section["planningStatus"] = "teacher_plan_generated"

    manifest["generatedAt"] = GENERATED_AT
    manifest["sourceAudit"] = {
        **(manifest.get("sourceAudit") or {}),
        "teacherPlanGeneration": {
            "generatedAt": GENERATED_AT,
            "format": "docx",
            "unitPlans": unit_count,
            "lessonPlans": lesson_count,
            "source": "reconstructed_from_moodle_content",
            "basis": [
                "localized OLC4O Lesson Expectations pages",
                "localized Moodle Lesson, Hands On, Consolidation, and Homework sections",
                "localized iSpring, H5P, worksheet, assignment, portfolio, final examination, and rubric records",
                "MDM4U-style teacher plan templates and existing generated course plan examples",
            ],
            "boundary": "Teacher planning aids only; no answer keys, rubrics, textbook excerpts, or Moodle-original teacher packet documents were created.",
            "culminatingPlanPlacement": "Stored on courseSections.final_examination_culminating as a section-level unit plan, not as a nested Teacher Packet card.",
        },
    }

    MANIFEST_PATH.write_text(json.dumps(manifest, indent=2, ensure_ascii=False) + "\n", encoding="utf-8")
    REPORT_PATH.parent.mkdir(parents=True, exist_ok=True)
    REPORT_PATH.write_text(
        json.dumps(
            {"course": "OLC4O", "unitPlans": unit_count, "lessonPlans": lesson_count, "written": written, "generatedAt": GENERATED_AT},
            indent=2,
            ensure_ascii=False,
        )
        + "\n",
        encoding="utf-8",
    )
    print(json.dumps({"course": "OLC4O", "unitPlans": unit_count, "lessonPlans": lesson_count, "generatedAt": GENERATED_AT}, indent=2))


if __name__ == "__main__":
    main()
