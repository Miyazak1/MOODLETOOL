import json
import re
from datetime import datetime, timezone
from html import unescape
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


WORKSPACE_ROOT = Path("D:/工作文件/SUNNYBROOK")
REPO_ROOT = WORKSPACE_ROOT / "ossd-course-portal"
COURSE_ROOT = WORKSPACE_ROOT / "courseware" / "SNC2D"
MANIFEST_PATH = COURSE_ROOT / "course-manifest.json"
REPORT_PATH = REPO_ROOT / "deployment" / "SNC2D-docx-plans-report.json"
GENERATED_AT = datetime.now(timezone.utc).replace(microsecond=0).isoformat()
TEMPLATE_ROOT = REPO_ROOT / "templates" / "teacher-plans"
LESSON_PLAN_TEMPLATE = TEMPLATE_ROOT / "mdm4u-style-lesson-plan-template.docx"
UNIT_PLAN_TEMPLATE = TEMPLATE_ROOT / "mdm4u-style-unit-plan-template.docx"
MDM4U_LESSON_PLAN_REFERENCE = (
    WORKSPACE_ROOT / "courseware" / "MDM4U" / "plans" / "source" / "Unit 1" / "Lesson Plan Unit 1 lesson 1.docx"
)
MDM4U_UNIT_PLAN_REFERENCE = WORKSPACE_ROOT / "courseware" / "MDM4U" / "plans" / "source" / "Unit 1" / "Unit 1 Plan .docx"


STRANDS = {
    1: {
        "code": "B",
        "title": "Biology: Tissues, Organs, and Systems of Living Things",
        "focus": "cell specialization, tissues, organs, body systems, microscopy, homeostasis, and related health or technology applications.",
        "essential": [
            "How do specialized cells work together to form tissues, organs, and body systems?",
            "How can evidence from models, microscopes, and investigations explain body-system function?",
            "How do choices, technologies, and diseases affect human body systems?",
        ],
    },
    2: {
        "code": "C",
        "title": "Chemistry: Chemical Reactions",
        "focus": "atomic structure review, ionic and molecular compounds, chemical nomenclature, reaction types, balancing equations, acids/bases, and chemical safety.",
        "essential": [
            "How do atoms, ions, and molecules account for the properties and reactions of substances?",
            "How can formulas, names, and equations communicate chemical change clearly?",
            "How do chemical reactions affect everyday life, health, technology, and the environment?",
        ],
    },
    3: {
        "code": "E",
        "title": "Physics: Light and Geometric Optics",
        "focus": "the electromagnetic spectrum, visible light, reflection, refraction, lenses, optical technologies, and vision.",
        "essential": [
            "How does light behave as it travels, reflects, refracts, and forms images?",
            "How can ray diagrams and models predict what optical devices do?",
            "How do optical technologies extend human ability to observe and communicate?",
        ],
    },
    4: {
        "code": "D",
        "title": "Earth and Space Science: Climate Change",
        "focus": "weather, climate systems, greenhouse gases, climate evidence, feedbacks, and social, economic, and political responses.",
        "essential": [
            "How do natural systems and human activities influence climate?",
            "How can climate evidence be interpreted, evaluated, and communicated responsibly?",
            "What scientific, social, economic, and political choices shape climate responses?",
        ],
    },
}


def clean_text(value):
    text = str(value or "")
    text = re.sub(r"<script[\s\S]*?</script>", " ", text, flags=re.I)
    text = re.sub(r"<style[\s\S]*?</style>", " ", text, flags=re.I)
    text = re.sub(r"<[^>]+>", " ", text)
    text = unescape(text)
    return re.sub(r"\s+", " ", text).strip()


def read_local_text(relative_path):
    if not relative_path:
        return ""
    path = COURSE_ROOT / relative_path
    if not path.exists():
        return ""
    return clean_text(path.read_text(encoding="utf-8", errors="ignore"))


def short_text(value, limit=320):
    text = clean_text(value)
    if len(text) <= limit:
        return text
    clipped = text[:limit]
    cut = max(clipped.rfind(". "), clipped.rfind("; "), clipped.rfind(", "))
    if cut > 150:
        clipped = clipped[: cut + 1]
    return clipped.strip() + "..."


def find_section(lesson, pattern):
    regex = re.compile(pattern, re.I)
    for section in lesson.get("bookSections") or []:
        label = section.get("sectionLabel") or section.get("label") or ""
        if regex.search(label):
            return section
    return None


def section_text(lesson, pattern):
    section = find_section(lesson, pattern)
    if not section:
        return ""
    return read_local_text(section.get("path")) or clean_text(section.get("textPreview"))


def expectation_codes(text):
    return sorted(set(re.findall(r"\b[A-E][1-3](?:\.\d+)?\b", text or "")))[:18]


def split_sentences(text, limit):
    text = clean_text(text)
    if not text:
        return []
    parts = re.split(r"(?<=[.;:])\s+|(?=\b[A-E][1-3](?:\.\d+)?\b)", text)
    return [short_text(part, 260) for part in parts if len(clean_text(part)) > 12][:limit] or [short_text(text, 260)]


def extract_expectation_sections(text):
    text = clean_text(text)
    if not text:
        return {
            "overall": ["Use the localized Lesson Expectations page for exact wording."],
            "specific": ["Use the localized Lesson Expectations page for exact wording."],
        }
    lower = text.lower()
    overall = text
    specific = ""
    for marker in ["specific lesson expectations", "specific expectations", "specific expectation"]:
        pos = lower.find(marker)
        if pos >= 0:
            overall = text[:pos]
            specific = text[pos:]
            break
    for marker in ["learning goals", "success criteria", "lesson "]:
        pos = overall.lower().find(marker)
        if pos > 25:
            overall = overall[:pos]
    for marker in ["learning goals", "success criteria"]:
        pos = specific.lower().find(marker)
        if pos >= 0:
            specific = specific[:pos]
    return {"overall": split_sentences(overall, 6), "specific": split_sentences(specific or text, 10)}


def learning_goals(expectations_text, lesson_title):
    match = re.search(
        r"Learning Goals?:\s*(.+?)(Success Criteria|Lesson|Hands On|Consolidation|Homework|$)",
        expectations_text or "",
        flags=re.I,
    )
    if match:
        raw = match.group(1)
        parts = re.split(
            r"(?=Learn|Explain|Describe|Identify|Compare|Analyse|Analyze|Understand|Investigate|Evaluate|Use|Create|Demonstrate|Predict|Classify)",
            raw,
        )
        goals = [short_text(part, 150) for part in parts if len(clean_text(part)) > 8]
        if goals:
            return goals[:4]
    return [
        f"Explain the central science concept in {lesson_title} using accurate SNC2D vocabulary.",
        "Use observations, models, diagrams, data, or examples from the lesson to support a scientific explanation.",
        "Apply the concept through the Hands On activity, consolidation evidence, and homework task.",
    ]


def lesson_success_criteria(lesson_title, codes):
    code_text = ", ".join(codes[:4]) if codes else "the listed lesson expectations"
    return [
        f"I can describe the main idea in {lesson_title} using correct science vocabulary.",
        f"I can connect my explanation to {code_text} and use lesson evidence to support it.",
        "I can interpret a model, diagram, video, data set, or worked example from the lesson.",
        "I can use feedback from Hands On or consolidation work to improve my homework response.",
    ]


def all_resources(lesson):
    return [item for item in (lesson.get("downloads") or []) + (lesson.get("textExports") or []) if item]


def resources_matching(lesson, pattern):
    regex = re.compile(pattern, re.I)
    found = []
    for item in all_resources(lesson):
        haystack = " ".join(str(item.get(key, "")) for key in ("label", "type", "category", "role", "path", "previewPath"))
        if regex.search(haystack):
            found.append(item)
    return found


def resource_labels(items):
    return [item.get("label") or item.get("title") or "Untitled resource" for item in items]


def role_label(value):
    label = str(value or "course").replace("_", " ").strip()
    return " ".join(part.capitalize() for part in label.split())


def lesson_materials(lesson):
    media = []
    for item in lesson.get("ispring") or []:
        media.append(f"iSpring module: {item.get('label') or 'lesson presentation'}")
    for item in resources_matching(lesson, r"h5p"):
        media.append(f"{role_label(item.get('role'))} H5P: {item.get('label')}")
    for item in resources_matching(lesson, r"mp4|video"):
        media.append(f"{role_label(item.get('role'))} video: {item.get('label')}")
    files = resource_labels(resources_matching(lesson, r"docx?|doc\b|pdf|pptx?|xlsx?|txt|png|jpe?g|gif|tiff?"))[:10]
    return media, files


def activity_titles(unit, pattern):
    regex = re.compile(pattern, re.I)
    titles = []
    for lesson in unit.get("lessons") or []:
        for item in (lesson.get("downloads") or []) + (lesson.get("textExports") or []) + (lesson.get("ispring") or []):
            haystack = " ".join(str(item.get(key, "")) for key in ("label", "type", "category", "role", "section", "path"))
            if regex.search(haystack):
                titles.append(item.get("label") or item.get("title") or "Untitled")
    return titles


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_width(cell, width_dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_text(cell, text, bold=False, size=9.5):
    cell.text = ""
    lines = str(text or "").split("\n")
    for line_index, line in enumerate(lines):
        paragraph = cell.paragraphs[0] if line_index == 0 else cell.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(0)
        run = paragraph.add_run(line)
        run.bold = bold
        run.font.name = "Calibri"
        run.font.size = Pt(size)


def set_row_text(row, text, bold=False, size=9.5):
    set_text(row.cells[0], text, bold=bold, size=size)


def template_document(kind):
    if kind == "lesson":
        path = LESSON_PLAN_TEMPLATE if LESSON_PLAN_TEMPLATE.exists() else MDM4U_LESSON_PLAN_REFERENCE
    elif kind == "unit":
        path = UNIT_PLAN_TEMPLATE if UNIT_PLAN_TEMPLATE.exists() else MDM4U_UNIT_PLAN_REFERENCE
    else:
        raise ValueError(f"Unknown plan template kind: {kind}")
    return Document(str(path))


def style_document(doc):
    section = doc.sections[0]
    section.top_margin = Inches(0.9)
    section.bottom_margin = Inches(0.9)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.12
    for name, size, color in [("Heading 1", 16, "002B57"), ("Heading 2", 12.5, "002B57"), ("Heading 3", 10.5, "002B57")]:
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)


def unit_summary(unit, strand, lessons):
    topics = ", ".join((lesson.get("title") or "") for lesson in lessons[:5])
    return (
        f"This unit develops the {strand['title']} strand through {len(lessons)} linked SNC2D lessons. "
        f"Students move from lesson expectations and guided media into Hands On practice, consolidation evidence, and homework submission. "
        f"Early topics include {topics}."
    )


def lesson_plan_doc(unit, lesson):
    strand = STRANDS.get(unit["unit"], {"code": "", "title": unit["title"], "focus": unit["title"], "essential": []})
    expectations = section_text(lesson, r"expectations")
    lesson_page = section_text(lesson, r"^Lesson$")
    codes = expectation_codes(expectations)
    media, files = lesson_materials(lesson)
    expectation_sections = extract_expectation_sections(expectations)
    goals = learning_goals(expectations or lesson_page, lesson["title"])
    criteria = lesson_success_criteria(lesson["title"], codes)
    doc = template_document("lesson")
    style_document(doc)

    overview = doc.tables[0]
    rows = overview.rows
    set_text(rows[0].cells[0], "Lesson Plan", bold=True)
    set_text(rows[0].cells[2], "Subject: SNC2D Science, Grade 10 Academic", bold=True)
    set_row_text(rows[1], f"Lesson Name: {lesson['title']}")
    set_row_text(rows[2], f"Unit of Study: {unit['title']}")
    prior = (
        f"Students should activate vocabulary and process knowledge from previous SNC2D work in {unit['title']}. "
        f"Begin by linking the previous lesson, KWL chart, or consolidation task to {lesson['title']} so students can name what they already know and what evidence they need next."
    )
    set_row_text(rows[3], "PRIOR KNOWLEDGE What do my students already know? What key questions or strategies will activate prior knowledge? What connections can I help students make?\n" + prior)
    expectation_text = "\n".join(["CURRICULUM EXPECTATIONS", "OVERALL", *expectation_sections["overall"], "SPECIFIC", *expectation_sections["specific"]])
    set_row_text(rows[4], expectation_text)
    set_row_text(rows[5], "\n".join(["LEARNING GOALS", "LEARNING GOALS What do I want students to know and be able to do?", *goals]))
    set_row_text(rows[6], "\n".join(["SUCCESS CRITERIA(S) How will students know they attained the goal? What tools can they use to self-assess?", *criteria]))
    set_text(rows[7].cells[0], "Assessment as Learning", bold=True)
    set_text(rows[7].cells[1], "Assessment for Learning", bold=True)
    set_text(rows[7].cells[2], "Assessment of Learning", bold=True)
    set_text(rows[8].cells[0], "☑ Observation\n☑ Anecdotal Notes\n☑ Exit Card\n☑ Self-Assessment checklist\n☐ Discussions")
    set_text(rows[8].cells[1], "☑ Worksheets\n☑ Homework\n☑ Strategic Questioning\n☑ Hands On/H5P checks\n☑ Consolidation response")
    set_text(rows[8].cells[2], "☑ Homework submission\n☑ Unit quiz/test/lab/assignment where listed\n☐ Oral Presentations\n☐ Portfolio")
    set_row_text(rows[9], "What will I do?\nConfer\nObserve\nGrade")
    set_row_text(
        rows[10],
        "Accommodations: How will you change the lesson to meet the needs of individual students?\n"
        "Chunk the lesson media and text into checkpoints; provide vocabulary support, diagrams, labelled models, formula or sentence frames where useful; allow oral rehearsal before written responses; and use Hands On/consolidation evidence to reteach one misconception.",
    )
    set_row_text(
        rows[11],
        "Materials and Resources: What do I need to prepare in advance? Equipment? Student resources? Teacher resources?\n"
        + "\n".join((media + files)[:10] or ["Localized lesson page, Lesson Expectations, Hands On, Consolidation, and Homework resources."])
        + f"\nSource status: reconstructed_from_moodle_content. Generated from localized SNC2D Moodle resources on {GENERATED_AT}.",
    )

    lesson_table = doc.tables[1]
    set_text(lesson_table.rows[0].cells[0], "DELIVERING THE LESSON", bold=True)
    lesson_rows = lesson_table.rows
    set_text(lesson_rows[1].cells[0], "Timing\n5-8\nminutes")
    set_text(lesson_rows[1].cells[1], "Grouping")
    set_text(
        lesson_rows[1].cells[4],
        "Minds On!\nOpen with a retrieval question, KWL prompt, or quick model/diagram interpretation tied to the previous lesson. Name the learning target and vocabulary students need before media viewing.",
    )
    set_text(lesson_rows[1].cells[5], "Materials/Resources\nKWL chart\nLesson Expectations page")
    phases = [
        (
            2,
            "Timing\n20-30\nminutes",
            "Action!\nUse the Lesson page and iSpring module to model the core idea. Pause after each major concept for a short check using a diagram, equation, observation, or example.",
            "Materials/Resources\nLocalized lesson page\niSpring module",
        ),
        (
            3,
            "Timing\n15-25\nminutes",
            "Action!\nFrame the Hands On activity or H5P as formative practice. Circulate, name misconceptions, and connect the task to the success criteria.",
            "Materials/Resources\n" + ("\n".join(media[:4]) if media else "Hands On page and activity resources"),
        ),
        (
            4,
            "Timing\n10-15\nminutes",
            "Consolidation!\nUse the video, exit slip, or consolidation prompt to check whether students can explain the concept without copying. Confirm homework files and submission expectations.",
            "Materials/Resources\n" + ("\n".join(files[:6]) if files else "Homework and consolidation resources"),
        ),
    ]
    for row_index, timing, action, materials in phases:
        row = lesson_rows[row_index]
        set_text(row.cells[0], timing)
        set_text(row.cells[1], "W\n☐")
        set_text(row.cells[2], "S\n☐")
        set_text(row.cells[3], "I\n☐")
        set_text(row.cells[4], action)
        set_text(row.cells[5], materials)

    notes = doc.tables[2]
    set_text(
        notes.rows[0].cells[0],
        "Notes:\nSupport: pre-teach vocabulary, chunk media, provide labelled diagrams/models or sentence frames, and allow oral rehearsal before written explanation.\n"
        "Extension: connect the lesson concept to a health, environmental, technological, or societal application where the expectations support it.\n"
        "Teacher reflection: identify the misconception that should launch the next lesson.",
    )
    return doc


def unit_plan_doc(unit):
    strand = STRANDS.get(unit["unit"], {"code": "", "title": unit["title"], "focus": unit["title"], "essential": []})
    lessons = unit.get("lessons") or []
    all_codes = []
    outlines = []
    for lesson in lessons:
        expectations = section_text(lesson, r"expectations")
        all_codes.extend(expectation_codes(expectations))
        media, files = lesson_materials(lesson)
        task = "; ".join((media[:2] or files[:2] or ["Lesson page, Hands On, consolidation, and homework resources"]))
        focus = short_text(expectations, 220) or f"Develop core understanding for {lesson['title']}."
        outlines.append(f"{lesson['id']} - {lesson['title']}\nFocus: {focus}\nAssessment: Hands On/H5P checks, consolidation evidence, and homework review.\nResources: {task}")
    codes = sorted(set(all_codes))[:24]
    evaluations = resource_labels((unit.get("unitResources") or {}).get("evaluations") or [])

    doc = template_document("unit")
    style_document(doc)
    if doc.paragraphs:
        doc.paragraphs[0].text = f"UNIT {unit['unit']}"
        doc.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    table = doc.tables[0]
    rows = table.rows
    for row_index, label in {0: "Unit Author", 5: "Unit Overview", 14: "Unit Foundation", 19: "Assessment Plan", 22: "Unit Details", 25: "Materials and Resources"}.items():
        set_text(rows[row_index].cells[0], label, bold=True)
    for row_index, label in [(1, "Name:"), (2, "School District:"), (3, "School Name:"), (4, "School City, Province:")]:
        set_row_text(rows[row_index], label)
    values = [
        (6, "Unit Title Name"),
        (7, unit["title"]),
        (8, "Unit Summary"),
        (9, unit_summary(unit, strand, lessons)),
        (10, "Year Level"),
        (11, "Grade 10, Academic Science"),
        (12, "Approximate Time Needed"),
        (13, f"{max(len(lessons) * 3, 6)} instructional hours, plus homework and evaluation time"),
        (15, "Targeted Curriculum Expectations"),
        (
            16,
            "\n".join(
                [
                    f"{strand['code']} - {strand['title']}",
                    f"Expectation evidence found across lessons: {', '.join(codes) if codes else 'see each Lesson Expectations page.'}",
                    f"Unit focus: {strand['focus']}",
                ]
            ),
        ),
        (17, "Learning Goals"),
        (
            18,
            "\n".join(
                [
                    f"Students will explain and apply major concepts in {unit['title']} using accurate SNC2D science vocabulary.",
                    "Students will use observations, diagrams, models, data, videos, and investigations to support scientific explanations.",
                    "Students will complete Hands On, consolidation, homework, reflection, and evaluation tasks as evidence of learning.",
                    "Essential Questions",
                    *(strand.get("essential") or []),
                ]
            ),
        ),
    ]
    for row_index, value in values:
        set_row_text(rows[row_index], value)
    set_text(rows[20].cells[0], "Assessment as Learning (ASL)", bold=True)
    set_text(rows[20].cells[2], "Assessment for Learning (AFL)", bold=True)
    set_text(rows[20].cells[3], "Assessment of Learning (AoL)", bold=True)
    set_text(rows[21].cells[0], "☑ KWL chart\n☑ Reflection routines\n☑ Learning Log\n☑ Student next-step notes")
    set_text(rows[21].cells[2], "☑ Hands On/H5P checks\n☑ Strategic questioning\n☑ Homework review\n☑ Consolidation responses")
    set_text(rows[21].cells[3], "☑ Unit quiz/test/lab/assignment where listed\n" + ("☑ " + "\n☑ ".join(evaluations[:10]) if evaluations else "☑ Unit quiz, test, lab, or culminating activity listed in the Evaluation section."))
    set_row_text(rows[23], "Lesson and Assessment Outlines", bold=True)
    set_row_text(rows[24], "\n\n".join(outlines))
    set_text(rows[26].cells[0], "Technology", bold=True)
    tech = "\n".join(sorted(set(activity_titles(unit, r"ispring|h5p|mp4|video")))[:24] or ["Localized lesson pages, iSpring presentations, H5P activities, and consolidation videos listed in the manifest."])
    set_text(rows[26].cells[1], tech)
    set_text(rows[27].cells[0], "Printed", bold=True)
    printed = "\n".join(sorted(set(activity_titles(unit, r"docx?|doc\b|pdf|worksheet|homework|handout|rubric|lab")))[:24] or ["Homework, lab, rubric, and activity files listed in the manifest."])
    set_text(rows[27].cells[1], printed)
    set_text(rows[28].cells[0], "Other Resources", bold=True)
    set_text(rows[28].cells[1], "Anecdotal notes of observation\nExit slips and consolidation responses\nStudent self-assessment and reflection\nKWL chart and Learning Log routines where present")
    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    return doc


def save_doc(doc, relative_path):
    path = COURSE_ROOT / relative_path
    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(path)
    return path.stat().st_size


def file_record(label, relative_path, category, role):
    path = COURSE_ROOT / relative_path
    preview = f"previews-html/{relative_path}.html"
    return {
        "label": label,
        "type": "docx",
        "category": category,
        "role": role,
        "path": relative_path.replace("\\", "/"),
        "bytes": path.stat().st_size,
        "source": "reconstructed_from_moodle_content",
        "sourceStatus": "locally_authored_from_course_materials",
        "teacherUse": "Teacher planning aid reconstructed from localized SNC2D course content; not an original Moodle teacher packet.",
        "templateReference": "MDM4U/ENG3U teacher plan format and local mdm4u-style templates.",
        "previewPath": preview.replace("\\", "/"),
    }


def main():
    manifest = json.loads(MANIFEST_PATH.read_text(encoding="utf-8"))
    unit_count = 0
    lesson_count = 0
    written = []
    for unit in manifest.get("units") or []:
        unit_rel = f"plans/generated/unit-plans/U{unit['unit']:02d}-unit-plan.docx"
        save_doc(unit_plan_doc(unit), unit_rel)
        unit["unitPlan"] = file_record(f"Unit Plan - {unit['title']}", unit_rel, "unit_plan", "unit_plan")
        unit_count += 1
        written.append(unit_rel)
        for lesson in unit.get("lessons") or []:
            lesson_rel = f"plans/generated/lesson-plans/{lesson['id']}-lesson-plan.docx"
            save_doc(lesson_plan_doc(unit, lesson), lesson_rel)
            lesson["lessonPlan"] = file_record(
                f"Lesson Plan - Unit {unit['unit']} Lesson {lesson['lesson']}",
                lesson_rel,
                "lesson_plan",
                "lesson_plan",
            )
            lesson["resourceCounts"] = {**(lesson.get("resourceCounts") or {}), "lessonPlan": 1}
            lesson_count += 1
            written.append(lesson_rel)

    manifest["generatedAt"] = GENERATED_AT
    manifest["sourceAudit"] = {
        **(manifest.get("sourceAudit") or {}),
        "teacherPlanGeneration": {
            "generatedAt": GENERATED_AT,
            "format": "docx",
            "unitPlans": unit_count,
            "lessonPlans": lesson_count,
            "source": "reconstructed_from_moodle_content",
            "basis": [
                "localized Moodle Lesson Expectations pages",
                "localized Moodle Lesson, Hands On, Consolidation, and Homework sections",
                "localized iSpring, H5P, video, worksheet, exit slip, KWL/reflection, lab, assignment, and unit Evaluation records",
                "SNC2D strand and expectation evidence already present in the localized course pages",
                "MDM4U/ENG3U teacher plan formatting conventions and local mdm4u-style plan templates",
            ],
            "boundary": "Teacher planning aids only; no answer keys, rubrics, textbook excerpts, or Moodle-original teacher packet documents were created.",
        },
    }
    MANIFEST_PATH.write_text(json.dumps(manifest, indent=2, ensure_ascii=False) + "\n", encoding="utf-8")
    REPORT_PATH.parent.mkdir(parents=True, exist_ok=True)
    REPORT_PATH.write_text(
        json.dumps({"course": "SNC2D", "unitPlans": unit_count, "lessonPlans": lesson_count, "written": written, "generatedAt": GENERATED_AT}, indent=2, ensure_ascii=False) + "\n",
        encoding="utf-8",
    )
    print(json.dumps({"course": "SNC2D", "unitPlans": unit_count, "lessonPlans": lesson_count, "generatedAt": GENERATED_AT}, indent=2))


if __name__ == "__main__":
    main()
