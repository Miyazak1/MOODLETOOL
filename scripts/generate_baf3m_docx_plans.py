import html
import json
import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENTATION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


REPO_ROOT = Path(__file__).resolve().parents[1]
COURSE_ROOT = (REPO_ROOT.parent / "courseware" / "BAF3M").resolve()
MANIFEST_PATH = COURSE_ROOT / "course-manifest.json"
REPORT_PATH = REPO_ROOT / "deployment" / "BAF3M-docx-plans-report.json"

GENERATED_SOURCE = (
    "Generated from St.Mary Moodle book sections, Moodle activities, local course files, "
    "Ontario Business Studies curriculum, and public BAF3M course sequencing references."
)
GENERATED_NOTE = (
    "Original Moodle export did not include a separate unit plan or lesson plan file for this item. "
    "This DOCX follows the local Unit Plan / Lesson Plan template used by existing courses and is "
    "reconstructed from indexed course content."
)

CURRICULUM = {
    1: {
        "strand": "A. Fundamental Accounting Practices",
        "overall": [
            "Describe the discipline of accounting and its importance for business.",
            "Describe the differences among the various forms of business organization.",
            "Demonstrate an understanding of the basic procedures and principles of the accounting cycle for a service business.",
        ],
    },
    2: {
        "strand": "B. Advanced Accounting Practices",
        "overall": [
            "Demonstrate an understanding of the procedures and principles of the accounting cycle for a merchandising business.",
            "Demonstrate an understanding of the accounting practices for sales tax.",
            "Apply accounting practices in a computerized environment.",
        ],
    },
    3: {
        "strand": "C. Internal Control, Financial Analysis, and Decision Making",
        "overall": [
            "Demonstrate an understanding of internal control procedures in the financial management of a business.",
            "Evaluate the financial status of a business by analysing performance measures and financial statements.",
            "Explain how accounting information is used in decision making.",
        ],
    },
    4: {
        "strand": "D. Ethics, Impact of Technology, and Careers",
        "overall": [
            "Assess the role of ethics in, and the impact of current issues on, the practice of accounting.",
            "Assess the impact of technology on the accounting functions in business.",
            "Describe professional accounting designations and career opportunities.",
        ],
    },
}

UNIT_PEDAGOGY = {
    1: {
        "big_idea": "Accounting is a decision-making language. Students move from the purpose and users of accounting to the accounting cycle for a service business.",
        "sequence": "The unit begins with vocabulary and purpose, then builds toward GAAP, forms of business ownership, transaction analysis, trial balances, financial statements, and year-end procedures.",
        "teacher_intent": "Teach students to slow down before recording: identify the business event, decide which accounts are affected, apply the appropriate principle, and check that the accounting equation remains balanced.",
        "goals": [
            "Students will explain why businesses keep accounting records and who uses the information.",
            "Students will apply GAAP and business-organization concepts when interpreting basic accounting situations.",
            "Students will analyse service-business transactions and connect journal/ledger work to the accounting equation.",
            "Students will prepare and interpret a trial balance and basic financial statements for a service business.",
            "Students will describe year-end procedures and explain how closing entries prepare records for the next fiscal period.",
        ],
    },
    2: {
        "big_idea": "Merchandising businesses require additional records because inventory, cost of goods sold, sales tax, and repeated transactions change the accounting cycle.",
        "sequence": "The unit extends service-business accounting into inventory systems, merchandising entries, closing procedures, subsidiary ledgers, special journals, sales tax, and accounting software.",
        "teacher_intent": "Help students see what changes when a business sells goods: inventory must be tracked, sales and purchase activity must be summarized efficiently, and tax/software systems must still produce reliable financial statements.",
        "goals": [
            "Students will identify the accounting elements unique to a merchandising business.",
            "Students will record merchandising transactions using appropriate journals, ledgers, and inventory methods.",
            "Students will prepare closing entries and financial statements that include cost of goods sold.",
            "Students will explain how subsidiary ledgers, special journals, and sales tax accounts support accurate reporting.",
            "Students will connect manual accounting procedures to computerized accounting software and reporting outputs.",
        ],
    },
    3: {
        "big_idea": "Financial information becomes useful when it is controlled, budgeted, analysed, and used to support decisions.",
        "sequence": "The unit starts with internal control and budgeting, then shifts to auditors, current assets and liabilities, liquidity/solvency, profitability, and financial decision making.",
        "teacher_intent": "Teach students to move beyond recording transactions into asking whether information is reliable, what the numbers reveal, and what decision should follow from the evidence.",
        "goals": [
            "Students will explain why internal controls and auditors increase the reliability of accounting information.",
            "Students will use budgeted financial statements to support planning and evaluation.",
            "Students will classify current assets and liabilities and interpret their effect on financial position.",
            "Students will calculate and interpret liquidity, solvency, and profitability measures.",
            "Students will use financial analysis to make and justify business decisions for internal and external users.",
        ],
    },
    4: {
        "big_idea": "Accounting work is shaped by ethical responsibility, technology, current issues, and professional pathways.",
        "sequence": "The unit connects ethics and current issues to technology, compares manual and computerized systems, and finishes with accounting careers and designations.",
        "teacher_intent": "Help students understand that accountants do more than process numbers: they protect trust, adapt to technology, communicate with stakeholders, and make professional judgements.",
        "goals": [
            "Students will apply ethical standards to accounting situations involving pressure, accuracy, confidentiality, and fraud risk.",
            "Students will describe current issues that affect accounting practice and stakeholder trust.",
            "Students will compare manual and computerized accounting systems, including benefits, risks, and controls.",
            "Students will assess how technology affects accounting information, workflow, security, and professional judgement.",
            "Students will investigate accounting designations and career pathways connected to course skills.",
        ],
    },
}


def main():
    if not MANIFEST_PATH.exists():
        raise FileNotFoundError(MANIFEST_PATH)
    manifest = json.loads(MANIFEST_PATH.read_text(encoding="utf-8"))
    generated = {"unitPlans": 0, "lessonPlans": 0, "missingExpectationPages": [], "written": []}

    for unit in manifest.get("units", []):
        unit_no = int(unit["unit"])
        unit_rel = f"plans/source/Unit {unit_no}/BAF3M - Unit {unit_no} - Unit Plan.docx"
        unit_preview_rel = f"previews-html/{unit_rel}.html"
        make_unit_docx(manifest, unit, unit_rel)
        generated["written"].append(unit_rel)
        unit["unitPlan"] = plan_resource(
            label=f"Unit Plan - {unit['title']}",
            role="unit_plan",
            category="unit_plan",
            rel_path=unit_rel,
            preview_rel=unit_preview_rel,
        )
        generated["unitPlans"] += 1

        for lesson in unit.get("lessons", []):
            lesson_rel = (
                f"plans/source/Unit {unit_no}/"
                f"BAF3M - Unit {unit_no} - Lesson {lesson['lesson']} Lesson Plan.docx"
            )
            lesson_preview_rel = f"previews-html/{lesson_rel}.html"
            exp = expectation_for(lesson)
            if not exp["exists"]:
                generated["missingExpectationPages"].append(lesson["id"])
            make_lesson_docx(manifest, unit, lesson, exp, lesson_rel)
            generated["written"].append(lesson_rel)
            lesson["lessonPlan"] = plan_resource(
                label=f"Lesson Plan - Unit {unit_no} Lesson {lesson['lesson']}",
                role="lesson_plan",
                category="lesson_plan",
                rel_path=lesson_rel,
                preview_rel=lesson_preview_rel,
            )
            generated["lessonPlans"] += 1

    manifest["sourceAudit"] = {
        **manifest.get("sourceAudit", {}),
        "generatedPlans": {
            "sourceStatus": "reconstructed_from_moodle_content",
            "source": GENERATED_SOURCE,
            "sourceNote": GENERATED_NOTE,
            "templateReference": "MDM4U and BOH4M DOCX Unit Plan / Lesson Plan table templates.",
            "curriculumSource": (
                "The Ontario Curriculum, Grades 11 and 12: Business Studies, 2006 (Revised), "
                "BAF3M Financial Accounting Fundamentals."
            ),
            "externalCourseReference": (
                "TVO Learn BAF3M public course overview and lesson sequence were used only to "
                "check accounting topic coverage; local Moodle content remains the primary source."
            ),
            "counts": {"unitPlans": generated["unitPlans"], "lessonPlans": generated["lessonPlans"]},
        },
    }
    MANIFEST_PATH.write_text(json.dumps(manifest, indent=2) + "\n", encoding="utf-8")
    REPORT_PATH.parent.mkdir(parents=True, exist_ok=True)
    REPORT_PATH.write_text(json.dumps(generated, indent=2) + "\n", encoding="utf-8")
    print(json.dumps({"courseRoot": str(COURSE_ROOT), **generated, "reportPath": str(REPORT_PATH)}, indent=2))


def plan_resource(label, role, category, rel_path, preview_rel):
    return {
        "label": label,
        "type": "docx",
        "category": category,
        "role": role,
        "path": rel_path,
        "bytes": (COURSE_ROOT / rel_path).stat().st_size,
        "source": GENERATED_SOURCE,
        "sourceStatus": "reconstructed_from_moodle_content",
        "sourceNote": GENERATED_NOTE,
        "generated": True,
        "previewPath": preview_rel,
    }


def setup_doc(title):
    doc = Document()
    section = doc.sections[0]
    section.orientation = WD_ORIENTATION.PORTRAIT
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)
    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(10)
    p = doc.add_paragraph()
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(14)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return doc


def make_unit_docx(manifest, unit, rel_path):
    unit_no = int(unit["unit"])
    curriculum = CURRICULUM[unit_no]
    lesson_data = [(lesson, expectation_for(lesson)) for lesson in unit.get("lessons", [])]
    title = f"UNIT {unit_no}"
    doc = setup_doc(title)
    table = doc.add_table(rows=0, cols=4)
    table.style = "Table Grid"

    add_span_row(table, "Unit Author")
    add_single_row(table, "Name:")
    add_single_row(table, "School District:")
    add_single_row(table, "School Name:")
    add_single_row(table, "School City, Province:")
    add_span_row(table, "Unit Overview")
    add_label_value(table, "Unit Title Name", unit["title"])
    add_label_value(table, "Unit Summary", unit_summary(unit))
    add_label_value(table, "Year Level", "Grade 11, University/College Preparation")
    add_label_value(table, "Approximate Time Needed", approximate_hours(unit))
    add_span_row(table, "Unit Foundation")
    add_label_value(
        table,
        "Targeted Curriculum Expectations",
        curriculum["strand"] + "\n" + bullet_text(curriculum["overall"]) + "\n\nSpecific\n" + bullet_text(
            unique(flatten([exp["specific"] for _, exp in lesson_data]))[:20]
        ),
    )
    add_label_value(
        table,
        "Learning Goals",
        bullet_text(unit_learning_goals(lesson_data)),
    )
    add_unit_assessment_rows(table, unit)
    add_span_row(table, "Unit Details")
    add_label_value(table, "Lesson and Assessment Outlines", unit_outline(unit, lesson_data))
    add_span_row(table, "Materials and Resources")
    add_label_value(table, "Technology", unit_technology(unit))
    add_label_value(table, "Printed", unit_printed(unit))
    add_label_value(table, "Other Resources", "Anecdotal notes of observation\nExit slips\nTeacher and student checklists\nStudent peer and/or self reflection")
    add_label_value(table, "Notes", GENERATED_NOTE)

    style_table(table)
    save_doc(doc, rel_path)


def make_lesson_docx(manifest, unit, lesson, exp, rel_path):
    unit_no = int(unit["unit"])
    curriculum = CURRICULUM[unit_no]
    doc = setup_doc("Lesson Plan")
    table = doc.add_table(rows=0, cols=3)
    table.style = "Table Grid"

    add_two_cell(table, "Lesson Plan", f"Subject: Financial Accounting Fundamentals, Grade 11")
    add_full_value(table, "Lesson Name:", lesson["title"])
    add_full_value(table, "Unit of Study:", unit["title"])
    add_label_value3(
        table,
        "PRIOR KNOWLEDGE What do my students already know? What key questions or strategies will activate prior knowledge? What connections can I help students make?",
        prior_knowledge(unit, lesson),
    )
    add_label_value3(
        table,
        "CURRICULUM EXPECTATIONS",
        "OVERALL\n"
        + bullet_text(exp["overall"] or curriculum["overall"])
        + "\n\nSPECIFIC\n"
        + bullet_text(exp["specific"]),
    )
    add_label_value3(
        table,
        "LEARNING GOALS What do I want students to know and be able to do?",
        bullet_text(lesson_learning_goals(unit, lesson, exp)),
    )
    add_label_value3(
        table,
        "SUCCESS CRITERIA(S) How will students know they attained the goal? What tools can they use to monitor their progress?",
        bullet_text(lesson_success_criteria(unit, lesson, exp)),
    )
    add_assessment_row3(table, unit, lesson)
    add_label_value3(table, "What will I do?", teacher_assessment_action(unit, lesson))
    add_label_value3(
        table,
        "Accommodations: How will you change the lesson to meet the needs of individual students?",
        accommodations_for(lesson),
    )
    add_label_value3(
        table,
        "Materials and Resources: What do I need to prepare in advance? Equipment? Student Materials? Teacher Resources? Human Resources?",
        materials_for(lesson),
    )
    style_table(table)

    doc.add_paragraph()
    flow = doc.add_table(rows=0, cols=6)
    flow.style = "Table Grid"
    add_span_row_cols(flow, "DELIVERING THE LESSON", 6)
    for row in lesson_flow_rows(unit, lesson, exp):
        add_flow_row(flow, *row)
    add_span_row_cols(flow, "Notes:\n" + GENERATED_NOTE, 6)
    style_table(flow)
    save_doc(doc, rel_path)


def make_unit_preview(manifest, unit, preview_rel, source_rel):
    unit_no = int(unit["unit"])
    curriculum = CURRICULUM[unit_no]
    lesson_data = [(lesson, expectation_for(lesson)) for lesson in unit.get("lessons", [])]
    rows = [
        ["Unit Author", ""],
        ["Name:", ""],
        ["School District:", ""],
        ["School Name:", ""],
        ["School City, Province:", ""],
        ["Unit Overview", ""],
        ["Unit Title Name", unit["title"]],
        ["Unit Summary", unit_summary(unit)],
        ["Year Level", "Grade 11, University/College Preparation"],
        ["Approximate Time Needed", approximate_hours(unit)],
        ["Unit Foundation", ""],
        [
            "Targeted Curriculum Expectations",
            curriculum["strand"] + "\n" + bullet_text(curriculum["overall"]) + "\n\nSpecific\n" + bullet_text(unique(flatten([exp["specific"] for _, exp in lesson_data]))[:20]),
        ],
        ["Learning Goals", bullet_text(unit_learning_goals(lesson_data))],
        ["Assessment Plan", assessment_text(unit)],
        ["Unit Details", ""],
        ["Lesson and Assessment Outlines", unit_outline(unit, lesson_data)],
        ["Notes", GENERATED_NOTE],
    ]
    write_preview(preview_rel, f"BAF3M Unit {unit_no} Plan", rows, source_rel)


def make_lesson_preview(manifest, unit, lesson, exp, preview_rel, source_rel):
    unit_no = int(unit["unit"])
    curriculum = CURRICULUM[unit_no]
    rows = [
        ["Lesson Plan", "Subject: Financial Accounting Fundamentals, Grade 11"],
        ["Lesson Name:", lesson["title"]],
        ["Unit of Study:", unit["title"]],
        ["PRIOR KNOWLEDGE", prior_knowledge(unit, lesson)],
        ["CURRICULUM EXPECTATIONS", "OVERALL\n" + bullet_text(exp["overall"] or curriculum["overall"]) + "\n\nSPECIFIC\n" + bullet_text(exp["specific"])],
        ["LEARNING GOALS", bullet_text(exp["goals"] or convert_to_lesson_goals(exp["specific"]))],
        ["SUCCESS CRITERIA(S)", bullet_text(exp["success"] or convert_to_success(exp["goals"]))],
        ["Assessment", assessment_text(unit)],
        ["What will I do?", "Confer\nObserve\nGrade"],
        ["Accommodations", "Increase time\nSimplifying Language\nDecrease time\nScribe/Oral explanation\nPeer tutor/partner\nExtend\nInclude visuals"],
        ["Materials and Resources", materials_for(lesson)],
        ["DELIVERING THE LESSON", "Minds On: review expectations and activate prior knowledge.\nAction: teach with Moodle lesson page, iSpring, and Hands On/H5P.\nConsolidation: complete consolidation, exit slip/self-check, homework, and submission workflow."],
        ["Notes", GENERATED_NOTE],
    ]
    write_preview(preview_rel, f"BAF3M Unit {unit_no} Lesson {lesson['lesson']} Lesson Plan", rows, source_rel)


def save_doc(doc, rel_path):
    path = COURSE_ROOT / rel_path
    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(path)


def write_preview(rel_path, title, rows, source_rel):
    path = COURSE_ROOT / rel_path
    path.parent.mkdir(parents=True, exist_ok=True)
    trs = "\n".join(
        f"<tr><th>{html.escape(k)}</th><td>{html.escape(v).replace(chr(10), '<br>')}</td></tr>"
        for k, v in rows
    )
    body = f"""<!doctype html>
<html lang="en">
<head>
  <meta charset="utf-8">
  <meta name="viewport" content="width=device-width, initial-scale=1">
  <title>{html.escape(title)}</title>
  <style>
    body {{ margin: 0; background: #eef4fb; color: #001f42; font-family: Arial, Helvetica, sans-serif; line-height: 1.45; }}
    main {{ max-width: 980px; margin: 48px auto; background: #fff; border: 1px solid #d5e3f4; border-radius: 6px; padding: 28px 32px; }}
    h1 {{ margin: 0 0 20px; text-align: center; font-size: 22px; }}
    table {{ width: 100%; border-collapse: collapse; }}
    th, td {{ border: 1px solid #9fb9d6; padding: 9px 10px; vertical-align: top; font-size: 14px; }}
    th {{ width: 28%; background: #eaf2fb; text-align: left; }}
    .source {{ margin-top: 18px; font-size: 13px; color: #465f7f; }}
  </style>
</head>
<body>
  <main>
    <h1>{html.escape(title)}</h1>
    <table>{trs}</table>
    <p class="source">DOCX source: {html.escape(source_rel)}</p>
  </main>
</body>
</html>
"""
    path.write_text(body, encoding="utf-8")


TOPIC_DETAILS = [
    (
        ["accounting & bookkeeping", "bookkeeping"],
        {
            "application": "distinguish accounting from bookkeeping, connect source documents to financial records, and explain why accurate record keeping matters to business decisions",
            "question": "What information would a business owner lose if transactions were not recorded clearly?",
        },
    ),
    (
        ["users & uses", "uses of accounting"],
        {
            "application": "identify internal and external users of accounting information and match each user to the decisions they need to make",
            "question": "How would a bank, owner, employee, and government agency use the same accounting information differently?",
        },
    ),
    (
        ["gaap"],
        {
            "application": "apply generally accepted accounting principles to simple business situations and justify which principle supports the decision",
            "question": "Which accounting principle would keep this record fair, consistent, and understandable?",
        },
    ),
    (
        ["forms of business"],
        {
            "application": "compare sole proprietorships, partnerships, and corporations through ownership, liability, financing, and accounting implications",
            "question": "How does the form of business change what the owner records and reports?",
        },
    ),
    (
        ["accounting transactions"],
        {
            "application": "classify accounts, analyse transactions, and show the effect on the accounting equation before students record entries",
            "question": "What accounts change, and do assets still equal liabilities plus owner's equity?",
        },
    ),
    (
        ["trial balance", "financial statements"],
        {
            "application": "prepare a trial balance, locate errors, and connect account balances to the financial statements of a service business",
            "question": "What does a trial balance prove, and what does it not prove?",
        },
    ),
    (
        ["year end"],
        {
            "application": "sequence year-end procedures, closing entries, and post-closing checks so students see how temporary accounts reset for the next fiscal period",
            "question": "Why must revenue and expense accounts be closed before the next fiscal year begins?",
        },
    ),
    (
        ["inventory"],
        {
            "application": "explain inventory principles for merchandising businesses, including inventory flow, cost of goods sold, and the effect of inventory errors",
            "question": "How does the value assigned to inventory affect profit?",
        },
    ),
    (
        ["merchandising"],
        {
            "application": "model purchases, sales, returns, discounts, and cost of goods sold for a merchandising business",
            "question": "What extra records does a merchandising business need compared with a service business?",
        },
    ),
    (
        ["closing entries of merchandising"],
        {
            "application": "complete closing entries for a merchandising company and distinguish temporary accounts from permanent accounts",
            "question": "Which merchandising accounts must be closed, and where do the balances go?",
        },
    ),
    (
        ["susidiary", "subsidiary", "special journals"],
        {
            "application": "use subsidiary ledgers and special journals to organize repeated transactions and reconcile control accounts",
            "question": "Why would a growing business separate customer, supplier, sales, and cash records?",
        },
    ),
    (
        ["sales tax"],
        {
            "application": "record sales tax collected and paid, identify liability balances, and explain the flow of HST/sales tax through the accounts",
            "question": "When a business collects sales tax from a customer, why is it not business revenue?",
        },
    ),
    (
        ["accounting software"],
        {
            "application": "connect manual accounting steps to computerized accounting workflows, reports, and audit trails",
            "question": "Which parts of the accounting cycle become faster with software, and which still require judgement?",
        },
    ),
    (
        ["internal control"],
        {
            "application": "analyse internal control procedures, separation of duties, authorization, documentation, and physical controls for financial management",
            "question": "What could go wrong if one person controls every step of cash handling?",
        },
    ),
    (
        ["budgeted"],
        {
            "application": "prepare and interpret budgeted financial statements, linking assumptions to projected income, expenses, cash, and financial position",
            "question": "Which assumption would change the budget the most, and how would management respond?",
        },
    ),
    (
        ["auditors"],
        {
            "application": "explain the role of auditors, audit evidence, independence, materiality, and how audits support trust in financial statements",
            "question": "What evidence would convince an auditor that the financial information is reliable?",
        },
    ),
    (
        ["current assets", "liabilities"],
        {
            "application": "classify current assets and current liabilities and connect them to working capital and short-term obligations",
            "question": "Can a profitable business still have trouble paying its bills on time?",
        },
    ),
    (
        ["liquidity", "solvency"],
        {
            "application": "calculate and interpret liquidity and solvency measures to evaluate short-term and long-term financial health",
            "question": "What is the difference between being liquid and being solvent?",
        },
    ),
    (
        ["profitability"],
        {
            "application": "analyse profitability using margins, return measures, and trend comparisons so students support conclusions with evidence",
            "question": "Is higher sales revenue always evidence of better performance?",
        },
    ),
    (
        ["financial analysis", "decision making"],
        {
            "application": "use financial statement information and ratios to make reasoned business decisions and identify limitations in the data",
            "question": "What decision would you recommend, and what financial evidence supports it?",
        },
    ),
    (
        ["ethics"],
        {
            "application": "apply ethical decision-making to accounting scenarios involving accuracy, confidentiality, pressure, and professional responsibility",
            "question": "What should an accountant do when accurate reporting conflicts with pressure from a manager?",
        },
    ),
    (
        ["current issues"],
        {
            "application": "connect current accounting issues to regulation, reporting quality, stakeholder trust, and changes in business practice",
            "question": "How can a current issue change what stakeholders expect from accountants?",
        },
    ),
    (
        ["manual & computerized"],
        {
            "application": "compare manual and computerized accounting systems, including speed, accuracy, controls, reporting, and risk of input errors",
            "question": "What controls are still needed when accounting is computerized?",
        },
    ),
    (
        ["technology & accounting"],
        {
            "application": "assess how technology changes accounting tasks, data security, reporting, collaboration, and professional judgement",
            "question": "Which accounting tasks can technology improve, and which still need human judgement?",
        },
    ),
    (
        ["careers"],
        {
            "application": "investigate accounting designations, career pathways, workplace skills, and how course concepts connect to business careers",
            "question": "Which accounting career pathway best fits different strengths and interests?",
        },
    ),
]


def lesson_topic_detail(lesson):
    title = lesson.get("title", "").lower()
    for keys, detail in TOPIC_DETAILS:
        if any(key in title for key in keys):
            return detail
    return {
        "application": "apply the lesson expectations to a realistic accounting scenario and explain the reasoning behind each decision",
        "question": f"What accounting decision is central to {lesson.get('title', 'this topic')}?",
    }


def lesson_learning_goals(unit, lesson, exp):
    detail = lesson_topic_detail(lesson)
    goals = [
        f"Students will understand how {lesson['title']} fits into the accounting work developed in {unit['title']}.",
        f"Students will learn to {detail['application']}.",
    ]
    for item in exp["specific"][:2]:
        goals.append("Students will connect their practice to this expectation: " + lower_initial(item) + ".")
    if lesson.get("h5p"):
        goals.append("Students will use the Hands On/H5P practice to test understanding before completing the worksheet independently.")
    return unique(goals)


def lesson_success_criteria(unit, lesson, exp):
    detail = lesson_topic_detail(lesson)
    criteria = [
        "I can explain the lesson topic in my own words and use the correct accounting vocabulary.",
        "I can work through a teacher-modelled example and explain why each accounting step is needed.",
        f"I can {detail['application']} in a short practice scenario.",
    ]
    for item in exp["specific"][:2]:
        criteria.append("I can show evidence that I can " + lower_initial(item) + ".")
    if lesson.get("h5p"):
        criteria.append("I can use feedback from the Hands On/H5P activity to correct mistakes before submitting homework.")
    if any("Exit Slip" in item.get("label", "") for item in lesson.get("downloads", [])):
        criteria.append("I can complete the exit slip with one accurate concept, one example, and one question if I still need help.")
    return unique(criteria)[:7]


def teacher_assessment_action(unit, lesson):
    actions = [
        "Confer with students while they move from the modelled example to guided practice.",
        "Observe whether students can explain the accounting reasoning, not only the final answer.",
        "Use H5P/Hands On results, worksheet attempts, and homework submissions to identify misconceptions for the next lesson.",
    ]
    if any("Exit Slip" in item.get("label", "") for item in lesson.get("downloads", [])):
        actions.append("Review exit slips for patterns that need whole-class reteaching or individual follow-up.")
    if any("Self-Check" in item.get("label", "") for item in lesson.get("downloads", [])):
        actions.append("Use the self-check/reflection to confirm readiness before the next unit concept.")
    if unit.get("unitResources", {}).get("evaluations"):
        actions.append("Connect feedback to the unit AOL task so students know how practice evidence prepares them for assessment of learning.")
    return "\n".join(actions)


def accommodations_for(lesson):
    detail = lesson_topic_detail(lesson)
    items = [
        "Provide extra time for multi-step accounting procedures and calculations.",
        "Simplify language by pre-teaching key vocabulary and showing a completed model before independent work.",
        "Chunk the task into identify accounts, choose the procedure, complete the entry/calculation, and check the result.",
        "Allow oral explanation, scribing support, or a peer partner for students who can reason through the accounting but need writing support.",
        "Use visuals such as T-accounts, accounting equation charts, journal/ledger templates, ratio formula cards, or comparison tables.",
        f"Extend by asking students to explain a second scenario: {detail['question']}",
    ]
    return "\n".join(items)


def lesson_flow_rows(unit, lesson, exp):
    detail = lesson_topic_detail(lesson)
    lesson_text = section_text(lesson, "Lesson")
    hands_on = section_text(lesson, "Hands On")
    consolidation = section_text(lesson, "Consolidation")
    homework = section_text(lesson, "Homework")
    goals = lesson_learning_goals(unit, lesson, exp)
    success = lesson_success_criteria(unit, lesson, exp)
    scenario = lesson_scenario(unit, lesson)

    minds_on = (
        f"Begin with a short accounting scenario: {scenario} "
        f"Ask the lesson question: {detail['question']} "
        f"Activate prior knowledge through {opening_source(unit, lesson)}. "
        f"Students preview the learning goal and identify what evidence would prove they understand it: {sentence_text(first_or_default(goals, 'understand the lesson concept'))}."
    )
    if lesson_text:
        minds_on += " Use the Moodle opening prompt only as the launch point: " + short_summary(lesson_text, 160)

    action = (
        f"Teach the concept through an I do / We do / You do progression. First, model how to {detail['application']} using the scenario. "
        "Then pause the iSpring/Moodle lesson at natural checkpoints so students explain the accounting reasoning, identify the accounts or evidence involved, and catch common errors before they practise. "
        f"Success check: {sentence_text(first_or_default(success, 'I can explain and apply the accounting concept'))}."
    )

    hands_on_note = activity_summary(hands_on, "Hands On")
    consolidation_note = activity_summary(consolidation, "Consolidation")
    homework_note = activity_summary(homework, "Homework")

    practice = (
        "Students complete the Hands On/local H5P activity as guided practice. Circulate while students justify their choices, then pull one common error for a quick reteach before students move into independent worksheet work. "
        + (hands_on_note if hands_on_note else "Use strategic questioning to have students explain each accounting choice.")
    )

    consolidation_action = (
        "Students complete the Moodle Consolidation item and write a brief evidence statement: what they can now do, what step they must check carefully, and what question remains. "
        + (consolidation_note if consolidation_note else "Use an exit slip, self-check, or short written reflection when listed in the lesson resources.")
        + " Assign the matching homework worksheet and submit through the Homework Submission Folder."
    )
    if homework_note:
        consolidation_action += " Homework instruction basis: " + homework_note

    return [
        ("15 minutes", "Minds On!", minds_on, "Lesson Expectations page\nKWL/opening prompt"),
        ("45 minutes", "Action!", action, "Moodle Lesson page\niSpring presentation\nTeacher worked example"),
        ("45 minutes", "Practice!", practice, "Hands On / H5P\nWorksheet\nAnecdotal notes"),
        ("45 minutes", "Consolidation!", consolidation_action, "Consolidation page\nExit slip/self-check when provided\nHomework worksheet\nStep by step guide"),
    ]


def lesson_scenario(unit, lesson):
    title = lesson.get("title", "").lower()
    if "bookkeeping" in title:
        return "a small business owner has receipts, invoices, and bank deposits but no clear records yet."
    if "users" in title:
        return "the same set of financial statements is reviewed by an owner, lender, employee, and tax authority."
    if "gaap" in title:
        return "two students record the same transaction differently and must decide which treatment follows GAAP."
    if "forms of business" in title:
        return "three entrepreneurs choose between a sole proprietorship, partnership, and corporation."
    if "transactions" in title:
        return "a service business buys supplies, earns revenue, and pays an expense during the same week."
    if "trial balance" in title:
        return "a trial balance does not balance and the class must trace the likely error."
    if "year end" in title or "closing entries" in title:
        return "a business reaches year end and needs to prepare records for the next fiscal period."
    if "inventory" in title:
        return "a store buys inventory at different costs and must decide how inventory and cost of goods sold are reported."
    if "merchandising" in title:
        return "a retail business records purchases, sales, returns, and cost of goods sold."
    if "ledger" in title or "journal" in title:
        return "a growing business has many customer and supplier transactions that no longer fit comfortably in one general ledger view."
    if "sales tax" in title:
        return "a business collects tax from customers and later remits it to the government."
    if "software" in title or "computerized" in title or "technology" in title:
        return "a business moves from manual records to accounting software and must protect accuracy and security."
    if "internal control" in title:
        return "one employee receives cash, records it, and reconciles the account without review."
    if "budget" in title:
        return "management compares budgeted results with actual results and must explain the variance."
    if "auditor" in title:
        return "an auditor asks for evidence before trusting a reported balance."
    if "current assets" in title:
        return "a company has profit on paper but limited cash to pay bills due this month."
    if "liquidity" in title or "solvency" in title:
        return "two companies show different current ratios and debt ratios, and investors must judge risk."
    if "profitability" in title:
        return "sales increased, but profit margin decreased, and the class must explain the performance."
    if "decision" in title:
        return "an investor and a manager use the same financial analysis to make different decisions."
    if "ethics" in title:
        return "an accountant is asked to delay recording an expense to make results look better."
    if "current issues" in title:
        return "a news event or reporting issue changes what stakeholders expect from accounting information."
    if "career" in title:
        return "students compare accounting career pathways and the education/skills each one requires."
    return f"a business situation connected to {lesson.get('title', 'the lesson topic')}."


def opening_source(unit, lesson):
    if lesson["lesson"] == 1:
        return "the unit KWL chart/opening Moodle activity"
    previous = next((item for item in unit.get("lessons", []) if item.get("lesson") == lesson["lesson"] - 1), None)
    if previous:
        return f"a review of Unit {unit['unit']} Lesson {previous['lesson']} ({previous['title']})"
    return "a review of the previous accounting concept"


def section_text(lesson, label):
    section = next((s for s in lesson.get("bookSections", []) if s.get("sectionLabel") == label), None)
    if not section:
        return ""
    text = ""
    if section.get("path"):
        page = COURSE_ROOT / section["path"]
        if page.exists():
            text = strip_html(page.read_text(encoding="utf-8", errors="ignore"))
    if not text:
        text = clean(section.get("textPreview", ""))
    return clean(remove_common_page_text(text, lesson, label))


def remove_common_page_text(text, lesson, label):
    unit_no = lesson.get("unit", "")
    lesson_no = lesson.get("lesson", "")
    title = str(lesson.get("title", ""))
    patterns = [
        rf"\bBAF3M\s+Unit\s+{unit_no}\s+Lesson\s+{lesson_no}\s*-\s*{re.escape(label)}\b",
        rf"\bLesson\s+{lesson_no}\s*:\s*{re.escape(title)}\s+{re.escape(label)}\b",
        rf"\b{re.escape(title)}\s+{re.escape(label)}\b",
        r"^Moodle activity\s+",
    ]
    out = clean(text)
    for pattern in patterns:
        out = re.sub(pattern, " ", out, flags=re.I).strip()
    out = re.sub(r"\bclick\s+HERE\b", "use the linked local resource", out, flags=re.I)
    out = re.sub(r"\bHERE\b", "the linked local resource", out)
    return clean(out)


def activity_summary(text, label):
    text = clean(text)
    lower = text.lower()
    if not text:
        return ""
    if label == "Hands On" and "post skills check quiz" in lower:
        return "Moodle identifies this as a Post Skills Check practice quiz; students test understanding, repeat as needed, and use feedback before moving on."
    if label == "Consolidation" and "post skills check quiz" in lower:
        return "Moodle uses a Consolidation Activity / Post Skills Check quiz so students can confirm understanding before submitting homework."
    if label == "Consolidation" and "summary video" in lower:
        return "Moodle directs students to review the summary video to reinforce the concepts learned in the lesson."
    if label == "Consolidation" and "exit slip" in lower:
        return "Moodle includes an Exit Slip / record-your-work step for students to summarize evidence of learning."
    if label == "Consolidation" and "self-check" in lower:
        return "Moodle includes a self-check/reflection step so students can monitor readiness before the next lesson."
    if label == "Homework" and "download, print and complete" in lower:
        return "Download, print and complete the worksheet; scan or photograph the completed work, save it as a Word/PDF file, and submit it in the matching folder."
    return short_summary(text, 220)


def short_summary(text, limit=220):
    text = clean(text)
    if len(text) <= limit:
        return text
    sentences = re.split(r"(?<=[.!?])\s+", text)
    summary = ""
    for sentence in sentences:
        candidate = clean((summary + " " + sentence).strip())
        if len(candidate) > limit:
            break
        summary = candidate
    if summary:
        return summary
    return text[: limit - 3].rstrip() + "..."


def sentence_text(value):
    return clean(value).rstrip(".")


def first_or_default(items, default):
    for item in items or []:
        if clean(item):
            return clean(item)
    return default


def expectation_for(lesson):
    section = next((s for s in lesson.get("bookSections", []) if s.get("sectionLabel") == "Lesson Expectations"), None)
    if not section and lesson.get("bookSections"):
        section = lesson["bookSections"][0]
    text = ""
    exists = False
    if section and section.get("path"):
        page = COURSE_ROOT / section["path"]
        exists = page.exists()
        if exists:
            text = strip_html(page.read_text(encoding="utf-8", errors="ignore"))
    if not text and section:
        text = clean(section.get("textPreview", ""))
    return {
        "exists": exists,
        "text": text,
        "overall": extract_items(text, "Overall Expectations", ["Specific Lesson Expectations", "Learning Goals", "Success Criteria"]),
        "specific": extract_items(text, "Specific Lesson Expectations", ["Learning Goals", "Success Criteria"]),
        "goals": extract_items(text, "Learning Goals", ["Success Criteria"]),
        "success": extract_items(text, "Success Criteria", []),
    }


def strip_html(value):
    value = re.sub(r"<script[\s\S]*?</script>", " ", value, flags=re.I)
    value = re.sub(r"<style[\s\S]*?</style>", " ", value, flags=re.I)
    value = re.sub(r"<br\s*/?>", "\n", value, flags=re.I)
    value = re.sub(r"</(p|div|li|tr|h[1-6])>", "\n", value, flags=re.I)
    value = re.sub(r"<[^>]+>", " ", value)
    return clean(html.unescape(value))


def clean(value):
    value = re.sub(r"([A-Za-z])-\s+([A-Za-z])", r"\1\2", str(value or ""))
    return re.sub(r"\s+", " ", value).strip()


def extract_items(text, label, next_labels):
    lower = text.lower()
    start = lower.find(label.lower())
    if start < 0:
        return []
    start += len(label)
    if start < len(text) and text[start] == ":":
        start += 1
    end = len(text)
    for nxt in next_labels:
        index = lower.find(nxt.lower(), start)
        if index >= 0:
            end = min(end, index)
    block = clean(text[start:end])
    block = re.sub(r"\s*;\s*", "\n", block)
    block = re.sub(r"\s+\d+\.\s+", "\n", block)
    items = [clean(x).rstrip(".;") for x in re.split(r"\n+|(?<=\.)\s+(?=I am|I will|Students|Explain|Describe|Apply|Assess|Demonstrate|Evaluate)", block)]
    return unique([x for x in items if x])[:12]


def unit_summary(unit):
    pedagogy = UNIT_PEDAGOGY[int(unit["unit"])]
    evals = ", ".join(item.get("label", "") for item in unit.get("unitResources", {}).get("evaluations", []))
    return (
        f"{pedagogy['big_idea']}\n\n"
        f"Instructional sequence: {pedagogy['sequence']}\n\n"
        f"Teacher planning focus: {pedagogy['teacher_intent']}\n\n"
        "Students work through the Moodle lesson pages, local iSpring presentations, localized H5P/Hands On practice, "
        "homework worksheets, and reflection/self-check activities. "
        + (f"Assessment of learning is gathered through {evals}." if evals else "Assessment of learning is gathered through the unit evaluation tasks.")
    )


def approximate_hours(unit):
    return f"{len(unit.get('lessons', [])) * 4} hrs"


def unit_learning_goals(lesson_data):
    if not lesson_data:
        return []
    unit_no = int(lesson_data[0][0]["unit"])
    return UNIT_PEDAGOGY[unit_no]["goals"]


def unit_outline(unit, lesson_data):
    chunks = []
    for lesson, exp in lesson_data:
        detail = lesson_topic_detail(lesson)
        assessments = ["Evidence: homework worksheet submitted through the matching Homework Submission Folder"]
        if lesson["lesson"] == 1:
            assessments.insert(0, "Evidence: unit KWL/opening reflection activates prior knowledge")
        if lesson.get("h5p"):
            assessments.append("Evidence: localized Hands On/H5P practice checks understanding before homework")
        if lesson.get("downloads"):
            if any("Exit Slip" in d.get("label", "") for d in lesson.get("downloads", [])):
                assessments.append("Evidence: exit slip captures consolidation and questions for follow-up")
            if any("Self-Check" in d.get("label", "") for d in lesson.get("downloads", [])):
                assessments.append("Evidence: self-check/reflection monitors readiness for the next unit step")
        if not exp["specific"]:
            expectation_line = "Target: apply the lesson topic to a practical accounting decision."
        else:
            expectation_line = "Target: " + "; ".join(exp["specific"][:2])
        chunks.append(
            f"Lesson {lesson['lesson']} - {lesson['title']}\n"
            f"{expectation_line}\n"
            f"Teacher plan: model how students will {detail['application']}; then move students from guided examples into Moodle Hands On/H5P practice and worksheet evidence.\n"
            + "\n".join(assessments)
        )
    for item in unit.get("unitResources", {}).get("evaluations", []):
        chunks.append(
            f"Assessment of Learning - {item.get('label', '')}\n"
            "Students use the unit concepts in an assessed task. Teacher feedback should focus on accounting accuracy, reasoning, communication, and application of the relevant procedures."
        )
    return "\n\n".join(chunks)


def prior_knowledge(unit, lesson):
    detail = lesson_topic_detail(lesson)
    if lesson["lesson"] == 1:
        return (
            "Students activate existing knowledge through the unit KWL/opening activity and connect prior business, "
            f"personal finance, or workplace experience to the lesson focus. Key activating question: {detail['question']}"
        )
    previous = next((item for item in unit.get("lessons", []) if item.get("lesson") == lesson["lesson"] - 1), None)
    previous_text = f"Unit {unit['unit']} Lesson {previous['lesson']} ({previous['title']})" if previous else "the previous lesson"
    return (
        f"Students build from {previous_text} and use the current Lesson Expectations page to connect prior accounting concepts "
        f"to {lesson['title']}. Key activating question: {detail['question']}"
    )


def materials_for(lesson):
    detail = lesson_topic_detail(lesson)
    items = [
        "Teacher preparation: review the Lesson Expectations page and choose one modelled example for the lesson scenario.",
        "Teacher preparation: prepare a quick misconception check connected to this question: " + detail["question"],
    ]
    if lesson.get("ispring"):
        items.append(f"Instructional presentation: Unit {lesson['unit']} Lesson {lesson['lesson']} local iSpring presentation")
    labels = [section.get("sectionLabel") for section in lesson.get("bookSections", []) if section.get("sectionLabel")]
    if labels:
        items.append("Moodle sequence: " + ", ".join(labels))
    if lesson.get("h5p"):
        items.append("Guided practice: localized Hands On/H5P activity")
    worksheet_labels = []
    submission_labels = []
    for item in lesson.get("downloads", []):
        label = item.get("label", item.get("path", "Resource"))
        if re.search(r"worksheet|step by step|guide|kwl|reflection|self-check|exit slip", label, flags=re.I):
            worksheet_labels.append(label)
        else:
            submission_labels.append(label)
    if worksheet_labels:
        items.append("Student files: " + "; ".join(unique(worksheet_labels)))
    if submission_labels:
        items.append("Linked/local activity resources: " + "; ".join(unique(submission_labels)))
    return "\n".join(unique(items))


def assessment_text(unit):
    return (
        "Assessment as Learning\n"
        "[x] Observation\n[x] Anecdotal Notes\n[x] Exit Card\n[x] Self-Assessment checklist\n[x] Reflection / Learning Log\n\n"
        "Assessment for Learning\n"
        "[x] Worksheets\n[x] Homework\n[x] Strategic Questioning\n[x] Hands On / H5P activities\n\n"
        "Assessment of Learning\n"
        + "\n".join(f"[x] {item.get('label', '')}" for item in unit.get("unitResources", {}).get("evaluations", []))
    )


def convert_to_lesson_goals(items):
    return [f"By the end of this lesson, students will be able to {lower_initial(x)}" for x in items[:6]]


def convert_to_success(items):
    return [re.sub(r"^I am learning to ", "I can ", x, flags=re.I) for x in items[:6]]


def lower_initial(value):
    value = clean(value)
    return value[:1].lower() + value[1:] if value else value


def add_span_row(table, text):
    add_span_row_cols(table, text, 4)


def add_span_row_cols(table, text, cols):
    cells = table.add_row().cells
    cell = cells[0]
    for other in cells[1:]:
        cell = cell.merge(other)
    set_cell_text(cells[0], text, bold=True)


def add_single_row(table, text):
    cells = table.add_row().cells
    cells[0].merge(cells[-1])
    set_cell_text(cells[0], text)


def add_label_value(table, label, value):
    cells = table.add_row().cells
    cells[0].merge(cells[1])
    cells[2].merge(cells[3])
    set_cell_text(cells[0], label, bold=True)
    set_cell_text(cells[2], value)


def add_unit_assessment_rows(table, unit):
    add_span_row(table, "Assessment Plan")
    header = table.add_row().cells
    header[2].merge(header[3])
    set_cell_text(header[0], "Assessment as Learning (ASL)", bold=True)
    set_cell_text(header[1], "Assessment for Learning (AFL)", bold=True)
    set_cell_text(header[2], "Assessment of Learning (AoL)", bold=True)
    content = table.add_row().cells
    content[2].merge(content[3])
    set_cell_text(content[0], "☑ Observation\n☑ Anecdotal Notes\n☑ Exit Card\n☑ Self-Assessment checklist\n☑ Discussions")
    set_cell_text(content[1], "☑ Worksheets\n☑ Homework\n☑ Strategic Questioning\n☑ Hands-On Activity")
    aol = "\n".join(f"☑ {item.get('label', '')}" for item in unit.get("unitResources", {}).get("evaluations", []))
    set_cell_text(content[2], aol or "☑ Assignments\n☑ Portfolio")


def add_two_cell(table, left, right):
    cells = table.add_row().cells
    cells[1].merge(cells[2])
    set_cell_text(cells[0], left, bold=True)
    set_cell_text(cells[1], right)


def add_full_value(table, label, value):
    cells = table.add_row().cells
    cells[1].merge(cells[2])
    set_cell_text(cells[0], label, bold=True)
    set_cell_text(cells[1], value)


def add_label_value3(table, label, value):
    add_full_value(table, label, value)


def add_assessment_row3(table, unit, lesson):
    asl_items = ["☑ Observation", "☑ Anecdotal Notes"]
    if any("Exit Slip" in item.get("label", "") for item in lesson.get("downloads", [])):
        asl_items.append("☑ Exit Card")
    if any("Self-Check" in item.get("label", "") for item in lesson.get("downloads", [])):
        asl_items.append("☑ Self-Assessment checklist")
    if any("Reflection" in item.get("label", "") for item in lesson.get("downloads", [])):
        asl_items.append("☑ Reflection / Learning Log")
    if len(asl_items) == 2:
        asl_items.append("❐ Exit Card")

    afl_items = ["☑ Worksheets", "☑ Homework", "☑ Strategic Questioning"]
    if lesson.get("h5p"):
        afl_items.append("☑ Hands On / H5P")
    else:
        afl_items.append("❐ Hands On / H5P")

    aol_labels = [item.get("label", "") for item in unit.get("unitResources", {}).get("evaluations", []) if item.get("label")]
    aol_items = ["❐ Quizzes", "❐ Unit Test"]
    aol_items.extend(f"☑ {label}" for label in aol_labels)
    aol_items.append("❐ Portfolio")

    cells = table.add_row().cells
    set_cell_text(cells[0], "Assessment as Learning\n" + "\n".join(asl_items))
    set_cell_text(cells[1], "Assessment for Learning\n" + "\n".join(afl_items))
    set_cell_text(cells[2], "Assessment of Learning\n" + "\n".join(aol_items))


def add_flow_row(table, timing, phase, action, materials):
    cells = table.add_row().cells
    set_cell_text(cells[0], timing)
    set_cell_text(cells[1], "W\n❐")
    set_cell_text(cells[2], "S\n❐")
    set_cell_text(cells[3], "I\n❐")
    set_cell_text(cells[4], phase + "\n" + action)
    set_cell_text(cells[5], materials)


def set_cell_text(cell, text, bold=False):
    cell.text = ""
    for index, line in enumerate(str(text or "").split("\n")):
        p = cell.paragraphs[0] if index == 0 else cell.add_paragraph()
        run = p.add_run(line)
        run.bold = bold and index == 0
        run.font.name = "Arial"
        run.font.size = Pt(9)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP


def style_table(table):
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(2)
                for run in p.runs:
                    run.font.name = "Arial"
                    if run.font.size is None:
                        run.font.size = Pt(9)


def bullet_text(items):
    return "\n".join(clean(item) for item in items if clean(item))


def unit_technology(unit):
    lines = []
    unit_no = int(unit["unit"])
    for lesson in unit.get("lessons", []):
        if lesson.get("ispring"):
            lines.append(f"Unit {unit_no} - Lesson {lesson['lesson']} - {lesson['title']} - iSpring presentation")
        if lesson.get("h5p"):
            lines.append(f"Unit {unit_no} - Lesson {lesson['lesson']} - {lesson['title']} - Hands on activity")
        if any(item.get("type") in {"mp4", "video"} or str(item.get("label", "")).lower().endswith(".mp4") for item in lesson.get("downloads", [])):
            lines.append(f"Unit {unit_no} - Lesson {lesson['lesson']} - {lesson['title']} - Consolidation video")
    return "\n".join(unique(lines))


def unit_printed(unit):
    lines = []
    unit_no = int(unit["unit"])
    for lesson in unit.get("lessons", []):
        for item in lesson.get("downloads", []):
            label = item.get("label", "")
            if item.get("type") in {"doc", "docx", "pdf"} and label:
                lines.append(f"Unit {unit_no} - Lesson {lesson['lesson']} - {label}")
    for item in unit.get("unitResources", {}).get("evaluations", []):
        for attachment in item.get("attachments", []):
            if attachment.get("label"):
                lines.append(f"{item.get('label', '')} - {attachment.get('label')}")
    return "\n".join(unique(lines))


def flatten(items):
    return [x for sub in items for x in sub]


def unique(items):
    out = []
    seen = set()
    for item in items:
        key = clean(item).lower()
        if key and key not in seen:
            seen.add(key)
            out.append(clean(item))
    return out


if __name__ == "__main__":
    sys.exit(main())
