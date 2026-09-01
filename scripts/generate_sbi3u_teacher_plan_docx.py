import json
import re
from datetime import datetime, timezone
from html import unescape
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


WORKSPACE_ROOT = Path("D:/工作文件/SUNNYBROOK")
REPO_ROOT = WORKSPACE_ROOT / "ossd-course-portal"
COURSE_ROOT = WORKSPACE_ROOT / "courseware" / "SBI3U"
MANIFEST_PATH = COURSE_ROOT / "course-manifest.json"
GENERATED_AT = datetime.now(timezone.utc).replace(microsecond=0).isoformat()
TEMPLATE_ROOT = REPO_ROOT / "templates" / "teacher-plans"
LESSON_PLAN_TEMPLATE = TEMPLATE_ROOT / "mdm4u-style-lesson-plan-template.docx"
UNIT_PLAN_TEMPLATE = TEMPLATE_ROOT / "mdm4u-style-unit-plan-template.docx"
MDM4U_LESSON_PLAN_REFERENCE = (
    WORKSPACE_ROOT / "courseware" / "MDM4U" / "plans" / "source" / "Unit 1" / "Lesson Plan Unit 1 lesson 1.docx"
)
MDM4U_UNIT_PLAN_REFERENCE = WORKSPACE_ROOT / "courseware" / "MDM4U" / "plans" / "source" / "Unit 1" / "Unit 1 Plan .docx"


STRANDS = {
    1: {
        "code": "D",
        "title": "Genetic Processes",
        "focus": "heredity, mitosis and meiosis, Mendelian inheritance, chromosomal conditions, and ethical issues in genetic research.",
    },
    2: {
        "code": "E",
        "title": "Animals: Structure and Function",
        "focus": "digestive, respiratory, and circulatory systems, their relationships, and health or technology applications.",
    },
    3: {
        "code": "B",
        "title": "Diversity of Living Things",
        "focus": "classification, biodiversity, evolutionary relationships, and the characteristics of major groups of living things.",
    },
    4: {
        "code": "C",
        "title": "Evolution",
        "focus": "evolutionary theory, natural selection, variation, speciation, and applications such as drug resistance.",
    },
    5: {
        "code": "F",
        "title": "Plants: Anatomy, Growth, and Function",
        "focus": "plant tissues, growth factors, reproduction, structures, succession, and sustainable use of plants.",
    },
}


def clean_text(value):
    text = str(value or "")
    text = re.sub(r"<script[\s\S]*?</script>", " ", text, flags=re.I)
    text = re.sub(r"<style[\s\S]*?</style>", " ", text, flags=re.I)
    text = re.sub(r"<[^>]+>", " ", text)
    text = unescape(text)
    text = re.sub(r"\s+", " ", text).strip()
    return text


def read_local_text(relative_path):
    if not relative_path:
        return ""
    path = COURSE_ROOT / relative_path
    if not path.exists():
        return ""
    return clean_text(path.read_text(encoding="utf-8", errors="ignore"))


def short_text(value, limit=360):
    text = clean_text(value)
    if len(text) <= limit:
        return text
    clipped = text[:limit]
    cut = max(clipped.rfind(". "), clipped.rfind("; "), clipped.rfind(", "))
    if cut > 160:
        clipped = clipped[: cut + 1]
    return clipped.strip() + "..."


def find_section(lesson, pattern):
    regex = re.compile(pattern, re.I)
    for section in lesson.get("bookSections") or []:
        label = section.get("label") or section.get("sectionLabel") or ""
        if regex.search(label):
            return section
    return None


def section_text(lesson, pattern):
    section = find_section(lesson, pattern)
    if not section:
        return ""
    return read_local_text(section.get("path")) or clean_text(section.get("textPreview"))


def expectation_codes(text):
    return sorted(set(re.findall(r"\b[A-F][1-3](?:\.\d+)?\b", text)))[:18]


def learning_goals(expectations_text, lesson_title):
    match = re.search(
        r"Learning Goals?:\s*(.+?)(Success Criteria|Lesson|Hands On|Consolidation|Homework|$)",
        expectations_text or "",
        flags=re.I,
    )
    if match:
        raw = match.group(1)
        parts = re.split(
            r"(?=Learn|Explain|Describe|Identify|Compare|Analyse|Analyze|Understand|Investigate|Evaluate|Use|Create|Demonstrate)",
            raw,
        )
        goals = [short_text(part, 150) for part in parts if len(clean_text(part)) > 8]
        if goals:
            return goals[:4]
    return [
        f"Explain the core biology concept in {lesson_title} using precise course vocabulary.",
        "Use evidence from the lesson page, iSpring module, and activity resources to support an explanation.",
        "Apply the lesson concept through the Hands On activity, consolidation task, and homework submission.",
    ]


def all_resources(lesson):
    return [item for item in (lesson.get("downloads") or []) + (lesson.get("textExports") or []) if item]


def resources_matching(lesson, pattern):
    regex = re.compile(pattern, re.I)
    found = []
    for item in all_resources(lesson):
        haystack = " ".join(str(item.get(key, "")) for key in ("label", "type", "category", "role", "path"))
        if regex.search(haystack):
            found.append(item)
    return found


def resource_labels(items):
    return [item.get("label") or item.get("title") or "Untitled resource" for item in items]


def lesson_materials(lesson):
    media = []
    for item in lesson.get("ispring") or []:
        media.append(f"iSpring module: {item.get('label') or 'lesson presentation'}")
    for item in resources_matching(lesson, r"h5p"):
        media.append(f"{role_label(item.get('role'))} H5P: {item.get('label')}")
    for item in resources_matching(lesson, r"mp4|video"):
        media.append(f"{role_label(item.get('role'))} video: {item.get('label')}")
    files = resource_labels(resources_matching(lesson, r"docx?|pdf|pptx?|xlsx?|txt|png|jpe?g|gif|tiff?"))[:10]
    return media, files


def role_label(value):
    label = str(value or "course").replace("_", " ").strip()
    return " ".join(part.capitalize() for part in label.split())


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_width(cell, width_dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def style_document(doc):
    section = doc.sections[0]
    section.top_margin = Inches(0.9)
    section.bottom_margin = Inches(0.9)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.12

    for name, size, color in [
        ("Heading 1", 16, "002B57"),
        ("Heading 2", 12.5, "002B57"),
        ("Heading 3", 10.5, "002B57"),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(12)
        style.paragraph_format.space_after = Pt(5)


def add_title(doc, title, subtitle):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(title)
    run.bold = True
    run.font.name = "Calibri"
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor.from_string("002B57")
    p.paragraph_format.space_after = Pt(2)

    p = doc.add_paragraph()
    run = p.add_run(subtitle)
    run.font.name = "Calibri"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor.from_string("49607A")
    p.paragraph_format.space_after = Pt(8)


def add_callout(doc, text):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    cell = table.cell(0, 0)
    set_cell_width(cell, 8200)
    set_cell_shading(cell, "FFF2CC")
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor.from_string("5F3B00")


def add_bullets(doc, items):
    if not items:
        items = ["No separate local item is listed in the manifest; use the Moodle section page and teacher judgement."]
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.left_indent = Inches(0.24)
        p.paragraph_format.first_line_indent = Inches(-0.12)
        run = p.add_run(str(item))
        run.font.name = "Calibri"
        run.font.size = Pt(10)


def add_numbered(doc, items):
    for idx, item in enumerate(items or [], start=1):
        p = doc.add_paragraph(f"{idx}. {item}")
        p.paragraph_format.left_indent = Inches(0.16)
        p.paragraph_format.first_line_indent = Inches(-0.16)


def add_contents(doc, items):
    doc.add_heading("Contents", level=1)
    for item in items:
        p = doc.add_paragraph(item)
        p.paragraph_format.left_indent = Inches(0.18)


def add_field(doc, label, value="", empty=False):
    doc.add_heading(label, level=3)
    if empty:
        doc.add_paragraph("")
        return
    if isinstance(value, list):
        add_bullets(doc, value)
    else:
        doc.add_paragraph(str(value or ""))


def extract_expectation_sections(text):
    text = clean_text(text)
    if not text:
        return {
            "overall": ["Use the localized Lesson Expectations page for exact wording."],
            "specific": ["Use the localized Lesson Expectations page for exact wording."],
        }
    lower = text.lower()
    overall = text
    specific = ""
    for marker in ["specific expectations", "specific expectation"]:
        pos = lower.find(marker)
        if pos >= 0:
            overall = text[:pos]
            specific = text[pos:]
            break
    for marker in ["learning goals", "success criteria"]:
        pos = overall.lower().find(marker)
        if pos >= 0:
            overall = overall[:pos]
    if not specific:
        specific = text
    for marker in ["learning goals", "success criteria"]:
        pos = specific.lower().find(marker)
        if pos >= 0:
            specific = specific[:pos]
    return {
        "overall": split_sentences(overall, 6),
        "specific": split_sentences(specific, 10),
    }


def split_sentences(text, limit):
    text = clean_text(text)
    if not text:
        return []
    parts = re.split(r"(?<=[.;:])\s+|(?=\([A-F][1-3](?:\.\d+)?\))", text)
    cleaned = [short_text(part, 260) for part in parts if len(clean_text(part)) > 12]
    return cleaned[:limit] or [short_text(text, 260)]


def lesson_success_criteria(lesson_title, codes):
    code_text = ", ".join(codes[:4]) if codes else "the listed lesson expectations"
    return [
        f"I can explain the central idea in {lesson_title} using accurate SBI3U vocabulary.",
        f"I can connect my explanation to {code_text} and cite evidence from the lesson activity.",
        "I can interpret a model, diagram, data set, or example from the lesson and describe what it shows.",
        "I can use feedback from the Hands On or consolidation task to improve my homework response.",
    ]


def unit_summary(unit, strand, lessons):
    first_topics = ", ".join((lesson.get("title") or "") for lesson in lessons[:4])
    return (
        f"This unit develops the {strand['title']} strand through {len(lessons)} linked lessons. "
        f"Students begin with core vocabulary and guided lesson media, practise through Hands On and consolidation tasks, "
        f"and build toward unit evaluation evidence. Early topics include {first_topics}."
    )


def activity_titles(unit, role_pattern):
    regex = re.compile(role_pattern, re.I)
    titles = []
    for lesson in unit.get("lessons") or []:
        for item in (lesson.get("downloads") or []) + (lesson.get("textExports") or []):
            haystack = " ".join(str(item.get(key, "")) for key in ("label", "type", "category", "role", "section"))
            if regex.search(haystack):
                titles.append(item.get("label") or item.get("title") or "Untitled")
    return titles


def add_kv_table(doc, pairs):
    table = doc.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    table.autofit = False
    for label, value in pairs:
        row = table.add_row()
        row.cells[0].text = str(label)
        row.cells[1].text = str(value)
        set_cell_width(row.cells[0], 2100)
        set_cell_width(row.cells[1], 6100)
        set_cell_shading(row.cells[0], "F5F9FD")
        for cell in row.cells:
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.name = "Calibri"
                    run.font.size = Pt(9.3)
    return table


def add_structured_rows(doc, rows):
    table = doc.add_table(rows=0, cols=1)
    table.style = "Table Grid"
    table.autofit = False
    for value in rows:
        row = table.add_row()
        cell = row.cells[0]
        cell.text = str(value or "")
        set_cell_width(cell, 8600)
        for paragraph in cell.paragraphs:
            paragraph.paragraph_format.space_after = Pt(0)
            for run in paragraph.runs:
                run.font.name = "Calibri"
                run.font.size = Pt(10)
    return table


def set_text(cell, text, bold=False, size=9.5):
    cell.text = ""
    lines = str(text or "").split("\n")
    for line_index, line in enumerate(lines):
        paragraph = cell.paragraphs[0] if line_index == 0 else cell.add_paragraph()
        run = paragraph.add_run(line)
        run.bold = bold
        run.font.name = "Calibri"
        run.font.size = Pt(size)


def format_grid_table(table, widths=None, header_rows=()):
    table.style = "Table Grid"
    table.autofit = False
    for row_index, row in enumerate(table.rows):
        for cell_index, cell in enumerate(row.cells):
            if widths:
                set_cell_width(cell, widths[min(cell_index, len(widths) - 1)])
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            if row_index in header_rows:
                set_cell_shading(cell, "E6EEF7")
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                for run in paragraph.runs:
                    run.font.name = "Calibri"
                    run.font.size = Pt(9)
                    if row_index in header_rows:
                        run.bold = True


def merge_cells(cells):
    merged = cells[0]
    for cell in cells[1:]:
        merged = merged.merge(cell)
    return merged


def template_document(kind):
    if kind == "lesson":
        path = LESSON_PLAN_TEMPLATE if LESSON_PLAN_TEMPLATE.exists() else MDM4U_LESSON_PLAN_REFERENCE
    elif kind == "unit":
        path = UNIT_PLAN_TEMPLATE if UNIT_PLAN_TEMPLATE.exists() else MDM4U_UNIT_PLAN_REFERENCE
    else:
        raise ValueError(f"Unknown plan template kind: {kind}")
    return Document(str(path))


def set_row_text(row, text, bold=False, size=9.5):
    set_text(row.cells[0], text, bold=bold, size=size)


def add_sequence_table(doc, rows):
    table = doc.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    table.autofit = False
    headers = ["Lesson", "Topic", "Teacher Focus", "Student Task", "Assessment / Feedback"]
    widths = [850, 1650, 2450, 2450, 2050]
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        cell.text = header
        set_cell_width(cell, widths[idx])
        set_cell_shading(cell, "E6EEF7")
    for lesson_id, topic, focus, task, feedback in rows:
        row = table.add_row()
        values = [lesson_id, topic, focus, task, feedback]
        for idx, value in enumerate(values):
            cell = row.cells[idx]
            cell.text = short_text(value, 300 if idx >= 2 else 120)
            set_cell_width(cell, widths[idx])
    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(0)
                for run in p.runs:
                    run.font.name = "Calibri"
                    run.font.size = Pt(8)
    return table


def add_phase_table(doc, rows):
    table = doc.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    table.autofit = False
    headers = ["Phase", "Time", "Teacher Moves", "Student Work", "Evidence to Check"]
    widths = [1000, 750, 2650, 2450, 2100]
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        cell.text = header
        set_cell_width(cell, widths[idx])
        set_cell_shading(cell, "E6EEF7")
    for phase, time, teacher, student, evidence in rows:
        row = table.add_row()
        for idx, value in enumerate([phase, time, teacher, student, evidence]):
            cell = row.cells[idx]
            cell.text = value
            set_cell_width(cell, widths[idx])
    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(0)
                for run in p.runs:
                    run.font.name = "Calibri"
                    run.font.size = Pt(8)
    return table


def lesson_plan_doc(unit, lesson):
    strand = STRANDS.get(unit["unit"], {"code": "", "title": unit["title"], "focus": unit["title"]})
    expectations = section_text(lesson, r"expectations")
    lesson_page = section_text(lesson, r"^Lesson -")
    hands_on = section_text(lesson, r"Hands On")
    consolidation = section_text(lesson, r"Consolidation")
    homework = section_text(lesson, r"Homework")
    codes = expectation_codes(expectations)
    media, files = lesson_materials(lesson)
    doc = template_document("lesson")

    prior = (
        f"Students should bring forward vocabulary and process knowledge from earlier SBI3U work in {unit['title']}. "
        f"The opening prompt should connect the previous lesson or KWL entry to {lesson['title']} so students can name what they already understand before the iSpring lesson and activity."
    )
    expectation_sections = extract_expectation_sections(expectations)
    goals = learning_goals(expectations or lesson_page, lesson["title"])
    criteria = lesson_success_criteria(lesson["title"], codes)

    overview = doc.tables[0]
    rows = overview.rows
    set_text(rows[0].cells[0], "Lesson Plan", bold=True)
    set_text(rows[0].cells[2], "Subject: SBI3U Biology", bold=True)
    set_row_text(rows[1], f"Lesson Name: {lesson['title']}")
    set_row_text(rows[2], f"Unit of Study: {unit['title']}")
    prior_text = (
        "PRIOR KNOWLEDGE What do my students already know? What key questions or strategies will activate prior knowledge? "
        "What connections can I help students make?\n" + prior
    )
    set_row_text(rows[3], prior_text)
    expectation_text = "\n".join(
        [
            "CURRICULUM EXPECTATIONS",
            "OVERALL",
            *expectation_sections["overall"],
            "SPECIFIC",
            *expectation_sections["specific"],
        ]
    )
    set_row_text(rows[4], expectation_text)
    goal_text = "\n".join(["LEARNING GOALS", "LEARNING GOALS What do I want students to know and be able to do?", *goals])
    set_row_text(rows[5], goal_text)
    success_text = "\n".join(
        [
            "SUCCESS CRITERIA(S) How will students know they attained the goal? What tools can they use to self-assess?",
            *criteria,
        ]
    )
    set_row_text(rows[6], success_text)
    set_text(rows[7].cells[0], "Assessment as Learning", bold=True)
    set_text(rows[7].cells[1], "Assessment for Learning", bold=True)
    set_text(rows[7].cells[2], "Assessment of Learning", bold=True)
    set_text(rows[8].cells[0], "☑ Observation\n☑ Anecdotal Notes\n☑ Exit Card\n☑ Self-Assessment checklist\n☐ Discussions")
    set_text(rows[8].cells[1], "☑ Worksheets\n☑ Homework\n☑ Strategic Questioning\n☑ Hands On/H5P checks\n☑ Consolidation response")
    set_text(rows[8].cells[2], "☑ Homework submission\n☑ Unit quiz/test/lab/assignment where listed\n☐ Oral Presentations\n☐ Portfolio")
    set_row_text(rows[9], "What will I do?\nConfer\nObserve\nGrade")
    set_row_text(
        rows[10],
        "Accommodations: How will you change the lesson to meet the needs of individual students?\n"
        "Chunk the iSpring and lesson text into checkpoints; provide vocabulary support, diagrams, or sentence frames; "
        "allow oral rehearsal before written responses; and use Hands On/consolidation evidence to reteach one misconception.",
    )
    set_row_text(
        rows[11],
        "Materials and Resources: What do I need to prepare in advance? Equipment? Student resources? Teacher resources?\n"
        + "\n".join((media + files)[:10] or ["Localized lesson page, Lesson Expectations, Hands On, Consolidation, and Homework resources."])
        + f"\nSource status: reconstructed_from_moodle_content. Generated from localized SBI3U Moodle resources on {GENERATED_AT}.",
    )

    lesson_table = doc.tables[1]
    set_text(lesson_table.rows[0].cells[0], "DELIVERING THE LESSON", bold=True)
    lesson_rows = lesson_table.rows
    set_text(lesson_rows[1].cells[0], "Timing\n5-8\nminutes")
    set_text(lesson_rows[1].cells[1], "Grouping")
    set_text(
        lesson_rows[1].cells[4],
        "Minds On!\nOpen with a retrieval prompt tied to the previous lesson or unit KWL. "
        "Name the learning target and vocabulary students need before media viewing.",
    )
    set_text(lesson_rows[1].cells[5], "Materials/Resources\nKWL chart\nLesson Expectations page")
    phase_rows = [
        (
            2,
            "Timing\n20-30\nminutes",
            "Action!\nUse the Lesson page and iSpring module to model the core idea. Pause after each major concept for a short verbal or written check.",
            "Materials/Resources\nLocalized lesson page\niSpring module",
        ),
        (
            3,
            "Timing\n15-25\nminutes",
            "Action!\nFrame the Hands On activity or H5P as formative practice. Circulate, name misconceptions, and connect the task to success criteria.",
            "Materials/Resources\n" + ("\n".join(media[:4]) if media else "Hands On page and activity resources"),
        ),
        (
            4,
            "Timing\n10-15\nminutes",
            "Consolidation!\nUse the video, exit slip, or consolidation prompt to check whether students can explain the concept without copying. Confirm the homework file and submission location.",
            "Materials/Resources\n" + ("\n".join(files[:6]) if files else "Homework and consolidation resources"),
        ),
    ]
    for row_index, timing, action, materials in phase_rows:
        row = lesson_rows[row_index]
        set_text(row.cells[0], timing)
        set_text(row.cells[1], "W\n❐")
        set_text(row.cells[2], "S\n❐")
        set_text(row.cells[3], "I\n❐")
        set_text(row.cells[4], action)
        set_text(row.cells[5], materials)

    notes = doc.tables[2]
    set_text(
        notes.rows[0].cells[0],
        "Notes:\nSupport: pre-teach vocabulary, chunk media, provide labelled diagrams or sentence frames, and allow oral rehearsal before written explanation.\n"
        "Extension: connect the biology concept to a health, environmental, technological, or ethical application where the lesson expectations support it.\n"
        "Teacher reflection: identify the misconception that should launch the next lesson.",
    )
    return doc


def unit_plan_doc(unit):
    strand = STRANDS.get(unit["unit"], {"code": "", "title": unit["title"], "focus": unit["title"]})
    lessons = unit.get("lessons") or []
    all_codes = []
    sequence_rows = []
    for lesson in lessons:
        expectations = section_text(lesson, r"expectations")
        all_codes.extend(expectation_codes(expectations))
        media, files = lesson_materials(lesson)
        focus = short_text(expectations, 220) or f"Develop core understanding for {lesson['title']}."
        task = "; ".join((media[:2] or files[:2] or ["Lesson page, Hands On, and homework resources"]))
        feedback = "Hands On/H5P checks, consolidation evidence, and homework review."
        sequence_rows.append((lesson["id"], lesson["title"], focus, task, feedback))
    codes = sorted(set(all_codes))[:22]

    doc = template_document("unit")
    if doc.paragraphs:
        doc.paragraphs[0].text = f"UNIT {unit['unit']}"
    evaluations = resource_labels(unit.get("unitResources", {}).get("evaluations") or [])
    unit_table = doc.tables[0]
    rows = unit_table.rows
    section_rows = {
        0: "Unit Author",
        5: "Unit Overview",
        14: "Unit Foundation",
        19: "Assessment Plan",
        22: "Unit Details",
        25: "Materials and Resources",
    }
    for row_index, label in section_rows.items():
        set_text(rows[row_index].cells[0], label, bold=True)
    for row_index, label in [(1, "Name:"), (2, "School District:"), (3, "School Name:"), (4, "School City, Province:")]:
        set_row_text(rows[row_index], label)
    for row_index, label, value in [
        (6, "Unit Title Name", ""),
        (7, unit["title"], ""),
        (8, "Unit Summary", ""),
        (9, unit_summary(unit, strand, lessons), ""),
        (10, "Year Level", ""),
        (11, "Grade 11, University Preparation", ""),
        (12, "Approximate Time Needed", ""),
        (13, f"{max(len(lessons) * 3, 6)} instructional hours, plus homework and evaluation time", ""),
        (15, "Targeted Curriculum Expectations", ""),
        (
            16,
            "\n".join(
                [
                    f"{strand['code']} - {strand['title']}",
                    f"Expectation evidence found across lessons: {', '.join(codes) if codes else 'see each Lesson Expectations page.'}",
                    f"Unit focus: {strand['focus']}",
                ]
            ),
            "",
        ),
        (17, "Learning Goals", ""),
        (
            18,
            "\n".join(
                [
                    f"Students will explain and apply major concepts in {unit['title']} using accurate biology vocabulary.",
                    "Students will use lesson evidence, models, diagrams, and investigations to support biological explanations.",
                    "Students will complete Hands On, consolidation, homework, and reflection tasks as preparation for unit evaluation evidence.",
                    "Essential Questions",
                    f"How do the structures, processes, and evidence in {unit['title']} help explain living systems?",
                    "How can students use models, investigations, and evidence to justify a biological explanation?",
                    "How do scientific ideas in this unit connect to health, technology, environment, ethics, or society?",
                ]
            ),
            "",
        ),
    ]:
        set_row_text(rows[row_index], label)
    set_text(rows[20].cells[0], "Assessment as Learning (ASL)", bold=True)
    set_text(rows[20].cells[2], "Assessment for Learning (AFL)", bold=True)
    set_text(rows[20].cells[3], "Assessment of Learning (AoL)", bold=True)
    set_text(rows[21].cells[0], "☑ KWL chart\n☑ Reflection routines\n☑ Learning Log\n☑ Student next-step notes")
    set_text(rows[21].cells[2], "☑ Hands On/H5P checks\n☑ Strategic questioning\n☑ Homework review\n☑ Consolidation responses")
    set_text(
        rows[21].cells[3],
        "☑ Unit quiz/test/lab/assignment where listed\n"
        + ("☑ " + "\n☑ ".join(evaluations[:10]) if evaluations else "☑ Unit quiz, test, lab, or culminating activity listed in the Evaluation section."),
    )
    set_row_text(rows[23], "Lesson and Assessment Outlines", bold=True)
    outlines = []
    for lesson_id, topic, focus, task, feedback in sequence_rows:
        outlines.append(f"{lesson_id} - {topic}\nAssessment: {feedback}\nResources: {task}")
    set_row_text(rows[24], "\n\n".join(outlines))
    set_text(rows[26].cells[0], "Technology", bold=True)
    tech = "\n".join(
        sorted(set(activity_titles(unit, r"ispring|h5p|mp4|video")))[:24]
        or ["Localized lesson pages, iSpring presentations, H5P activities, and consolidation videos listed in the manifest."]
    )
    set_text(rows[26].cells[1], tech)
    set_text(rows[27].cells[0], "Printed", bold=True)
    printed = "\n".join(
        sorted(set(activity_titles(unit, r"docx?|pdf|worksheet|homework|handout|rubric|lab")))[:24]
        or ["Homework, lab, rubric, and activity files listed in the manifest."]
    )
    set_text(rows[27].cells[1], printed)
    set_text(rows[28].cells[0], "Other Resources", bold=True)
    set_text(rows[28].cells[1], "Anecdotal notes of observation\nExit slips and consolidation responses\nStudent self-assessment and reflection\nKWL chart and Learning Log routines where present")
    return doc


def save_doc(doc, relative_path):
    path = COURSE_ROOT / relative_path
    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(path)
    return path.stat().st_size


def file_record(label, relative_path, category, role):
    path = COURSE_ROOT / relative_path
    preview = f"previews-html/{relative_path}.html"
    return {
        "label": label,
        "type": "docx",
        "category": category,
        "role": role,
        "path": relative_path.replace("\\", "/"),
        "bytes": path.stat().st_size,
        "source": "reconstructed_from_moodle_content",
        "sourceStatus": "locally_authored_from_course_materials",
        "teacherUse": "Teacher planning aid reconstructed from localized course content; not an original Moodle teacher packet.",
        "previewPath": preview.replace("\\", "/"),
    }


def main():
    manifest = json.loads(MANIFEST_PATH.read_text(encoding="utf-8"))
    unit_count = 0
    lesson_count = 0
    for unit in manifest.get("units") or []:
        unit_rel = f"plans/generated/unit-plans/U{unit['unit']:02d}-unit-plan.docx"
        save_doc(unit_plan_doc(unit), unit_rel)
        unit["unitPlan"] = file_record(f"Unit Plan - {unit['title']}", unit_rel, "unit_plan", "unit_plan")
        unit_count += 1
        for lesson in unit.get("lessons") or []:
            lesson_rel = f"plans/generated/lesson-plans/{lesson['id']}-lesson-plan.docx"
            save_doc(lesson_plan_doc(unit, lesson), lesson_rel)
            lesson["lessonPlan"] = file_record(
                f"Lesson Plan - Unit {unit['unit']} Lesson {lesson['lesson']}",
                lesson_rel,
                "lesson_plan",
                "lesson_plan",
            )
            lesson["resourceCounts"] = {**(lesson.get("resourceCounts") or {}), "lessonPlan": 1}
            lesson_count += 1

    manifest["generatedAt"] = GENERATED_AT
    manifest["sourceAudit"] = {
        **(manifest.get("sourceAudit") or {}),
        "teacherPlanGeneration": {
            "generatedAt": GENERATED_AT,
            "format": "docx",
            "unitPlans": unit_count,
            "lessonPlans": lesson_count,
            "source": "reconstructed_from_moodle_content",
            "basis": [
                "localized Moodle Lesson Expectations pages",
                "localized Moodle Lesson, Hands On, Consolidation, and Homework sections",
                "localized iSpring, H5P, video, worksheet, exit slip, KWL/reflection, and unit Evaluation records",
                "Ontario Science curriculum guidance included in course texts",
                "McGraw-Hill Ryerson Biology 11 textbook metadata included in course texts",
            ],
            "boundary": "Teacher planning aids only; no answer keys, rubrics, textbook excerpts, or Moodle-original teacher packet documents were created.",
        },
    }
    MANIFEST_PATH.write_text(json.dumps(manifest, indent=2, ensure_ascii=False) + "\n", encoding="utf-8")
    print(json.dumps({"course": "SBI3U", "unitPlans": unit_count, "lessonPlans": lesson_count, "generatedAt": GENERATED_AT}, indent=2))


if __name__ == "__main__":
    main()
