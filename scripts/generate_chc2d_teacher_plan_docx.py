import json
import re
from datetime import datetime, timezone
from html import unescape
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


WORKSPACE_ROOT = Path("D:/工作文件/SUNNYBROOK")
REPO_ROOT = WORKSPACE_ROOT / "ossd-course-portal"
COURSE_ROOT = WORKSPACE_ROOT / "courseware" / "CHC2D"
MANIFEST_PATH = COURSE_ROOT / "course-manifest.json"
REPORT_PATH = REPO_ROOT / "deployment" / "CHC2D-docx-plans-report.json"
TEMPLATE_ROOT = REPO_ROOT / "templates" / "teacher-plans"
LESSON_PLAN_TEMPLATE = TEMPLATE_ROOT / "mdm4u-style-lesson-plan-template.docx"
UNIT_PLAN_TEMPLATE = TEMPLATE_ROOT / "mdm4u-style-unit-plan-template.docx"
GENERATED_AT = datetime.now(timezone.utc).replace(microsecond=0).isoformat()


GENERATED_SOURCE = (
    "Locally authored from the St.Mary CHC2D Moodle course, CHC2D lesson expectation pages, "
    "localized course files/media, course outline, the Think History core textbook record, and "
    "the Ontario Canadian and World Studies curriculum expectations."
)
GENERATED_NOTE = (
    "Generated teacher planning aid. The Moodle export did not provide complete lesson plan DOCX "
    "files for this course, so these plans were reconstructed in the local MDM4U-style teacher "
    "planning format from indexed course content."
)


STRANDS = {
    "A": {
        "title": "Historical Inquiry and Skill Development",
        "focus": "historical inquiry, evidence, historical thinking concepts, communication, and transferable skills.",
    },
    "B": {
        "title": "Canada, 1914-1929",
        "focus": "World War I, the home front, postwar change, identity, citizenship, and Canadian society in the 1920s.",
    },
    "C": {
        "title": "Canada, 1929-1945",
        "focus": "the Great Depression, World War II, social and political change, conflict, cooperation, and identity.",
    },
    "D": {
        "title": "Canada, 1945-1982",
        "focus": "postwar prosperity and tension, the Cold War, social movements, Quebec/national unity, and changing identity.",
    },
    "E": {
        "title": "Canada, 1982 to the Present",
        "focus": "constitutional change, rights, social and economic change, global relationships, reconciliation, and contemporary identity.",
    },
}


OVERALL_BY_UNIT = {
    1: [
        "A1. Historical Inquiry: use the historical inquiry process and the concepts of historical thinking when investigating aspects of Canadian history since 1914.",
        "A2. Developing Transferable Skills: apply skills developed through historical investigation in everyday contexts and identify related careers.",
        "B1. Social, Economic, and Political Context: describe key events, trends, and developments between 1914 and 1929 and assess their significance for groups in Canada.",
        "B2. Communities, Conflict, and Cooperation: analyse interactions within and between communities in Canada and between Canada and the international community from 1914 to 1929.",
        "B3. Identity, Citizenship, and Heritage: explain how individuals, organizations, and social changes between 1914 and 1929 contributed to Canadian identity, citizenship, and heritage.",
    ],
    2: [
        "A1. Historical Inquiry: use the historical inquiry process and the concepts of historical thinking when investigating aspects of Canadian history since 1914.",
        "A2. Developing Transferable Skills: apply skills developed through historical investigation in everyday contexts and identify related careers.",
        "C1. Social, Economic, and Political Context: describe key events, trends, and developments between 1929 and 1945 and assess their impact on groups in Canada.",
        "C2. Communities, Conflict, and Cooperation: analyse interactions within and between communities in Canada and between Canada and the international community from 1929 to 1945.",
        "C3. Identity, Citizenship, and Heritage: explain how individuals, groups, events, and major international events contributed to Canadian identity, citizenship, and heritage from 1929 to 1945.",
    ],
    3: [
        "A1. Historical Inquiry: use the historical inquiry process and the concepts of historical thinking when investigating aspects of Canadian history since 1914.",
        "A2. Developing Transferable Skills: apply skills developed through historical investigation in everyday contexts and identify related careers.",
        "D1. Social, Economic, and Political Context: describe key social, economic, and political developments in Canada between 1945 and 1982 and assess their significance.",
        "D2. Communities, Conflict, and Cooperation: analyse experiences of and interactions between communities in Canada and interactions between Canada and the international community from 1945 to 1982.",
        "D3. Identity, Citizenship, and Heritage: analyse how events, individuals, and groups, including Aboriginal peoples, Quebecois, and immigrants, contributed to identity, citizenship, and heritage from 1945 to 1982.",
    ],
    4: [
        "A1. Historical Inquiry: use the historical inquiry process and the concepts of historical thinking when investigating aspects of Canadian history since 1914.",
        "A2. Developing Transferable Skills: apply skills developed through historical investigation in everyday contexts and identify related careers.",
        "E1. Social, Economic, and Political Context: describe key events, trends, and developments in Canada from 1982 to the present and assess their significance.",
        "E2. Communities, Conflict, and Cooperation: analyse significant interactions within and between communities in Canada and between Canada and the international community from 1982 to the present.",
        "E3. Identity, Citizenship, and Heritage: analyse how individuals, groups, organizations, and events have contributed to Canadian identity, citizenship, and heritage from 1982 to the present.",
    ],
    5: [
        "A1. Historical Inquiry: use the historical inquiry process and historical thinking concepts to investigate Canadian history since 1914.",
        "A2. Developing Transferable Skills: apply historical investigation skills in everyday contexts and identify related careers.",
        "B1-B3, C1-C3, D1-D3, E1-E3: synthesize evidence from the chronological course strands to demonstrate understanding of change, continuity, cause, consequence, perspective, and significance.",
    ],
}


UNIT_INTENT = {
    1: {
        "big_idea": "The First World War and the 1920s reshaped Canada's political autonomy, economy, social roles, and sense of identity.",
        "teacher_intent": "Move students beyond a list of events by having them explain causes, consequences, perspectives, and historical significance for different communities in Canada.",
        "essential": [
            "Why did Canada become involved in World War I, and how did that involvement affect Canadian society?",
            "How did technology, battle experiences, financing, and the home front change the lives of people in Canada?",
            "In what ways did postwar social change contribute to emerging Canadian identity, citizenship, and conflict?",
        ],
    },
    2: {
        "big_idea": "The Great Depression and World War II produced economic hardship, political responses, conflict, and lasting changes in citizenship and identity.",
        "teacher_intent": "Help students connect economic crisis, global conflict, government response, human rights, and home-front experiences through evidence-based historical reasoning.",
        "essential": [
            "How did the Great Depression affect people and government in Canada?",
            "What caused World War II, and how did Canada contribute to the conflict at home and abroad?",
            "How should students evaluate difficult histories such as the Holocaust, internment, and wartime decision making?",
        ],
    },
    3: {
        "big_idea": "Between 1945 and 1982, Canada changed through postwar prosperity, Cold War pressures, social movements, immigration, Quebec nationalism, and constitutional debates.",
        "teacher_intent": "Teach students to interpret postwar Canada as a period of rapid change where identity, rights, economic policy, and national unity were contested.",
        "essential": [
            "How did postwar economic and political changes affect different groups in Canada?",
            "How did the 1950s, 1960s, and 1970s alter ideas about rights, identity, and social responsibility?",
            "How did Trudeau-era policy and national unity debates shape modern Canada?",
        ],
    },
    4: {
        "big_idea": "From 1982 to the present, Canada has continued to debate rights, identity, reconciliation, economic change, and its role in an interconnected world.",
        "teacher_intent": "Help students connect contemporary Canada to historical roots, especially the Charter, regional change, social movements, and Truth and Reconciliation.",
        "essential": [
            "How did the Charter era shape citizenship, rights, and debates over Canadian identity?",
            "How did major events in the 1980s, 1990s, and 2000s affect people living in Canada?",
            "What responsibilities follow from learning about residential schools and Truth and Reconciliation?",
        ],
    },
    5: {
        "big_idea": "The final evaluation asks students to synthesize course knowledge, historical inquiry skills, and communication into culminating evidence of learning.",
        "teacher_intent": "Guide students through planning, evidence selection, revision, and final submission while checking readiness against course expectations.",
        "essential": [
            "What evidence best demonstrates historical thinking across the course?",
            "How can students communicate a defensible historical interpretation clearly?",
            "How can reflection and revision improve final course evidence?",
        ],
    },
}


TOPIC_DETAILS = [
    (
        ["causes of world war 1", "causes of world war i"],
        {
            "application": "explain the long- and short-term causes of World War I and connect them to Canada's relationship with Britain",
            "question": "Which cause best explains why a European conflict became a Canadian historical turning point?",
        },
    ),
    (
        ["canada become involved"],
        {
            "application": "explain Canada's automatic entry into the war and assess the significance of enlistment, imperial ties, and public opinion",
            "question": "Was Canada's entry into World War I a choice, an obligation, or both?",
        },
    ),
    (
        ["technological advances", "world war 1"],
        {
            "application": "analyse how military technology changed combat, casualties, strategy, and historical memory",
            "question": "How did technology alter both the experience and consequences of war?",
        },
    ),
    (
        ["major battles", "vimy"],
        {
            "application": "evaluate the historical significance of major battles using evidence, consequence, and perspective",
            "question": "Why do some battles become national symbols while others remain less remembered?",
        },
    ),
    (
        ["paying for the war"],
        {
            "application": "connect wartime finance, taxation, victory bonds, and government policy to experiences on the home front",
            "question": "Who paid the costs of war, and how were those costs justified?",
        },
    ),
    (
        ["roles of women"],
        {
            "application": "analyse how women's wartime work, suffrage, and reform activism affected citizenship and social change",
            "question": "Did wartime change create lasting equality, or only temporary opportunity?",
        },
    ),
    (
        ["social changes after world war"],
        {
            "application": "explain postwar social change, immigration, labour unrest, and cultural shifts in relation to identity and citizenship",
            "question": "How did the 1920s reveal both progress and conflict in Canadian society?",
        },
    ),
    (
        ["world war 2 timeline"],
        {
            "application": "sequence major events of World War II and use chronology to explain cause, consequence, and turning points",
            "question": "How does a timeline help historians decide which events were turning points?",
        },
    ),
    (
        ["causes of world war 2"],
        {
            "application": "explain how the Treaty of Versailles, depression, fascism, appeasement, and militarism contributed to World War II",
            "question": "Could World War II have been prevented, or were its causes already too deeply connected?",
        },
    ),
    (
        ["great depression"],
        {
            "application": "assess how the Great Depression affected employment, families, government response, and regional experiences in Canada",
            "question": "How did the Great Depression change what Canadians expected government to do?",
        },
    ),
    (
        ["world war 2 technology"],
        {
            "application": "analyse how wartime technologies affected military strategy, civilian life, and postwar society",
            "question": "Which wartime technologies produced the most lasting historical consequences?",
        },
    ),
    (
        ["holocaust"],
        {
            "application": "examine the Holocaust using respectful historical inquiry, evidence, human rights language, and attention to responsibility",
            "question": "How should historians and citizens respond to evidence of genocide and discrimination?",
        },
    ),
    (
        ["pacific war"],
        {
            "application": "analyse Canada's connections to the Pacific War and the impact of wartime racism and internment on citizenship",
            "question": "How did war in the Pacific expose conflicts between security, rights, and identity?",
        },
    ),
    (
        ["home front"],
        {
            "application": "explain how rationing, labour, propaganda, women's work, and internment shaped Canada's home-front experience",
            "question": "How did people in Canada experience World War II differently depending on identity and community?",
        },
    ),
    (
        ["post war years"],
        {
            "application": "connect postwar prosperity, suburbanization, immigration, consumer culture, and international commitments to Canadian society",
            "question": "Why did postwar prosperity not affect all Canadians in the same way?",
        },
    ),
    (
        ["1950s"],
        {
            "application": "analyse political and economic change in the 1950s through evidence about prosperity, culture, Cold War pressure, and social expectations",
            "question": "What made the 1950s appear stable, and where were tensions already visible?",
        },
    ),
    (
        ["1960s"],
        {
            "application": "explain how youth culture, rights movements, immigration policy, Quebec nationalism, and social programs changed Canada in the 1960s",
            "question": "Why are the 1960s often remembered as a decade of change?",
        },
    ),
    (
        ["1970s"],
        {
            "application": "analyse economic challenges, social change, Indigenous activism, Quebec politics, and Canada's international role in the 1970s",
            "question": "How did economic pressure and identity debates reshape Canada in the 1970s?",
        },
    ),
    (
        ["trudeau", "national unity"],
        {
            "application": "evaluate Trudeau-era policy, bilingualism, multiculturalism, patriation, and national unity debates",
            "question": "How did debates over national unity change ideas about Canada?",
        },
    ),
    (
        ["after ww2 review"],
        {
            "application": "synthesize postwar evidence to review continuity, change, cause, consequence, and historical significance",
            "question": "Which postwar change most clearly shaped modern Canada?",
        },
    ),
    (
        ["1982-present timeline"],
        {
            "application": "sequence events from 1982 to the present and use the Charter as an anchor for rights, identity, and citizenship debates",
            "question": "How does the Charter era change the way Canadians discuss rights and citizenship?",
        },
    ),
    (
        ["1980s"],
        {
            "application": "analyse the 1980s through constitutional change, free trade, social issues, economic shifts, and public figures",
            "question": "Which 1980s developments still shape Canada today?",
        },
    ),
    (
        ["1990s"],
        {
            "application": "explain the 1990s through recession, globalization, constitutional debates, Quebec referendum, and changing social identities",
            "question": "How did the 1990s test Canadian unity and identity?",
        },
    ),
    (
        ["2000s"],
        {
            "application": "connect contemporary events, technology, security, globalization, and changing communities to Canada's evolving identity",
            "question": "How did the early twenty-first century change Canada's opportunities and responsibilities?",
        },
    ),
    (
        ["truth", "reconciliation"],
        {
            "application": "examine residential schools, survivor testimony, the TRC, Calls to Action, and responsibilities for reconciliation",
            "question": "What does reconciliation require from governments, institutions, and individuals?",
        },
    ),
    (
        ["culminating", "summative"],
        {
            "application": "select evidence from across the course and communicate a defensible historical argument or final response",
            "question": "What evidence best proves course learning across knowledge, inquiry, communication, and application?",
        },
    ),
]


def clean_text(value):
    text = str(value or "")
    text = re.sub(r"<script[\s\S]*?</script>", " ", text, flags=re.I)
    text = re.sub(r"<style[\s\S]*?</style>", " ", text, flags=re.I)
    text = re.sub(r"<br\s*/?>", "\n", text, flags=re.I)
    text = re.sub(r"</(p|div|li|tr|h[1-6])>", "\n", text, flags=re.I)
    text = re.sub(r"<[^>]+>", " ", text)
    text = unescape(text)
    text = text.replace("\u00a0", " ")
    text = re.sub(r"([A-E])\s+(\d(?:\.\d+)?)", r"\1\2", text)
    return re.sub(r"\s+", " ", text).strip()


def short_text(value, limit=360):
    text = clean_text(value)
    if len(text) <= limit:
        return text
    clipped = text[:limit]
    cut = max(clipped.rfind(". "), clipped.rfind("; "), clipped.rfind(", "))
    if cut > 140:
        clipped = clipped[: cut + 1]
    return clipped.strip() + "..."


def unique(items):
    out = []
    seen = set()
    for item in items:
        value = clean_text(item)
        key = value.lower()
        if value and key not in seen:
            seen.add(key)
            out.append(value)
    return out


def read_local_text(relative_path):
    if not relative_path:
        return ""
    path = COURSE_ROOT / relative_path
    if not path.exists():
        return ""
    return clean_text(path.read_text(encoding="utf-8", errors="ignore"))


def find_section(lesson, pattern):
    regex = re.compile(pattern, re.I)
    for section in lesson.get("bookSections") or []:
        label = section.get("sectionLabel") or section.get("label") or ""
        if regex.search(label):
            return section
    return None


def section_text(lesson, pattern):
    section = find_section(lesson, pattern)
    if not section:
        return ""
    return read_local_text(section.get("path")) or clean_text(section.get("textPreview"))


def strip_repeated_heading(text, lesson):
    out = clean_text(text)
    title = re.escape(str(lesson.get("title") or ""))
    patterns = [
        rf"^Lesson\s+\d+\s*:\s*{title}\s+Lesson Expectations",
        rf"^Lesson\s+\d+\s*:\s*{title}",
        r"^CHC2D\s+Unit\s+\d+\s+Lesson\s+\d+",
    ]
    changed = True
    while changed:
        changed = False
        for pattern in patterns:
            new = re.sub(pattern, "", out, flags=re.I).strip()
            if new != out:
                out = clean_text(new)
                changed = True
    return out


def extract_between(text, start_labels, end_labels):
    lower = text.lower()
    start = -1
    label_len = 0
    for label in start_labels:
        idx = lower.find(label.lower())
        if idx >= 0 and (start < 0 or idx < start):
            start = idx
            label_len = len(label)
    if start < 0:
        return ""
    start += label_len
    if start < len(text) and text[start] == ":":
        start += 1
    end = len(text)
    for label in end_labels:
        idx = lower.find(label.lower(), start)
        if idx >= 0:
            end = min(end, idx)
    return clean_text(text[start:end])


def split_items(block, limit=10):
    text = clean_text(block)
    if not text:
        return []
    pieces = re.split(r"(?=\b[A-E]\s*\d(?:\.\d+)?\b)|(?<=\.)\s+(?=[A-Z][a-z])", text)
    return [short_text(piece, 300).rstrip(".") for piece in pieces if len(clean_text(piece)) > 12][:limit]


def expectation_codes(text):
    codes = re.findall(r"\b[A-E]\s*\d(?:\.\d+)?\b", text or "")
    return unique(code.replace(" ", "") for code in codes)[:24]


def expectation_for(lesson):
    raw = strip_repeated_heading(section_text(lesson, r"expectations"), lesson)
    overall = extract_between(
        raw,
        ["Overall Expectations", "Overall Expectation"],
        ["Specific Lesson Expectations", "Specific Expectations", "Learning Goals", "Success Criteria"],
    )
    specific = extract_between(
        raw,
        ["Specific Lesson Expectations", "Specific Expectations", "Specific Lesson Expectation"],
        ["Learning Goals", "Success Criteria"],
    )
    goals = extract_between(raw, ["Learning Goals", "Learning Goal"], ["Success Criteria"])
    success = extract_between(raw, ["Success Criteria", "Success Criterion"], [])
    if not overall and raw:
        overall = raw
    if not specific and raw:
        specific = raw
    return {
        "raw": raw,
        "codes": expectation_codes(raw),
        "overall": split_items(overall, 8),
        "specific": split_items(specific, 10),
        "goals": split_items(goals, 6),
        "success": split_items(success, 6),
    }


def resource_items(lesson):
    return [item for item in (lesson.get("downloads") or []) + (lesson.get("textExports") or []) if item]


def item_label(item):
    return clean_text(item.get("label") or item.get("title") or Path(str(item.get("path") or "")).name)


def resource_labels(items):
    return unique(item_label(item) for item in items)


def lesson_materials(lesson):
    media = []
    files = []
    for item in lesson.get("ispring") or []:
        media.append(f"iSpring lesson presentation: {item.get('label') or 'localized presentation'}")
    for item in lesson.get("h5p") or []:
        media.append(f"H5P/interactive activity: {item.get('label') or 'localized H5P activity'}")
    for item in resource_items(lesson):
        label = item_label(item)
        haystack = " ".join(str(item.get(k, "")) for k in ("type", "category", "role", "path", "label"))
        if re.search(r"mp4|video", haystack, re.I):
            media.append(f"Video/consolidation media: {label}")
        elif re.search(r"h5p|interactive", haystack, re.I):
            media.append(f"Interactive practice: {label}")
        elif re.search(r"docx?|pdf|pptx?|xlsx?|worksheet|rubric|assignment|kwl|reflection|learning log|exit", haystack, re.I):
            files.append(label)
    return unique(media), unique(files)


def lesson_activity_summary(lesson):
    pieces = []
    for label, pattern, limit in [
        ("Lesson page", r"^Lesson$", 230),
        ("Hands On", r"Hands On", 190),
        ("Consolidation", r"Consolidation", 190),
        ("Homework", r"Homework", 190),
    ]:
        text = section_text(lesson, pattern)
        if text:
            pieces.append(f"{label}: {short_text(text, limit)}")
    return "\n".join(pieces) or "Use the localized Moodle activity page and attached culminating/final assignment files."


def lesson_topic_detail(lesson):
    title = clean_text(lesson.get("title", "")).lower()
    for keys, detail in TOPIC_DETAILS:
        if any(key in title for key in keys):
            return detail
    return {
        "application": f"use evidence to explain the historical significance of {lesson.get('title', 'the lesson topic')}",
        "question": f"What makes {lesson.get('title', 'this topic')} historically significant?",
    }


def lower_initial(text):
    text = clean_text(text)
    if not text:
        return text
    return text[0].lower() + text[1:]


def prior_knowledge(unit, lesson):
    unit_no = int(unit["unit"])
    lesson_no = int(lesson.get("lesson", 1))
    if unit_no == 5:
        return (
            "Students bring forward evidence and historical thinking practice from Units 1-4. Begin with a checklist of "
            "course strands, major periods, and final-evaluation requirements so students can identify what evidence they "
            "already have and what still needs revision."
        )
    if lesson_no == 1:
        return (
            f"Students begin {unit['title']} by activating prior knowledge through a KWL prompt, timeline prediction, "
            "image/source observation, or short discussion. Ask students what they already know, what evidence would verify it, "
            "and whose perspective might be missing."
        )
    prev = next((item for item in unit.get("lessons") or [] if item.get("lesson") == lesson_no - 1), None)
    previous = f"Unit {unit['unit']} Lesson {prev['lesson']} ({prev['title']})" if prev else "the previous lesson"
    return (
        f"Students build from {previous}. Start by revisiting one key event, source, or historical-thinking concept from "
        f"the previous lesson, then connect it to {lesson.get('title')} through a cause/consequence, continuity/change, "
        "or historical perspective question."
    )


def lesson_learning_goals(unit, lesson, exp):
    if exp["goals"]:
        return exp["goals"]
    detail = lesson_topic_detail(lesson)
    goals = [
        f"Students will understand how {lesson['title']} fits into the larger history of {unit['title']}.",
        f"Students will learn to {detail['application']}.",
        "Students will practise historical inquiry by using evidence, context, and historical thinking vocabulary in their responses.",
    ]
    for item in exp["specific"][:2]:
        goals.append("Students will connect their practice to this expectation: " + lower_initial(item) + ".")
    if lesson.get("ispring"):
        goals.append("Students will use the iSpring presentation to organize notes and check for main ideas before independent work.")
    return unique(goals)


def lesson_success_criteria(unit, lesson, exp):
    if exp["success"]:
        return exp["success"]
    detail = lesson_topic_detail(lesson)
    return unique(
        [
            "I can describe the key historical events, people, or ideas from the lesson using accurate vocabulary.",
            "I can support an interpretation with specific evidence from the lesson page, presentation, video, reading, or worksheet.",
            f"I can {detail['application']} in a short written, oral, or interactive response.",
            "I can identify at least one cause, consequence, continuity, change, perspective, or historically significant detail.",
            "I can complete the assigned homework/submission evidence using the correct course file or dropbox instructions.",
        ]
    )


def assessment_lists(unit, lesson):
    labels = " ".join(resource_labels(resource_items(lesson))).lower()
    asl = ["[x] Observation", "[x] Anecdotal notes", "[x] Student self-check against success criteria"]
    if re.search(r"kwl|reflection|learning log|exit", labels):
        asl.append("[x] KWL/reflection/exit evidence")
    else:
        asl.append("[ ] KWL/reflection/exit evidence")
    afl = ["[x] Strategic questioning", "[x] Guided note or worksheet check", "[x] Homework preparation check"]
    if lesson.get("h5p") or section_text(lesson, r"Hands On"):
        afl.append("[x] Hands On/H5P or practice check")
    if lesson.get("ispring"):
        afl.append("[x] iSpring note-taking checkpoint")
    unit_evals = resource_labels((unit.get("unitResources") or {}).get("evaluations") or [])
    aol = [f"[x] {label}" for label in unit_evals[:6]]
    if not aol and int(unit["unit"]) == 5:
        aol = ["[x] Culminating assignment", "[x] Final summative evidence"]
    if not aol:
        aol = ["[ ] No formal AoL in this lesson unless identified by the Moodle Evaluation area"]
    return "\n".join(asl), "\n".join(afl), "\n".join(aol)


def teacher_assessment_action(unit, lesson):
    if int(unit["unit"]) == 5:
        return (
            "Confer with students about topic choice, evidence quality, organization, and final-evaluation readiness.\n"
            "Observe planning/revision habits and provide descriptive feedback before final submission.\n"
            "Grade only the Moodle-identified culminating/final assessment products."
        )
    return (
        "Confer with students during source analysis, note-taking, and worksheet completion.\n"
        "Observe whether students use historical evidence rather than unsupported opinion.\n"
        "Give descriptive feedback tied to the learning goal and success criteria.\n"
        "Grade only when the task appears in the Evaluation/AoL area; otherwise use work as practice evidence."
    )


def accommodations_for(lesson):
    title = clean_text(lesson.get("title", ""))
    return (
        "Chunk historical readings and videos into short checkpoints; pre-teach vocabulary, dates, and names; provide a timeline "
        "or cause/consequence organizer; allow oral rehearsal before written responses; offer sentence frames for evidence and "
        "historical significance; provide extra time for reading, note-taking, and uploading; use captions/transcripts where available; "
        f"extend by asking students to compare {title} with another Canadian or global historical example."
    )


def materials_for(lesson):
    media, files = lesson_materials(lesson)
    lines = []
    lines.extend(media[:8])
    lines.extend(files[:10])
    if not lines:
        lines.append("Localized Moodle lesson page and teacher-selected historical source/context materials.")
    return "\n".join(unique(lines))


def unit_codes(unit):
    codes = []
    for lesson in unit.get("lessons") or []:
        codes.extend(expectation_for(lesson)["codes"])
    return unique(codes)


def unit_learning_goals(unit):
    intent = UNIT_INTENT[int(unit["unit"])]
    if int(unit["unit"]) == 5:
        return [
            "Students will synthesize course evidence across the historical periods studied in CHC2D.",
            "Students will demonstrate historical inquiry, communication, and application in culminating/final assessment evidence.",
            "Students will revise final products using task criteria, source evidence, and teacher feedback.",
            *intent["essential"],
        ]
    return [
        f"Students will explain the historical significance of {unit['title']} using evidence and historical thinking concepts.",
        "Students will analyse cause and consequence, continuity and change, historical perspective, and ethical dimensions where appropriate.",
        "Students will communicate historical interpretations through guided notes, practice activities, consolidation responses, and homework submissions.",
        *intent["essential"],
    ]


def unit_summary(unit):
    lessons = unit.get("lessons") or []
    intent = UNIT_INTENT[int(unit["unit"])]
    sequence = ", ".join(lesson.get("title", "") for lesson in lessons)
    return (
        f"{intent['big_idea']}\n\n"
        f"Instructional sequence: {sequence}.\n\n"
        f"Teacher planning focus: {intent['teacher_intent']}\n\n"
        "Students work through localized Moodle lesson pages, iSpring presentations where present, Hands On/H5P practice, "
        "video/readings in consolidation, homework submission files, reflection routines, and Moodle Evaluation tasks."
    )


def unit_outline(unit):
    chunks = []
    for lesson in unit.get("lessons") or []:
        exp = expectation_for(lesson)
        detail = lesson_topic_detail(lesson)
        target = "; ".join(exp["specific"][:2]) or "; ".join(exp["overall"][:2]) or "Use the localized Lesson Expectations page for exact wording."
        media, files = lesson_materials(lesson)
        evidence = []
        if media:
            evidence.append("media/practice: " + "; ".join(media[:3]))
        if files:
            evidence.append("files: " + "; ".join(files[:4]))
        evidence.append("assessment: lesson practice, consolidation checks, and homework submission where listed")
        chunks.append(
            f"{lesson['id']} - {lesson['title']}\n"
            f"Historical question: {detail['question']}\n"
            f"Target: {target}\n"
            f"Teacher plan: activate prior knowledge, model historical evidence use, guide the Moodle lesson sequence, and close with a reflective or submission-based check.\n"
            + "\n".join(evidence)
        )
    for item in (unit.get("unitResources") or {}).get("evaluations") or []:
        chunks.append(
            f"Assessment of Learning - {item_label(item)}\n"
            "Students apply unit learning in a Moodle-identified assessed task. Feedback should address knowledge, historical inquiry/evidence, communication, and application."
        )
    return "\n\n".join(chunks)


def unit_technology(unit):
    lines = []
    for lesson in unit.get("lessons") or []:
        media, _files = lesson_materials(lesson)
        for item in media:
            lines.append(f"{lesson['id']} - {item}")
    return "\n".join(unique(lines)) or "Localized Moodle activity pages, course textbook PDF, and final-evaluation resources."


def unit_printed(unit):
    lines = []
    for lesson in unit.get("lessons") or []:
        _media, files = lesson_materials(lesson)
        for item in files:
            lines.append(f"{lesson['id']} - {item}")
    for item in (unit.get("unitResources") or {}).get("evaluations") or []:
        for attachment in item.get("attachments") or []:
            lines.append(f"{item_label(item)} - {item_label(attachment)}")
    return "\n".join(unique(lines)) or "Course outline, textbook, culminating/final summative files, and any Moodle attachments listed in the course."


def template_document(kind):
    if kind == "lesson":
        return Document(str(LESSON_PLAN_TEMPLATE))
    if kind == "unit":
        return Document(str(UNIT_PLAN_TEMPLATE))
    raise ValueError(f"Unknown plan template kind: {kind}")


def style_document(doc):
    section = doc.sections[0]
    section.top_margin = Inches(0.9)
    section.bottom_margin = Inches(0.9)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.12
    for name, size, color in [
        ("Heading 1", 16, "002B57"),
        ("Heading 2", 12.5, "002B57"),
        ("Heading 3", 10.5, "002B57"),
    ]:
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)


def set_text(cell, text, bold=False, size=9.5):
    cell.text = ""
    lines = str(text or "").split("\n")
    for line_index, line in enumerate(lines):
        paragraph = cell.paragraphs[0] if line_index == 0 else cell.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(0)
        run = paragraph.add_run(line)
        run.bold = bold
        run.font.name = "Calibri"
        run.font.size = Pt(size)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP


def set_row_text(row, text, bold=False, size=9.5):
    set_text(row.cells[0], text, bold=bold, size=size)


def lesson_plan_doc(unit, lesson):
    exp = expectation_for(lesson)
    media, files = lesson_materials(lesson)
    detail = lesson_topic_detail(lesson)
    doc = template_document("lesson")
    style_document(doc)

    overview = doc.tables[0]
    rows = overview.rows
    set_text(rows[0].cells[0], "Lesson Plan", bold=True)
    set_text(rows[0].cells[2], "Subject: CHC2D Canadian History since World War I, Grade 10 Academic", bold=True)
    set_row_text(rows[1], f"Lesson Name: {lesson['title']}")
    set_row_text(rows[2], f"Unit of Study: {unit['title']}")
    set_row_text(rows[3], "PRIOR KNOWLEDGE What do my students already know? What key questions or strategies will activate prior knowledge? What connections can I help students make?\n" + prior_knowledge(unit, lesson))
    set_row_text(
        rows[4],
        "CURRICULUM EXPECTATIONS\nOVERALL\n"
        + "\n".join(exp["overall"] or OVERALL_BY_UNIT[int(unit["unit"])])
        + "\n\nSPECIFIC\n"
        + "\n".join(exp["specific"] or ["Use the localized Lesson Expectations page and course outline for exact wording."]),
    )
    set_row_text(rows[5], "\n".join(["LEARNING GOALS What do I want students to know and be able to do?", *lesson_learning_goals(unit, lesson, exp)]))
    set_row_text(rows[6], "\n".join(["SUCCESS CRITERIA(S) How will students know they attained the goal? What tools can they use to self-assess?", *lesson_success_criteria(unit, lesson, exp)]))

    asl, afl, aol = assessment_lists(unit, lesson)
    set_text(rows[7].cells[0], "Assessment as Learning", bold=True)
    set_text(rows[7].cells[1], "Assessment for Learning", bold=True)
    set_text(rows[7].cells[2], "Assessment of Learning", bold=True)
    set_text(rows[8].cells[0], asl)
    set_text(rows[8].cells[1], afl)
    set_text(rows[8].cells[2], aol)
    set_row_text(rows[9], "What will I do?\n" + teacher_assessment_action(unit, lesson))
    set_row_text(rows[10], "Accommodations: How will you change the lesson to meet the needs of individual students?\n" + accommodations_for(lesson))
    set_row_text(
        rows[11],
        "Materials and Resources: What do I need to prepare in advance? Equipment? Student resources? Teacher resources?\n"
        + "\n".join((media + files)[:14] or ["Localized Moodle activity page and course outline resources."])
        + f"\nPlanning source: {GENERATED_SOURCE}",
    )

    flow = doc.tables[1]
    set_text(flow.rows[0].cells[0], "DELIVERING THE LESSON", bold=True)
    phases = [
        (
            "Timing\n5-8\nminutes",
            "Minds On!\nOpen with a timeline prompt, image/source observation, KWL question, or short discussion. Name the historical-thinking focus and the lesson question before students enter the Moodle sequence.",
            "Materials/Resources\nLesson Expectations page\nTimeline/source prompt",
        ),
        (
            "Timing\n20-30\nminutes",
            "Action!\nUse the Lesson page and iSpring presentation where present to model the key historical context. Pause for vocabulary and chronology checks, and require students to connect details to cause, consequence, perspective, or significance.",
            "Materials/Resources\nLocalized lesson page\n" + ("\n".join(media[:4]) if media else "Course activity page"),
        ),
        (
            "Timing\n15-25\nminutes",
            "Action!\nGuide students through the Hands On/H5P/practice task. Require students to support answers with a precise event, date, group, source, or consequence rather than a general opinion.",
            "Materials/Resources\nHands On section\nPractice/H5P resources\nTeacher observation notes",
        ),
        (
            "Timing\n10-15\nminutes",
            "Consolidation!\nUse the consolidation media/reading or homework instructions to check understanding. Close by asking students to answer: "
            + detail["question"]
            + " Confirm submission location and revision target before upload.",
            "Materials/Resources\n" + ("\n".join(files[:6]) if files else "Homework/consolidation instructions"),
        ),
    ]
    for idx, (timing, action, materials) in enumerate(phases, start=1):
        row = flow.rows[idx]
        set_text(row.cells[0], timing)
        set_text(row.cells[1], "W\n[ ]")
        set_text(row.cells[2], "S\n[ ]")
        set_text(row.cells[3], "I\n[ ]")
        set_text(row.cells[4], action + "\n\n" + short_text(lesson_activity_summary(lesson), 480))
        set_text(row.cells[5], materials)

    set_text(
        doc.tables[2].rows[0].cells[0],
        "Notes:\n"
        + GENERATED_NOTE
        + " Keep answer-key resources in the Teacher Packet area; use Homework Submission Folder resources for student-facing submission workflow. Teacher reflection: note which historical-thinking concept or evidence habit needs reteaching next lesson.",
    )
    return doc


def unit_plan_doc(unit):
    unit_no = int(unit["unit"])
    doc = template_document("unit")
    style_document(doc)
    if doc.paragraphs:
        doc.paragraphs[0].text = f"UNIT {unit_no}"
        doc.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    table = doc.tables[0]
    rows = table.rows
    for row_index, label in {
        0: "Unit Author",
        5: "Unit Overview",
        14: "Unit Foundation",
        19: "Assessment Plan",
        22: "Unit Details",
        25: "Materials and Resources",
    }.items():
        set_text(rows[row_index].cells[0], label, bold=True)
    for row_index, label in [(1, "Name:"), (2, "School District:"), (3, "School Name:"), (4, "School City, Province:")]:
        set_row_text(rows[row_index], label)

    codes = unit_codes(unit)
    strand_titles = []
    for code in codes:
        strand = STRANDS.get(code[0])
        if strand:
            strand_titles.append(f"{code[0]}. {strand['title']} - {strand['focus']}")
    target_expectations = [
        "Expectation evidence is drawn from the localized CHC2D Lesson Expectations pages and course outline.",
        f"Codes found in this unit: {', '.join(codes) if codes else 'culminating synthesis across course strands.'}",
        *unique(strand_titles),
        *OVERALL_BY_UNIT.get(unit_no, []),
    ]

    values = [
        (6, "Unit Title Name"),
        (7, unit["title"]),
        (8, "Unit Summary"),
        (9, unit_summary(unit)),
        (10, "Year Level"),
        (11, "Grade 10, Academic Canadian History since World War I"),
        (12, "Approximate Time Needed"),
        (13, f"{max(len(unit.get('lessons') or []) * 3, 4)} instructional hours, plus homework, reflection, and evaluation time"),
        (15, "Targeted Curriculum Expectations"),
        (16, "\n".join(unique(target_expectations))),
        (17, "Learning Goals"),
        (18, "\n".join(unit_learning_goals(unit))),
    ]
    for row_index, value in values:
        set_row_text(rows[row_index], value)

    evaluations = resource_labels((unit.get("unitResources") or {}).get("evaluations") or [])
    set_text(rows[20].cells[0], "Assessment as Learning (ASL)", bold=True)
    set_text(rows[20].cells[2], "Assessment for Learning (AFL)", bold=True)
    set_text(rows[20].cells[3], "Assessment of Learning (AoL)", bold=True)
    set_text(rows[21].cells[0], "[x] KWL/reflection routines\n[x] Exit slips\n[x] Learning log\n[x] Student self-assessment against success criteria")
    set_text(rows[21].cells[2], "[x] Hands On/H5P checks\n[x] Guided worksheet and note checks\n[x] Historical evidence questioning\n[x] Consolidation response review")
    set_text(rows[21].cells[3], "[x] " + "\n[x] ".join(evaluations[:10]) if evaluations else "[x] Culminating/final evaluation where listed")
    set_row_text(rows[23], "Lesson and Assessment Outlines", bold=True)
    set_row_text(rows[24], unit_outline(unit))
    set_text(rows[26].cells[0], "Technology", bold=True)
    set_text(rows[26].cells[1], unit_technology(unit))
    set_text(rows[27].cells[0], "Printed", bold=True)
    set_text(rows[27].cells[1], unit_printed(unit))
    set_text(
        rows[28].cells[0],
        "Other Resources",
        bold=True,
    )
    set_text(
        rows[28].cells[1],
        "Teacher observation notes\nExit cards and consolidation responses\nStudent self-assessment and learning log routines\nCourse outline and Think History textbook\nCurrent historical/context examples selected by the teacher\n" + GENERATED_NOTE,
    )
    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    return doc


def save_doc(doc, relative_path):
    path = COURSE_ROOT / relative_path
    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(path)
    return path.stat().st_size


def file_record(label, relative_path, category, role):
    path = COURSE_ROOT / relative_path
    preview = f"previews-html/{relative_path}.html"
    return {
        "label": label,
        "type": "docx",
        "category": category,
        "role": role,
        "path": relative_path.replace("\\", "/"),
        "bytes": path.stat().st_size,
        "source": GENERATED_SOURCE,
        "sourceStatus": "locally_authored_from_course_materials",
        "teacherUse": "Teacher planning aid reconstructed from localized CHC2D course content; not an original Moodle teacher packet.",
        "templateReference": "MDM4U teacher plan format and local mdm4u-style plan templates.",
        "previewPath": preview.replace("\\", "/"),
    }


def main():
    manifest = json.loads(MANIFEST_PATH.read_text(encoding="utf-8"))
    unit_count = 0
    lesson_count = 0
    written = []

    for unit in manifest.get("units") or []:
        unit_rel = f"plans/generated/unit-plans/U{int(unit['unit']):02d}-unit-plan.docx"
        save_doc(unit_plan_doc(unit), unit_rel)
        unit["unitPlan"] = file_record(f"Unit Plan - {unit['title']}", unit_rel, "unit_plan", "unit_plan")
        unit_count += 1
        written.append(unit_rel)

        for lesson in unit.get("lessons") or []:
            lesson_rel = f"plans/generated/lesson-plans/{lesson['id']}-lesson-plan.docx"
            save_doc(lesson_plan_doc(unit, lesson), lesson_rel)
            lesson["lessonPlan"] = file_record(
                f"Lesson Plan - Unit {unit['unit']} Lesson {lesson['lesson']}",
                lesson_rel,
                "lesson_plan",
                "lesson_plan",
            )
            lesson["resourceCounts"] = {**(lesson.get("resourceCounts") or {}), "lessonPlan": 1}
            lesson_count += 1
            written.append(lesson_rel)

    manifest["generatedAt"] = GENERATED_AT
    manifest["sourceAudit"] = {
        **(manifest.get("sourceAudit") or {}),
        "teacherPlanGeneration": {
            "generatedAt": GENERATED_AT,
            "format": "docx",
            "unitPlans": unit_count,
            "lessonPlans": lesson_count,
            "source": GENERATED_SOURCE,
            "basis": [
                "St.Mary CHC2D Moodle course structure and localized activity records",
                "localized Lesson Expectations pages",
                "localized Lesson, Hands On, Consolidation, and Homework sections",
                "localized iSpring/H5P/video/worksheet/assignment/reflection/evaluation resources",
                "CHC2D course outline and Ontario Canadian and World Studies expectations",
                "local MDM4U-style teacher plan templates",
            ],
            "boundary": "Teacher planning aids only; answer keys remain in Teacher Packet and are not treated as homework submission items.",
        },
    }
    MANIFEST_PATH.write_text(json.dumps(manifest, indent=2, ensure_ascii=False) + "\n", encoding="utf-8")
    REPORT_PATH.parent.mkdir(parents=True, exist_ok=True)
    REPORT_PATH.write_text(
        json.dumps(
            {
                "course": "CHC2D",
                "unitPlans": unit_count,
                "lessonPlans": lesson_count,
                "written": written,
                "generatedAt": GENERATED_AT,
            },
            indent=2,
            ensure_ascii=False,
        )
        + "\n",
        encoding="utf-8",
    )
    print(json.dumps({"course": "CHC2D", "unitPlans": unit_count, "lessonPlans": lesson_count, "generatedAt": GENERATED_AT}, indent=2))


if __name__ == "__main__":
    main()
