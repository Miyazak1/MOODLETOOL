import json
import re
from datetime import datetime, timezone
from html import unescape
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


WORKSPACE_ROOT = Path("D:/工作文件/SUNNYBROOK")
REPO_ROOT = WORKSPACE_ROOT / "ossd-course-portal"
COURSE_ROOT = WORKSPACE_ROOT / "courseware" / "CHV2O"
MANIFEST_PATH = COURSE_ROOT / "course-manifest.json"
REPORT_PATH = REPO_ROOT / "deployment" / "CHV2O-docx-plans-report.json"
GENERATED_AT = datetime.now(timezone.utc).replace(microsecond=0).isoformat()
TEMPLATE_ROOT = REPO_ROOT / "templates" / "teacher-plans"
LESSON_PLAN_TEMPLATE = TEMPLATE_ROOT / "mdm4u-style-lesson-plan-template.docx"
UNIT_PLAN_TEMPLATE = TEMPLATE_ROOT / "mdm4u-style-unit-plan-template.docx"


STRANDS = {
    "A": {
        "title": "Political Inquiry and Skill Development",
        "focus": "inquiry, evidence gathering, interpretation, civic communication, and informed judgement.",
    },
    "B": {
        "title": "Civic Awareness",
        "focus": "democratic values, governance in Canada, rights, responsibilities, and civic issues.",
    },
    "C": {
        "title": "Civic Engagement and Action",
        "focus": "civic contributions, inclusion and participation, media literacy, debate, and action planning.",
    },
}


UNIT_INTENT = {
    1: {
        "big_idea": "Democratic citizenship depends on understanding how Canadian government works, how rights and responsibilities are protected, and how citizens participate in elections and civic debate.",
        "teacher_intent": "Move students from naming civic institutions to explaining how those institutions affect issues that matter to people living in Canada.",
        "essential": [
            "How do government institutions and political parties shape civic decisions in Canada?",
            "How do rights, freedoms, and responsibilities influence democratic participation?",
            "How can voters use evidence, political spectrum language, and polling information to make informed choices?",
        ],
    },
    2: {
        "big_idea": "Citizenship is active: people contribute to the common good by understanding perspectives, using media critically, debating issues, and planning responsible civic action.",
        "teacher_intent": "Teach students to move from opinion to evidence-based civic participation, with special attention to perspective, media messages, and respectful public discourse.",
        "essential": [
            "What makes someone an active citizen?",
            "How do values, beliefs, identity, and media shape public perspectives on civic issues?",
            "How can students design a realistic plan of action for a civic issue?",
        ],
    },
    3: {
        "big_idea": "Civic action requires understanding how governments and outside actors influence democratic life, lawmaking, public revenue, and positive social change.",
        "teacher_intent": "Help students connect Canadian civic processes to real-world influence, including foreign interference, public finance, legislation, and social action.",
        "essential": [
            "How can foreign and domestic actors influence Canadian democracy?",
            "How do governments raise revenue and pass laws to respond to public needs?",
            "How can people living in Canada use their rights and freedoms to create positive social change?",
        ],
    },
    4: {
        "big_idea": "The culminating project asks students to synthesize civic knowledge, inquiry skills, communication, and action planning into final evidence of learning.",
        "teacher_intent": "Guide students through final project planning, evidence selection, revision, and submission while checking readiness for final evaluation.",
        "essential": [
            "How can students demonstrate informed citizenship through a final civic inquiry or action product?",
            "What evidence best shows understanding of civic awareness and civic engagement?",
            "How can students revise and communicate their final learning clearly?",
        ],
    },
}


def clean_text(value):
    text = str(value or "")
    text = re.sub(r"<script[\s\S]*?</script>", " ", text, flags=re.I)
    text = re.sub(r"<style[\s\S]*?</style>", " ", text, flags=re.I)
    text = re.sub(r"<br\s*/?>", "\n", text, flags=re.I)
    text = re.sub(r"</(p|div|li|tr|h[1-6])>", "\n", text, flags=re.I)
    text = re.sub(r"<[^>]+>", " ", text)
    text = unescape(text)
    text = text.replace("\u00a0", " ")
    text = re.sub(r"([A-Z]{1,2})\s+(\d(?:\.\d+)?)", r"\1\2", text)
    return re.sub(r"\s+", " ", text).strip()


def read_local_text(relative_path):
    if not relative_path:
        return ""
    path = COURSE_ROOT / relative_path
    if not path.exists():
        return ""
    return clean_text(path.read_text(encoding="utf-8", errors="ignore"))


def short_text(value, limit=320):
    text = clean_text(value)
    if len(text) <= limit:
        return text
    clipped = text[:limit]
    cut = max(clipped.rfind(". "), clipped.rfind("; "), clipped.rfind(", "))
    if cut > 140:
        clipped = clipped[: cut + 1]
    return clipped.strip() + "..."


def unique(items):
    out = []
    seen = set()
    for item in items:
        value = clean_text(item)
        key = value.lower()
        if value and key not in seen:
            seen.add(key)
            out.append(value)
    return out


def find_section(lesson, pattern):
    regex = re.compile(pattern, re.I)
    for section in lesson.get("bookSections") or []:
        label = section.get("sectionLabel") or section.get("label") or ""
        if regex.search(label):
            return section
    return None


def section_text(lesson, pattern):
    section = find_section(lesson, pattern)
    if not section:
        return ""
    return read_local_text(section.get("path")) or clean_text(section.get("textPreview"))


def strip_repeated_heading(text, lesson):
    title = re.escape(str(lesson.get("title") or ""))
    patterns = [
        rf"^CHV2O Unit \d+ Lesson \d+ - [A-Za-z ]+",
        rf"^Lesson\s+\d+\s*:\s*{title}\s+Lesson Expectations",
        rf"^Lesson\s+\d+\s*:\s*{title}",
    ]
    out = clean_text(text)
    changed = True
    while changed:
        changed = False
        for pattern in patterns:
            new = re.sub(pattern, "", out, flags=re.I).strip()
            if new != out:
                out = clean_text(new)
                changed = True
    return out


def expectation_codes(text):
    codes = re.findall(r"\b[A-C]\s*\d(?:\.\d+)?\b", text or "")
    return unique(code.replace(" ", "") for code in codes)[:18]


def split_expectation_items(block, limit=10):
    block = clean_text(block)
    if not block:
        return []
    pieces = re.split(r"(?=\b[A-C]\s*\d(?:\.\d+)?\b)|(?<=\.)\s+(?=[A-Z][a-z]+)", block)
    return [short_text(piece, 280).rstrip(".") for piece in pieces if len(clean_text(piece)) > 12][:limit]


def extract_between(text, start_labels, end_labels):
    lower = text.lower()
    start = -1
    label_len = 0
    for label in start_labels:
        idx = lower.find(label.lower())
        if idx >= 0 and (start < 0 or idx < start):
            start = idx
            label_len = len(label)
    if start < 0:
        return ""
    start += label_len
    if start < len(text) and text[start] == ":":
        start += 1
    end = len(text)
    for label in end_labels:
        idx = lower.find(label.lower(), start)
        if idx >= 0:
            end = min(end, idx)
    return clean_text(text[start:end])


def expectation_for(lesson):
    raw = strip_repeated_heading(section_text(lesson, r"expectations"), lesson)
    overall = extract_between(
        raw,
        ["Overall Expectations", "Overall Expectation"],
        ["Specific Lesson Expectations", "Specific Expectations", "Learning Goals", "Success Criteria"],
    )
    specific = extract_between(
        raw,
        ["Specific Lesson Expectations", "Specific Expectations", "Specific Lesson Expectation"],
        ["Learning Goals", "Success Criteria"],
    )
    goals = extract_between(raw, ["Learning Goals", "Learning Goal"], ["Success Criteria"])
    success = extract_between(raw, ["Success Criteria", "Success Criterion"], [])
    if not overall and raw:
        overall = raw
    if not specific and raw:
        specific = raw
    return {
        "raw": raw,
        "codes": expectation_codes(raw),
        "overall": split_expectation_items(overall, 6),
        "specific": split_expectation_items(specific, 8),
        "goals": split_expectation_items(goals, 5),
        "success": split_expectation_items(success, 5),
    }


def resource_items(lesson):
    return [item for item in (lesson.get("downloads") or []) + (lesson.get("textExports") or []) if item]


def resource_labels(items):
    return unique(item.get("label") or item.get("title") or Path(str(item.get("path") or "")).name for item in items)


def lesson_materials(lesson):
    media = []
    files = []
    for item in lesson.get("ispring") or []:
        media.append(f"iSpring lesson presentation: {item.get('label') or 'localized presentation'}")
    for item in lesson.get("h5p") or []:
        media.append(f"Hands On/H5P activity: {item.get('label') or 'localized H5P activity'}")
    for item in resource_items(lesson):
        label = item.get("label") or item.get("title") or Path(str(item.get("path") or "")).name
        haystack = " ".join(str(item.get(k, "")) for k in ("type", "category", "role", "path", "label"))
        if re.search(r"mp4|video", haystack, re.I):
            media.append(f"Consolidation/video resource: {label}")
        elif re.search(r"docx?|pdf|pptx?|xlsx?|worksheet|rubric|assignment|kwl|reflection|learning log|exit", haystack, re.I):
            files.append(label)
    return unique(media), unique(files)


def lesson_activity_summary(lesson):
    lesson_page = section_text(lesson, r"^Lesson$")
    hands = section_text(lesson, r"Hands On")
    consolidation = section_text(lesson, r"Consolidation")
    homework = section_text(lesson, r"Homework")
    pieces = []
    if lesson_page:
        pieces.append("Lesson page: " + short_text(lesson_page, 210))
    if hands:
        pieces.append("Hands On: " + short_text(hands, 180))
    if consolidation:
        pieces.append("Consolidation: " + short_text(consolidation, 180))
    if homework:
        pieces.append("Homework: " + short_text(homework, 180))
    return "\n".join(pieces) or "Use the localized Moodle activity page and attached culminating assignment file."


def lesson_goals(unit, lesson, expectation):
    if expectation["goals"]:
        return expectation["goals"]
    title = lesson.get("title") or "the lesson topic"
    codes = ", ".join(expectation["codes"][:4]) or "the listed CHV2O expectations"
    if int(unit["unit"]) == 4:
        return [
            "Plan and complete the culminating civic inquiry/action product using evidence from the course.",
            "Demonstrate understanding of civic awareness, civic engagement, and responsible citizenship.",
            "Communicate final learning clearly and meet the submission requirements for the culminating task.",
        ]
    return [
        f"Explain the civic concept in {title} using accurate course vocabulary.",
        f"Connect the lesson topic to {codes} and to an issue of civic importance.",
        "Use evidence from the lesson page, presentation, Hands On activity, and homework task to support a civic judgement.",
    ]


def success_criteria(unit, lesson, expectation):
    if expectation["success"]:
        return expectation["success"]
    if int(unit["unit"]) == 4:
        return [
            "I can select relevant course evidence for my culminating product.",
            "I can explain my civic position/action clearly and responsibly.",
            "I can revise my work using the assignment requirements before final submission.",
        ]
    title = lesson.get("title") or "the lesson topic"
    return [
        f"I can describe the key idea in {title} in my own words.",
        "I can use at least one example, source, or case from the lesson to support my answer.",
        "I can complete the Hands On/consolidation task and use feedback before submitting homework.",
    ]


def prior_knowledge(unit, lesson):
    if int(unit["unit"]) == 4:
        return (
            "Students bring forward evidence from Units 1-3, including governance, rights and responsibilities, "
            "active citizenship, media, debate, lawmaking, revenue, and positive social change. Begin with a checklist "
            "of course evidence students can reuse responsibly in the culminating task."
        )
    if int(lesson.get("lesson", 1)) == 1:
        return (
            f"Students begin the unit by activating everyday understandings of {unit['title']} through a KWL prompt, "
            "current civic issue, or short discussion. Ask students what they already know, what they can verify, "
            "and what evidence would make their answer stronger."
        )
    prev = next((item for item in unit.get("lessons") or [] if item.get("lesson") == lesson.get("lesson") - 1), None)
    previous = f"Unit {unit['unit']} Lesson {prev['lesson']} ({prev['title']})" if prev else "the previous lesson"
    return (
        f"Students build from {previous}. Start by revisiting one key term or claim from the previous lesson, then "
        f"connect it to {lesson.get('title')} through a short question about civic rights, responsibilities, participation, or evidence."
    )


def assessment_lists(unit, lesson):
    asl = ["[x] Observation", "[x] Anecdotal Notes"]
    if any(re.search(r"exit|reflection|learning log|kwl", str(item.get("label", "")), re.I) for item in resource_items(lesson)):
        asl.extend(["[x] Exit/Reflection evidence", "[x] Self-assessment checklist"])
    else:
        asl.extend(["[ ] Exit Card", "[ ] Self-assessment checklist"])
    afl = ["[x] Strategic Questioning", "[x] Homework", "[x] Worksheet or activity evidence"]
    if lesson.get("h5p") or section_text(lesson, r"Hands On"):
        afl.append("[x] Hands On/H5P practice")
    else:
        afl.append("[ ] Hands On/H5P practice")
    evaluations = resource_labels((unit.get("unitResources") or {}).get("evaluations") or [])
    aol = [f"[x] {label}" for label in evaluations[:6]] or ["[x] Culminating assignment/final evaluation where listed"]
    return "\n".join(asl), "\n".join(afl), "\n".join(aol)


def all_unit_codes(unit):
    codes = []
    for lesson in unit.get("lessons") or []:
        codes.extend(expectation_for(lesson)["codes"])
    return unique(codes)


def unit_learning_goals(unit):
    intent = UNIT_INTENT[int(unit["unit"])]
    goals = [
        f"Students will explain the big idea of {unit['title']} through civic vocabulary and evidence.",
        "Students will interpret civic issues by considering rights, responsibilities, institutions, perspectives, and democratic values.",
        "Students will practise informed civic communication through discussion, debate, reflection, and written responses.",
    ]
    if int(unit["unit"]) == 4:
        goals = [
            "Students will synthesize evidence from the course into a final culminating product.",
            "Students will demonstrate informed citizenship through research, action planning, communication, and reflection.",
            "Students will prepare final evidence of learning that meets the course evaluation requirements.",
        ]
    return goals + intent["essential"]


def unit_summary(unit):
    lessons = unit.get("lessons") or []
    intent = UNIT_INTENT[int(unit["unit"])]
    sequence = ", ".join(lesson.get("title", "") for lesson in lessons)
    return (
        f"{intent['big_idea']}\n\n"
        f"Instructional sequence: {sequence}.\n\n"
        f"Teacher planning focus: {intent['teacher_intent']}\n\n"
        "Students work through localized Moodle lesson pages, iSpring presentations where present, Hands On/H5P or practice activities, "
        "consolidation/reflection evidence, homework folders, and assessment of learning tasks."
    )


def unit_outline(unit):
    chunks = []
    for lesson in unit.get("lessons") or []:
        exp = expectation_for(lesson)
        target = "; ".join(exp["specific"][:2]) or "; ".join(exp["overall"][:2]) or "Use the localized Lesson Expectations page for exact wording."
        media, files = lesson_materials(lesson)
        evidence = []
        if media:
            evidence.append("media/practice: " + "; ".join(media[:3]))
        if files:
            evidence.append("files: " + "; ".join(files[:4]))
        evidence.append("assessment: Hands On/consolidation checks and homework submission")
        chunks.append(
            f"{lesson['id']} - {lesson['title']}\n"
            f"Target: {target}\n"
            f"Teacher plan: introduce the civic question, model evidence-based reasoning, guide the Moodle activity sequence, and close with student reflection or submission evidence.\n"
            + "\n".join(evidence)
        )
    for item in (unit.get("unitResources") or {}).get("evaluations") or []:
        chunks.append(
            f"Assessment of Learning - {item.get('label', '')}\n"
            "Students apply unit learning in an assessed civic task. Feedback should address civic knowledge, inquiry/evidence, communication, and application/action."
        )
    return "\n\n".join(chunks)


def unit_technology(unit):
    lines = []
    for lesson in unit.get("lessons") or []:
        media, _files = lesson_materials(lesson)
        for item in media:
            lines.append(f"{lesson['id']} - {item}")
    return "\n".join(unique(lines)) or "Localized Moodle activity pages and downloadable culminating/course resources."


def unit_printed(unit):
    lines = []
    for lesson in unit.get("lessons") or []:
        _media, files = lesson_materials(lesson)
        for item in files:
            lines.append(f"{lesson['id']} - {item}")
    for item in (unit.get("unitResources") or {}).get("evaluations") or []:
        for attachment in item.get("attachments") or []:
            lines.append(f"{item.get('label', '')} - {attachment.get('label', '')}")
    return "\n".join(unique(lines)) or "Culminating assignment file and any Moodle attachments listed in the course."


def set_cell_width(cell, width_dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_text(cell, text, bold=False, size=9.5):
    cell.text = ""
    lines = str(text or "").split("\n")
    for line_index, line in enumerate(lines):
        paragraph = cell.paragraphs[0] if line_index == 0 else cell.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(0)
        run = paragraph.add_run(line)
        run.bold = bold
        run.font.name = "Calibri"
        run.font.size = Pt(size)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP


def set_row_text(row, text, bold=False, size=9.5):
    set_text(row.cells[0], text, bold=bold, size=size)


def template_document(kind):
    if kind == "lesson":
        return Document(str(LESSON_PLAN_TEMPLATE))
    if kind == "unit":
        return Document(str(UNIT_PLAN_TEMPLATE))
    raise ValueError(f"Unknown plan template kind: {kind}")


def style_document(doc):
    section = doc.sections[0]
    section.top_margin = Inches(0.9)
    section.bottom_margin = Inches(0.9)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.12
    for name, size, color in [("Heading 1", 16, "002B57"), ("Heading 2", 12.5, "002B57"), ("Heading 3", 10.5, "002B57")]:
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)


def lesson_plan_doc(unit, lesson):
    expectation = expectation_for(lesson)
    media, files = lesson_materials(lesson)
    doc = template_document("lesson")
    style_document(doc)

    overview = doc.tables[0]
    rows = overview.rows
    set_text(rows[0].cells[0], "Lesson Plan", bold=True)
    set_text(rows[0].cells[2], "Subject: CHV2O Civics and Citizenship, Grade 10 Open", bold=True)
    set_row_text(rows[1], f"Lesson Name: {lesson['title']}")
    set_row_text(rows[2], f"Unit of Study: {unit['title']}")
    set_row_text(rows[3], "PRIOR KNOWLEDGE What do my students already know? What key questions or strategies will activate prior knowledge? What connections can I help students make?\n" + prior_knowledge(unit, lesson))

    expectation_text = "\n".join(
        [
            "CURRICULUM EXPECTATIONS",
            "OVERALL",
            *(expectation["overall"] or ["Use the localized Lesson Expectations page for exact wording."]),
            "SPECIFIC",
            *(expectation["specific"] or ["Use the localized Lesson Expectations page for exact wording."]),
        ]
    )
    set_row_text(rows[4], expectation_text)
    set_row_text(rows[5], "\n".join(["LEARNING GOALS What do I want students to know and be able to do?", *lesson_goals(unit, lesson, expectation)]))
    set_row_text(rows[6], "\n".join(["SUCCESS CRITERIA(S) How will students know they attained the goal? What tools can they use to self-assess?", *success_criteria(unit, lesson, expectation)]))

    asl, afl, aol = assessment_lists(unit, lesson)
    set_text(rows[7].cells[0], "Assessment as Learning", bold=True)
    set_text(rows[7].cells[1], "Assessment for Learning", bold=True)
    set_text(rows[7].cells[2], "Assessment of Learning", bold=True)
    set_text(rows[8].cells[0], asl)
    set_text(rows[8].cells[1], afl)
    set_text(rows[8].cells[2], aol)
    set_row_text(rows[9], "What will I do?\nConfer\nObserve\nGive descriptive feedback\nGrade where the task is identified as assessment of learning")
    set_row_text(
        rows[10],
        "Accommodations: How will you change the lesson to meet the needs of individual students?\n"
        "Chunk civic texts and media into short checkpoints; pre-teach vocabulary; provide sentence frames for claims, evidence, and reasoning; offer oral rehearsal before written work; allow graphic organizers for rights/responsibilities or stakeholder perspectives; and extend with current Canadian civic examples.",
    )
    set_row_text(
        rows[11],
        "Materials and Resources: What do I need to prepare in advance? Equipment? Student resources? Teacher resources?\n"
        + "\n".join((media + files)[:12] or ["Localized Moodle activity page, culminating assignment file, and course outline resources."])
        + f"\nSource status: locally authored from CHV2O Moodle content on {GENERATED_AT}.",
    )

    lesson_table = doc.tables[1]
    set_text(lesson_table.rows[0].cells[0], "DELIVERING THE LESSON", bold=True)
    lesson_rows = lesson_table.rows
    phases = [
        (
            1,
            "Timing\n5-8\nminutes",
            "Minds On!\nOpen with a civic question, current issue, KWL prompt, or rights/responsibilities quick-write. Name the expectation codes and learning goal before students enter the Moodle activity sequence.",
            "Materials/Resources\nLesson Expectations page\nKWL or discussion prompt",
        ),
        (
            2,
            "Timing\n20-30\nminutes",
            "Action!\nUse the Lesson page and iSpring presentation where present to model the civic concept. Pause for vocabulary checks and ask students to connect the concept to a Canadian example or issue of civic importance.",
            "Materials/Resources\nLocalized lesson page\n" + ("\n".join(media[:4]) if media else "Course activity page"),
        ),
        (
            3,
            "Timing\n15-25\nminutes",
            "Action!\nGuide students through the Hands On/H5P/practice task. Require a claim, evidence, and reasoning response, and confer with students whose evidence is too general.",
            "Materials/Resources\nHands On section\nPractice/H5P resources\nTeacher observation notes",
        ),
        (
            4,
            "Timing\n10-15\nminutes",
            "Consolidation!\nUse the consolidation or homework instructions to check understanding. Confirm submission location and have students identify one revision target before uploading work.",
            "Materials/Resources\n" + ("\n".join(files[:6]) if files else "Homework/consolidation instructions"),
        ),
    ]
    for row_index, timing, action, materials in phases:
        row = lesson_rows[row_index]
        set_text(row.cells[0], timing)
        set_text(row.cells[1], "W\n[ ]")
        set_text(row.cells[2], "S\n[ ]")
        set_text(row.cells[3], "I\n[ ]")
        set_text(row.cells[4], action + "\n\n" + short_text(lesson_activity_summary(lesson), 420))
        set_text(row.cells[5], materials)

    notes = doc.tables[2]
    set_text(
        notes.rows[0].cells[0],
        "Notes:\nGenerated teacher planning aid reconstructed from localized CHV2O Moodle content. Keep answer-key resources in the teacher packet area; use Homework Submission Folder resources for student-facing submission workflow. Teacher reflection: note which civic term, institution, or evidence habit needs reteaching next lesson.",
    )
    return doc


def unit_plan_doc(unit):
    doc = template_document("unit")
    style_document(doc)
    if doc.paragraphs:
        doc.paragraphs[0].text = f"UNIT {unit['unit']}"
        doc.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    table = doc.tables[0]
    rows = table.rows
    for row_index, label in {0: "Unit Author", 5: "Unit Overview", 14: "Unit Foundation", 19: "Assessment Plan", 22: "Unit Details", 25: "Materials and Resources"}.items():
        set_text(rows[row_index].cells[0], label, bold=True)
    for row_index, label in [(1, "Name:"), (2, "School District:"), (3, "School Name:"), (4, "School City, Province:")]:
        set_row_text(rows[row_index], label)

    codes = all_unit_codes(unit)
    strand_titles = []
    for code in codes:
        strand = STRANDS.get(code[0])
        if strand:
            strand_titles.append(f"{code[0]}. {strand['title']} - {strand['focus']}")
    values = [
        (6, "Unit Title Name"),
        (7, unit["title"]),
        (8, "Unit Summary"),
        (9, unit_summary(unit)),
        (10, "Year Level"),
        (11, "Grade 10, Open Civics and Citizenship"),
        (12, "Approximate Time Needed"),
        (13, f"{max(len(unit.get('lessons') or []) * 3, 4)} instructional hours, plus homework and evaluation time"),
        (15, "Targeted Curriculum Expectations"),
        (
            16,
            "\n".join(
                [
                    "Expectation evidence is drawn from the localized CHV2O Lesson Expectations pages.",
                    f"Codes found in this unit: {', '.join(codes) if codes else 'see culminating assignment requirements.'}",
                    *unique(strand_titles),
                ]
            ),
        ),
        (17, "Learning Goals"),
        (18, "\n".join(unit_learning_goals(unit))),
    ]
    for row_index, value in values:
        set_row_text(rows[row_index], value)

    set_text(rows[20].cells[0], "Assessment as Learning (ASL)", bold=True)
    set_text(rows[20].cells[2], "Assessment for Learning (AFL)", bold=True)
    set_text(rows[20].cells[3], "Assessment of Learning (AoL)", bold=True)
    set_text(rows[21].cells[0], "[x] KWL/reflection routines\n[x] Exit cards\n[x] Learning log\n[x] Student next-step notes")
    set_text(rows[21].cells[2], "[x] Hands On/H5P checks\n[x] Strategic questioning\n[x] Homework review\n[x] Consolidation responses")
    evaluations = resource_labels((unit.get("unitResources") or {}).get("evaluations") or [])
    set_text(rows[21].cells[3], "[x] " + "\n[x] ".join(evaluations[:10]) if evaluations else "[x] Culminating assignment/final evaluation")
    set_row_text(rows[23], "Lesson and Assessment Outlines", bold=True)
    set_row_text(rows[24], unit_outline(unit))
    set_text(rows[26].cells[0], "Technology", bold=True)
    set_text(rows[26].cells[1], unit_technology(unit))
    set_text(rows[27].cells[0], "Printed", bold=True)
    set_text(rows[27].cells[1], unit_printed(unit))
    set_text(rows[28].cells[0], "Other Resources", bold=True)
    set_text(rows[28].cells[1], "Anecdotal notes of observation\nExit cards and consolidation responses\nStudent self-assessment and reflection\nKWL chart and Learning Log routines where present\nCurrent Canadian civic issue examples selected by the teacher")
    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    return doc


def save_doc(doc, relative_path):
    path = COURSE_ROOT / relative_path
    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(path)
    return path.stat().st_size


def file_record(label, relative_path, category, role):
    path = COURSE_ROOT / relative_path
    preview = f"previews-html/{relative_path}.html"
    return {
        "label": label,
        "type": "docx",
        "category": category,
        "role": role,
        "path": relative_path.replace("\\", "/"),
        "bytes": path.stat().st_size,
        "source": "reconstructed_from_moodle_content",
        "sourceStatus": "locally_authored_from_course_materials",
        "teacherUse": "Teacher planning aid reconstructed from localized CHV2O course content; not an original Moodle teacher packet.",
        "templateReference": "MDM4U teacher plan format and local mdm4u-style plan templates.",
        "previewPath": preview.replace("\\", "/"),
    }


def main():
    manifest = json.loads(MANIFEST_PATH.read_text(encoding="utf-8"))
    unit_count = 0
    lesson_count = 0
    written = []

    for unit in manifest.get("units") or []:
        unit_rel = f"plans/generated/unit-plans/U{unit['unit']:02d}-unit-plan.docx"
        save_doc(unit_plan_doc(unit), unit_rel)
        unit["unitPlan"] = file_record(f"Unit Plan - {unit['title']}", unit_rel, "unit_plan", "unit_plan")
        unit_count += 1
        written.append(unit_rel)

        for lesson in unit.get("lessons") or []:
            lesson_rel = f"plans/generated/lesson-plans/{lesson['id']}-lesson-plan.docx"
            save_doc(lesson_plan_doc(unit, lesson), lesson_rel)
            lesson["lessonPlan"] = file_record(
                f"Lesson Plan - Unit {unit['unit']} Lesson {lesson['lesson']}",
                lesson_rel,
                "lesson_plan",
                "lesson_plan",
            )
            lesson["resourceCounts"] = {**(lesson.get("resourceCounts") or {}), "lessonPlan": 1}
            lesson_count += 1
            written.append(lesson_rel)

    manifest["generatedAt"] = GENERATED_AT
    manifest["sourceAudit"] = {
        **(manifest.get("sourceAudit") or {}),
        "teacherPlanGeneration": {
            "generatedAt": GENERATED_AT,
            "format": "docx",
            "unitPlans": unit_count,
            "lessonPlans": lesson_count,
            "source": "reconstructed_from_moodle_content",
            "basis": [
                "localized Moodle Lesson Expectations pages",
                "localized Moodle Lesson, Hands On, Consolidation, and Homework sections",
                "localized iSpring, H5P, worksheet, assignment, reflection, exit card, and Evaluation records",
                "CHV2O course outline and local MDM4U-style teacher plan templates",
            ],
            "boundary": "Teacher planning aids only; no answer keys, rubrics, textbook excerpts, or Moodle-original teacher packet documents were created.",
        },
    }
    MANIFEST_PATH.write_text(json.dumps(manifest, indent=2, ensure_ascii=False) + "\n", encoding="utf-8")
    REPORT_PATH.parent.mkdir(parents=True, exist_ok=True)
    REPORT_PATH.write_text(
        json.dumps({"course": "CHV2O", "unitPlans": unit_count, "lessonPlans": lesson_count, "written": written, "generatedAt": GENERATED_AT}, indent=2, ensure_ascii=False) + "\n",
        encoding="utf-8",
    )
    print(json.dumps({"course": "CHV2O", "unitPlans": unit_count, "lessonPlans": lesson_count, "generatedAt": GENERATED_AT}, indent=2))


if __name__ == "__main__":
    main()
