import json
import re
from html import escape, unescape
from html.parser import HTMLParser
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


COURSE = "ENG1D"
REPO_ROOT = Path(__file__).resolve().parents[1]
WORKSPACE_ROOT = REPO_ROOT.parent
COURSE_ROOT = WORKSPACE_ROOT / "courseware" / COURSE
MANIFEST_PATH = COURSE_ROOT / "course-manifest.json"
REPORT_PATH = REPO_ROOT / "deployment" / "ENG1D-text-docx-report.json"


def clean_text(value):
    text = unescape(str(value or ""))
    text = re.sub(r"\s+", " ", text)
    return text.strip()


class HtmlBlocks(HTMLParser):
    block_tags = {"p", "li", "blockquote", "pre", "div"}
    heading_tags = {"h1", "h2", "h3", "h4", "h5", "h6", "title"}
    skip_tags = {"script", "style", "noscript", "svg"}

    def __init__(self):
        super().__init__(convert_charrefs=True)
        self.blocks = []
        self.stack = []
        self.buffer = []
        self.current_tag = None
        self.skip_depth = 0

    def flush(self):
        text = clean_text("".join(self.buffer))
        if text:
            tag = self.current_tag or "p"
            if not self.blocks or self.blocks[-1] != (tag, text):
                self.blocks.append((tag, text))
        self.buffer = []
        self.current_tag = None

    def handle_starttag(self, tag, attrs):
        tag = tag.lower()
        if tag in self.skip_tags:
            self.skip_depth += 1
            return
        if self.skip_depth:
            return
        if tag in self.block_tags or tag in self.heading_tags:
            self.flush()
            self.current_tag = tag
        elif tag == "br":
            self.buffer.append("\n")

    def handle_endtag(self, tag):
        tag = tag.lower()
        if tag in self.skip_tags and self.skip_depth:
            self.skip_depth -= 1
            return
        if self.skip_depth:
            return
        if tag in self.block_tags or tag in self.heading_tags:
            self.flush()

    def handle_data(self, data):
        if not self.skip_depth:
            self.buffer.append(data)

    def close(self):
        super().close()
        self.flush()


def html_blocks(path):
    html = path.read_text(encoding="utf-8", errors="ignore")
    parser = HtmlBlocks()
    parser.feed(html)
    parser.close()
    blocks = []
    seen_gutenberg_license = False
    for tag, text in parser.blocks:
        if len(text) > 18000:
            text = text[:18000] + "..."
        if "Project Gutenberg License" in text:
            if seen_gutenberg_license:
                continue
            seen_gutenberg_license = True
        if re.fullmatch(r"[\W_]{1,12}", text):
            continue
        blocks.append((tag, text))
    return blocks


def readable_html(text_entry, material, blocks):
    title = text_entry.get("title") or material.get("label") or "ENG1D Text"
    author = text_entry.get("author") or ""
    notes = text_entry.get("notes") or material.get("textPreview") or ""
    source = material.get("source") or text_entry.get("source") or ""
    status = "; ".join(part for part in [text_entry.get("copyrightStatus"), text_entry.get("sourceStatus")] if part)
    body = []
    skipped_title = False
    for tag, text in blocks:
        if not skipped_title and title.lower() in text.lower() and len(text) < 180:
            skipped_title = True
            continue
        safe = escape(text)
        if tag in {"h1", "title"}:
            body.append(f"<h1>{safe}</h1>")
        elif tag == "h2":
            body.append(f"<h2>{safe}</h2>")
        elif tag in {"h3", "h4", "h5", "h6"}:
            body.append(f"<h3>{safe}</h3>")
        elif tag == "li":
            body.append(f"<li>{safe}</li>")
        elif tag == "blockquote":
            body.append(f"<blockquote>{safe}</blockquote>")
        else:
            body.append(f"<p>{safe}</p>")
    body_html = "\n".join(body)
    return f"""<!doctype html>
<html lang="en" data-eng1d-readable-text="true">
<head>
  <meta charset="utf-8">
  <meta name="viewport" content="width=device-width, initial-scale=1">
  <title>{escape(title)}</title>
  <style>
    :root {{ color: #001f3f; background: #f3f6fa; font-family: Inter, "Segoe UI", Arial, Helvetica, sans-serif; line-height: 1.65; }}
    body {{ margin: 0; padding: 32px 18px 56px; background: #f3f6fa; }}
    main {{ max-width: 1160px; margin: 0 auto; background: #fff; border: 1px solid #d6e2f0; border-radius: 8px; padding: 30px 36px 44px; box-sizing: border-box; }}
    header {{ border-bottom: 1px solid #e0e8f2; margin-bottom: 28px; padding-bottom: 18px; }}
    .eyebrow {{ color: #55708f; font-size: 14px; font-weight: 700; letter-spacing: 0; margin: 0 0 8px; text-transform: uppercase; }}
    h1 {{ color: #001f3f; font-size: 32px; line-height: 1.22; margin: 0 0 10px; }}
    h2 {{ color: #001f3f; font-size: 24px; line-height: 1.3; margin: 34px 0 14px; }}
    h3 {{ color: #001f3f; font-size: 20px; line-height: 1.35; margin: 28px 0 12px; }}
    .meta {{ color: #3b5573; font-size: 15px; margin: 4px 0; }}
    .notice {{ background: #f6f9fd; border: 1px solid #d6e2f0; border-radius: 8px; color: #24425f; margin: 18px 0 0; padding: 12px 14px; }}
    .reading-body {{ max-width: 1080px; }}
    .reading-body p, .reading-body li, .reading-body blockquote {{ font-size: 18px; line-height: 1.72; margin: 0 0 18px; }}
    .reading-body li {{ margin-left: 24px; padding-left: 4px; }}
    blockquote {{ border-left: 4px solid #c7d8ea; color: #223f5c; margin-left: 0; padding-left: 18px; }}
    a {{ color: #064f9e; }}
    @media (max-width: 720px) {{
      body {{ padding: 0; }}
      main {{ border-left: 0; border-right: 0; border-radius: 0; padding: 24px 20px 38px; }}
      h1 {{ font-size: 27px; }}
      .reading-body p, .reading-body li, .reading-body blockquote {{ font-size: 17px; }}
    }}
  </style>
</head>
<body>
  <main>
    <header>
      <p class="eyebrow">ENG1D Text Reference</p>
      <h1>{escape(title)}</h1>
      {f'<p class="meta">{escape(author)}</p>' if author else ''}
      {f'<p class="meta">Source: {escape(source)}</p>' if source else ''}
      {f'<p class="meta">Status: {escape(status)}</p>' if status else ''}
      {f'<p class="notice">{escape(notes)}</p>' if notes else ''}
    </header>
    <article class="reading-body">
      {body_html}
    </article>
  </main>
</body>
</html>
"""


def write_readable_html(text_entry, material, source_path):
    blocks = html_blocks(source_path)
    source_path.write_text(readable_html(text_entry, material, blocks), encoding="utf-8")
    return source_path.stat().st_size


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def style_document(doc):
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)
    styles = doc.styles
    styles["Normal"].font.name = "Calibri"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    styles["Normal"].font.size = Pt(10.5)
    for name, size, color in [
        ("Heading 1", 20, "001F3F"),
        ("Heading 2", 15, "001F3F"),
        ("Heading 3", 12.5, "001F3F"),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)


def add_metadata_table(doc, text_entry, material):
    table = doc.add_table(rows=4, cols=2)
    table.style = "Table Grid"
    rows = [
        ("Course", "ENG1D - English, Grade 9, Academic"),
        ("Text", text_entry.get("title") or material.get("label") or "Text"),
        ("Author / Source", f"{text_entry.get('author') or ''} / {material.get('source') or text_entry.get('source') or ''}".strip(" /")),
        ("Status", f"{text_entry.get('copyrightStatus') or ''}; {text_entry.get('sourceStatus') or ''}".strip("; ")),
    ]
    for idx, (label, value) in enumerate(rows):
        table.rows[idx].cells[0].text = label
        table.rows[idx].cells[1].text = value
        set_cell_shading(table.rows[idx].cells[0], "EAF2FB")
        for paragraph in table.rows[idx].cells[0].paragraphs:
            for run in paragraph.runs:
                run.bold = True


def build_docx(text_entry, material, source_path, target_path):
    blocks = html_blocks(source_path)
    doc = Document()
    style_document(doc)

    title = text_entry.get("title") or material.get("label") or source_path.stem
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(title)
    r.bold = True
    r.font.size = Pt(22)
    r.font.color.rgb = RGBColor.from_string("001F3F")

    author = text_entry.get("author")
    if author:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(author)
        r.italic = True
        r.font.size = Pt(11)

    if text_entry.get("notes"):
        p = doc.add_paragraph(text_entry["notes"])
        p.paragraph_format.space_after = Pt(8)

    add_metadata_table(doc, text_entry, material)
    doc.add_paragraph()

    skipped_title = False
    for tag, text in blocks:
        if not skipped_title and title.lower() in text.lower() and len(text) < 180:
            skipped_title = True
            continue
        if tag in {"h1", "title"}:
            doc.add_paragraph(text, style="Heading 1")
        elif tag == "h2":
            doc.add_paragraph(text, style="Heading 2")
        elif tag in {"h3", "h4", "h5", "h6"}:
            doc.add_paragraph(text, style="Heading 3")
        elif tag == "li":
            doc.add_paragraph(text, style="List Bullet")
        elif tag == "blockquote":
            p = doc.add_paragraph(text)
            p.paragraph_format.left_indent = Inches(0.35)
            p.paragraph_format.right_indent = Inches(0.2)
            for run in p.runs:
                run.italic = True
        else:
            p = doc.add_paragraph(text)
            p.paragraph_format.space_after = Pt(6)

    target_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(target_path)
    return target_path.stat().st_size


def docx_rel_for(html_rel):
    source = Path(html_rel)
    without_suffix = source.with_suffix("")
    return str(Path("texts") / "docx" / without_suffix.relative_to("texts")).replace("\\", "/") + ".docx"


def material_key(item):
    return item.get("path") or item.get("downloadPath") or item.get("label")


def upsert_material(materials, item):
    key = material_key(item)
    for index, existing in enumerate(materials):
        if material_key(existing) == key:
            materials[index] = {**existing, **item}
            return
    materials.append(item)


def upsert_download(downloads, item):
    key = item.get("path")
    for index, existing in enumerate(downloads):
        if existing.get("path") == key:
            downloads[index] = {**existing, **item}
            return
    downloads.append(item)


def update_download_bytes(downloads, rel_path, bytes_written):
    for item in downloads:
        if item.get("path") == rel_path:
            item["bytes"] = bytes_written
            item["previewPath"] = rel_path
            item["downloadPath"] = rel_path


def main():
    manifest = json.loads(MANIFEST_PATH.read_text(encoding="utf-8"))
    report = []
    generated_by_source = {}
    normalized_html_by_source = {}

    for text_entry in manifest.get("texts") or []:
        materials = text_entry.get("materials") or []
        for material in list(materials):
            html_rel = material.get("path")
            if material.get("type") != "html" or not html_rel:
                continue
            source_path = COURSE_ROOT / html_rel
            if not source_path.exists():
                report.append({"textId": text_entry.get("id"), "source": html_rel, "status": "missing-html"})
                continue
            docx_rel = docx_rel_for(html_rel)
            target_path = COURSE_ROOT / docx_rel
            if html_rel in generated_by_source:
                bytes_written = generated_by_source[html_rel]
            else:
                bytes_written = build_docx(text_entry, material, source_path, target_path)
                generated_by_source[html_rel] = bytes_written
            if html_rel in normalized_html_by_source:
                html_bytes = normalized_html_by_source[html_rel]
            else:
                html_bytes = write_readable_html(text_entry, material, source_path)
                normalized_html_by_source[html_rel] = html_bytes
            material["bytes"] = html_bytes
            material["previewPath"] = html_rel
            update_download_bytes(manifest.setdefault("courseDownloads", []), html_rel, html_bytes)
            docx_material = {
                "label": re.sub(r"\s+-\s+public-domain", "", material.get("label") or text_entry.get("title") or "Text") + " - DOCX",
                "title": (material.get("title") or material.get("label") or text_entry.get("title") or "Text") + " - DOCX",
                "type": "docx",
                "category": material.get("category") or text_entry.get("category") or "text",
                "role": material.get("role") or text_entry.get("role") or "text_docx",
                "path": docx_rel,
                "previewPath": f"previews-html/{docx_rel}.html",
                "downloadPath": docx_rel,
                "bytes": bytes_written,
                "source": material.get("source") or text_entry.get("source"),
                "textPreview": f"DOCX version of {text_entry.get('title') or material.get('label')}.",
                "derivedFrom": html_rel,
            }
            upsert_material(materials, docx_material)
            upsert_download(manifest.setdefault("courseDownloads", []), {**docx_material, "textId": text_entry.get("id")})
            report.append({"textId": text_entry.get("id"), "source": html_rel, "docx": docx_rel, "bytes": bytes_written, "status": "generated"})
        text_entry["materials"] = materials

    manifest.setdefault("sourceAudit", {})
    manifest["sourceAudit"]["textDocxGeneration"] = {
        "generatedAt": __import__("datetime").datetime.now(__import__("datetime").timezone.utc).isoformat(),
        "generated": [item for item in report if item.get("status") == "generated"],
        "missing": [item for item in report if item.get("status") != "generated"],
        "note": "DOCX files are generated from the localized HTML text materials so teachers can download, print, and annotate them.",
    }
    manifest["generatedAt"] = manifest["sourceAudit"]["textDocxGeneration"]["generatedAt"]
    MANIFEST_PATH.write_text(json.dumps(manifest, indent=2) + "\n", encoding="utf-8")
    REPORT_PATH.write_text(json.dumps({"course": COURSE, "report": report}, indent=2) + "\n", encoding="utf-8")
    print(json.dumps({"course": COURSE, "generated": len([item for item in report if item.get("status") == "generated"]), "report": str(REPORT_PATH)}, indent=2))


if __name__ == "__main__":
    main()
