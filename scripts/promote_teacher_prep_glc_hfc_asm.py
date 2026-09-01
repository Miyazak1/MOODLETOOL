import json
import re
import shutil
from datetime import datetime, timezone
from html import unescape
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


WORKSPACE_ROOT = Path("D:/工作文件/SUNNYBROOK")
COURSEWARE_ROOT = WORKSPACE_ROOT / "courseware"
REPO_ROOT = WORKSPACE_ROOT / "ossd-course-portal"
REPORT_PATH = REPO_ROOT / "deployment" / "teacher-prep-promotion-GLC2O-HFC3M-ASM3M.json"
GENERATED_AT = datetime.now(timezone.utc).replace(microsecond=0).isoformat()


COURSE_PROFILES = {
    "GLC2O": {
        "subject": "Career Studies",
        "grade": "Grade 10, Open",
        "curriculum_title": "The Ontario Curriculum, Grade 10: Guidance and Career Education - Career Studies (GLC2O), 2024",
        "curriculum_url": "https://www.dcp.edu.gov.on.ca/en/curriculum/secondary-guidance-and-career-education/courses/glc2o/home",
        "source_url": "http://34.30.231.58/course/view.php?id=78",
        "course_lens": "career-life planning, transferable skills, postsecondary pathways, job search, financial literacy, resilience, stress management, and goal-setting.",
        "unit_focus": {
            1: "rights and responsibilities, transferable skills, personal values, strengths, interests, job search, social media, and postsecondary pathways.",
            2: "postsecondary planning, budgeting, consumer credit, transferable skills, and values-based decision making.",
            3: "resilience, stress management, balanced wellbeing, decision-making, and goal-setting strategies.",
            4: "final culminating project planning, evidence collection, reflection, and presentation of career-life planning growth.",
        },
        "essential_questions": [
            "How do students identify strengths, values, skills, and interests that can guide pathway decisions?",
            "How can career-life planning be grounded in evidence, reflection, and realistic next steps?",
            "How do financial decisions, wellbeing strategies, and digital presence affect postsecondary and work options?",
        ],
        "source_note": "This teacher-prep index uses the localized Moodle course, the official Ontario Digital Curriculum Platform course page, and locally generated planning documents. No commercial textbook copy was added.",
    },
    "HFC3M": {
        "subject": "Food and Culture",
        "grade": "Grade 11, University/College Preparation",
        "curriculum_title": "The Ontario Curriculum: Social Sciences and Humanities, Grades 9 to 12, 2013 (Revised)",
        "curriculum_url": "https://www.edu.gov.on.ca/eng/curriculum/secondary/ssciences9to122013.pdf",
        "curriculum_copy_from": COURSEWARE_ROOT / "HFA4U" / "texts" / "ontario-curriculum" / "ssciences9to122013.pdf",
        "curriculum_dest": Path("texts/ontario-curriculum/ssciences9to122013.pdf"),
        "source_url": "https://www.esunnybrook.com/course/view.php?id=56",
        "course_lens": "food safety, food preparation, cultural food practices, global food systems, Canadian food identity, nutrition, and practical kitchen decision-making.",
        "unit_focus": {
            1: "kitchen safety, sanitation, food preparation skills, measuring, recipe reading, and responsible lab habits.",
            2: "culture, identity, food traditions, celebrations, migration, and respectful inquiry into food practices.",
            3: "global food systems, regional cuisines, sustainability, food security, and comparative cultural analysis.",
            4: "food in Canada, Indigenous and newcomer foodways, local food systems, and Canadian food identity.",
            5: "culminating performance task planning, product/process evidence, and reflective presentation.",
            6: "final review, synthesis, assessment preparation, and transfer of course concepts.",
            7: "teacher observations of learning skills, collaboration, responsibility, initiative, and independent work habits.",
        },
        "essential_questions": [
            "How do food choices reflect culture, identity, availability, and social context?",
            "How can students prepare and evaluate food safely, respectfully, and sustainably?",
            "How can food studies connect practical skills with inquiry into families, communities, and global systems?",
        ],
        "source_note": "This teacher-prep index uses localized Moodle activities, the existing HFC3M course outline, and the official Social Sciences and Humanities curriculum PDF already present in local courseware.",
    },
    "ASM3M": {
        "subject": "Media Arts",
        "grade": "Grade 11, University/College Preparation",
        "curriculum_title": "The Ontario Curriculum, Grades 11 and 12: The Arts, 2010 (Revised)",
        "curriculum_url": "https://www.edu.gov.on.ca/eng/curriculum/secondary/arts1112curr2010.pdf",
        "curriculum_copy_from": COURSEWARE_ROOT / "AVI4M" / "texts" / "ontario-arts-curriculum-11-12" / "arts1112curr2010.pdf",
        "curriculum_dest": Path("texts/ontario-arts-curriculum-11-12/arts1112curr2010.pdf"),
        "source_url": "Moodle course id 66",
        "course_lens": "media arts design, elements and principles, character design, video art, animation, performance/installation, video game design, critique, portfolio evidence, and creative-process documentation.",
        "unit_focus": {
            1: "elements and principles of design, visual communication, conceptual design, logo/self-portrait options, critique, and learning-log reflection.",
            2: "character design, visual mapping, facial structure, project options, teacher observation, and self-reflection.",
            3: "continuity, storyboarding, narrative, contemporary art, video art, installation, performance, and assignment options.",
            4: "video game design, interactive media planning, character design, worksheet evidence, and project submission.",
            5: "independent study/final evaluation, portfolio evidence, exam preparation, and final submission expectations.",
        },
        "essential_questions": [
            "How do media artists use elements, principles, tools, and technologies to communicate ideas?",
            "How can students document creative process, critique, and revision as assessment evidence?",
            "How do media forms such as character design, video art, and games shape audience experience?",
        ],
        "source_note": "This teacher-prep index uses the localized Moodle shell, downloaded attachments, locally generated planning documents, and the official Ontario Arts curriculum. No separate commercial textbook was exposed in Moodle.",
    },
}


def clean_text(value):
    text = str(value or "")
    text = re.sub(r"<script[\s\S]*?</script>", " ", text, flags=re.I)
    text = re.sub(r"<style[\s\S]*?</style>", " ", text, flags=re.I)
    text = re.sub(r"<[^>]+>", " ", text)
    text = unescape(text)
    return re.sub(r"\s+", " ", text).strip()


def read_local_text(course_root, rel_path):
    if not rel_path:
        return ""
    path = course_root / rel_path
    if not path.exists() or path.suffix.lower() not in {".html", ".htm", ".md", ".txt"}:
        return ""
    return clean_text(path.read_text(encoding="utf-8", errors="ignore"))


def short_text(value, limit=280):
    text = clean_text(value)
    if len(text) <= limit:
        return text
    clipped = text[:limit]
    cut = max(clipped.rfind(". "), clipped.rfind("; "), clipped.rfind(", "))
    if cut > 120:
        clipped = clipped[: cut + 1]
    return clipped.strip() + "..."


def slug(value, limit=80):
    text = re.sub(r"[^A-Za-z0-9]+", "-", str(value or "")).strip("-")
    return (text[:limit].strip("-") or "item")


def rel(path, course_root):
    return path.relative_to(course_root).as_posix()


def ensure_dir(path):
    path.mkdir(parents=True, exist_ok=True)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False):
    cell.text = ""
    para = cell.paragraphs[0]
    run = para.add_run(str(text or ""))
    run.bold = bold
    run.font.name = "Aptos"
    run.font.size = Pt(9.5)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP


def add_table(document, rows):
    table = document.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    set_cell_text(hdr[0], "Planning Area", True)
    set_cell_text(hdr[1], "Teacher Notes", True)
    set_cell_shading(hdr[0], "D9EAF7")
    set_cell_shading(hdr[1], "D9EAF7")
    for label, value in rows:
        cells = table.add_row().cells
        set_cell_text(cells[0], label, True)
        if isinstance(value, list):
            set_cell_text(cells[1], "\n".join(f"- {item}" for item in value))
        else:
            set_cell_text(cells[1], value)
    return table


def configure_doc(document):
    section = document.sections[0]
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)
    styles = document.styles
    styles["Normal"].font.name = "Aptos"
    styles["Normal"].font.size = Pt(10)
    for name, size, color in [("Title", 18, "002B5C"), ("Heading 1", 15, "002B5C"), ("Heading 2", 12, "002B5C")]:
        style = styles[name]
        style.font.name = "Aptos"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)


def add_bullets(document, items):
    for item in items:
        para = document.add_paragraph(style="List Bullet")
        para.add_run(str(item or ""))


def all_lesson_resources(lesson):
    resources = []
    for key in ("bookSections", "ispring", "downloads", "handsOn", "textExports"):
        for item in lesson.get(key) or []:
            if isinstance(item, dict):
                resources.append(item)
    return resources


def resource_summary(course_root, lesson):
    rows = []
    for item in all_lesson_resources(lesson):
        label = item.get("label") or item.get("title") or item.get("path") or "Untitled resource"
        kind = item.get("type") or item.get("category") or item.get("role") or "resource"
        role = item.get("role") or item.get("sectionLabel") or item.get("category") or ""
        path = item.get("path") or item.get("previewPath") or item.get("downloadPath") or ""
        exists = "yes" if path and (course_root / path).exists() else ("n/a" if not path else "check")
        rows.append((short_text(label, 90), str(kind), str(role), exists))
    return rows[:24]


def section_text(course_root, lesson, section_pattern):
    regex = re.compile(section_pattern, re.I)
    for section in lesson.get("bookSections") or []:
        label = section.get("sectionLabel") or section.get("label") or ""
        if regex.search(label):
            return read_local_text(course_root, section.get("path")) or clean_text(section.get("textPreview"))
    return ""


def infer_codes(text):
    return sorted(set(re.findall(r"\b[A-F][1-3](?:\.\d+)?\b", text or "")))[:12]


def course_item(label, rel_path, role, category, item_type=None, source=None):
    p = Path(rel_path)
    return {
        "label": label,
        "type": item_type or p.suffix.lower().lstrip(".") or "document",
        "role": role,
        "category": category,
        "path": p.as_posix(),
        "downloadPath": p.as_posix(),
        "bytes": 0,
        "source": source or "locally generated teacher-prep resource",
        "sourceStatus": "local",
        "generatedAt": GENERATED_AT,
    }


def local_plan_item(label, rel_path, role):
    item = course_item(label, rel_path, role, role, "docx", "locally authored from indexed course materials")
    item["sourceStatus"] = "generated_from_local_course_materials"
    return item


def item_exists(course_root, item):
    if not item:
        return False
    for field in ("path", "downloadPath", "previewPath", "packagePath"):
        value = item.get(field)
        if value and (course_root / value).exists():
            return True
    return False


def upsert_path_item(items, item):
    path = (item.get("path") or item.get("downloadPath") or "").lower()
    label = (item.get("label") or "").lower()
    for index, existing in enumerate(items):
        existing_path = (existing.get("path") or existing.get("downloadPath") or "").lower()
        existing_label = (existing.get("label") or "").lower()
        if path and path == existing_path:
            items[index] = {**existing, **item}
            return "updated"
        if label and label == existing_label:
            items[index] = {**existing, **item}
            return "updated"
    items.append(item)
    return "added"


def update_bytes(course_root, item):
    path = item.get("path") or item.get("downloadPath")
    if path and (course_root / path).exists():
        item["bytes"] = (course_root / path).stat().st_size


def text_registry_item(course_code, label, material, all_unit_numbers, notes):
    item = dict(material)
    if item.get("path") and not item.get("previewPath") and item.get("type") in {"pdf", "html", "md"}:
        item["previewPath"] = item["path"]
    if item.get("path") and not item.get("downloadPath"):
        item["downloadPath"] = item["path"]
    category_role = f"{item.get('category', '')} {item.get('role', '')} {label}".lower()
    official = "curriculum" in category_role
    return {
        "id": f"{course_code.lower()}-{slug(label, 58).lower()}",
        "title": label,
        "label": label,
        "author": "Ontario Ministry of Education" if official else "Course resource",
        "type": "curriculum" if official else item.get("type", "document"),
        "units": all_unit_numbers,
        "copyrightStatus": "official_public_document" if official else "local_teacher_prep_reference",
        "sourceStatus": item.get("sourceStatus") or "local",
        "notes": notes,
        "materials": [item],
        "path": item.get("path"),
        "previewPath": item.get("previewPath"),
        "downloadPath": item.get("downloadPath"),
        "bytes": item.get("bytes"),
        "category": item.get("category"),
        "role": item.get("role"),
        "source": item.get("source"),
    }


def create_unit_plan(course_code, course_root, manifest, unit, profile, out_path):
    document = Document()
    configure_doc(document)
    document.add_heading(f"{course_code} Unit {unit.get('unit')} Plan", 0)
    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run(unit.get("title") or f"Unit {unit.get('unit')}").bold = True

    lessons = [lesson for lesson in unit.get("lessons") or [] if lesson.get("planningStatus") != "unit_overview"]
    focus = profile["unit_focus"].get(unit.get("unit"), profile["course_lens"])
    document.add_heading("Unit Snapshot", level=1)
    add_table(
        document,
        [
            ("Course / Strand Lens", f"{profile['subject']} ({profile['grade']}): {profile['course_lens']}"),
            ("Unit Focus", focus),
            ("Lesson Sequence", [f"Lesson {lesson.get('lesson')}: {lesson.get('title')}" for lesson in lessons]),
            ("Teacher Preparation Checklist", [
                "Open the localized lesson page and verify iSpring/H5P/video resources are playable before teaching.",
                "Keep ordinary files attached to their owning Moodle activity page; do not treat DOC/PDF/PPT as standalone playable resources.",
                "Prepare discussion prompts and assessment checkpoints from the localized Moodle text and attachments.",
                "Use the course health/audit report before packaging so missing local files are caught early.",
            ]),
        ],
    )

    document.add_heading("Big Ideas and Essential Questions", level=1)
    add_bullets(document, profile["essential_questions"])

    document.add_heading("Assessment and Feedback Arc", level=1)
    add_table(
        document,
        [
            ("Diagnostic / Minds On", "Use first-lesson prompts, KWL-style reflection, class discussion, or entry questions to surface prior knowledge."),
            ("Formative Evidence", "Use hands-on activities, practice questions, discussion posts, worksheets, and teacher observation to guide feedback."),
            ("Consolidation Evidence", "Use exit slips, summaries, reflections, H5P attempts, or short written responses to check progress."),
            ("Summative / AOL Evidence", "Use culminating tasks, major submissions, unit tests, rubrics, or project artifacts indexed in the course package."),
        ],
    )

    document.add_heading("Unit Sequence", level=1)
    seq = document.add_table(rows=1, cols=4)
    seq.style = "Table Grid"
    headers = ["Lesson", "Title", "Teacher Move", "Evidence to Collect"]
    for cell, header in zip(seq.rows[0].cells, headers):
        set_cell_text(cell, header, True)
        set_cell_shading(cell, "D9EAF7")
    for lesson in lessons:
        row = seq.add_row().cells
        set_cell_text(row[0], f"U{unit.get('unit')}L{lesson.get('lesson')}")
        set_cell_text(row[1], lesson.get("title"))
        set_cell_text(row[2], "Connect the localized page, playable resource, and attached task into a single teacher-led flow.")
        set_cell_text(row[3], "Student notes, activity result, attached submission, or consolidation response.")

    document.add_heading("Risk and Package QA Notes", level=1)
    add_bullets(
        document,
        [
            "Before upload, confirm this unit still has its unitPlan and every lesson has a lessonPlan in course-manifest.json.",
            "Check that H5P, video, and iSpring are standalone only when localized, while ordinary files remain attached to their page.",
            "Use ENG3U page shell conventions for page view, attached files, buttons, spacing, and playable-resource placement.",
        ],
    )
    document.save(out_path)


def create_lesson_plan(course_code, course_root, unit, lesson, profile, out_path):
    lesson_title = lesson.get("title") or f"Lesson {lesson.get('lesson')}"
    expectation_text = section_text(course_root, lesson, r"expectation")
    lesson_text = section_text(course_root, lesson, r"^lesson$")
    hands_text = section_text(course_root, lesson, r"hands")
    consolidation_text = section_text(course_root, lesson, r"consolidation")
    homework_text = section_text(course_root, lesson, r"homework")
    codes = infer_codes(expectation_text)
    resources = resource_summary(course_root, lesson)

    document = Document()
    configure_doc(document)
    document.add_heading(f"{course_code} U{unit.get('unit')}L{lesson.get('lesson')} Lesson Plan", 0)
    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run(lesson_title).bold = True

    document.add_heading("Lesson Identity", level=1)
    add_table(
        document,
        [
            ("Course", f"{course_code} - {profile['subject']} ({profile['grade']})"),
            ("Unit", f"Unit {unit.get('unit')}: {unit.get('title')}"),
            ("Lesson", f"Lesson {lesson.get('lesson')}: {lesson_title}"),
            ("Curriculum / Source Base", profile["curriculum_title"]),
            ("Planning Lens", profile["unit_focus"].get(unit.get("unit"), profile["course_lens"])),
        ],
    )

    document.add_heading("Learning Goals", level=1)
    goals = [
        f"Explain the central idea of {lesson_title} using precise course vocabulary.",
        "Use localized Moodle text, playable resources, and attachments as evidence for learning.",
        "Complete the hands-on, consolidation, and homework sequence with clear success criteria.",
    ]
    if codes:
        goals.append(f"Connect work to expectation code(s): {', '.join(codes)}.")
    add_bullets(document, goals)

    document.add_heading("Success Criteria", level=1)
    add_bullets(
        document,
        [
            "I can summarize the lesson concept in my own words and point to evidence from the course page.",
            "I can use the local iSpring/H5P/video or attached file without relying on an external resource.",
            "I can complete the assigned practice or submission using the instructions and files provided.",
            "I can use teacher feedback from the activity or consolidation checkpoint to improve my next response.",
        ],
    )

    document.add_heading("Indexed Lesson Resources", level=1)
    if resources:
        res_table = document.add_table(rows=1, cols=4)
        res_table.style = "Table Grid"
        for cell, header in zip(res_table.rows[0].cells, ["Resource", "Type", "Role / Section", "Local Path"]):
            set_cell_text(cell, header, True)
            set_cell_shading(cell, "D9EAF7")
        for label, kind, role, exists in resources:
            cells = res_table.add_row().cells
            set_cell_text(cells[0], label)
            set_cell_text(cells[1], kind)
            set_cell_text(cells[2], role)
            set_cell_text(cells[3], exists)
    else:
        document.add_paragraph("No lesson-level resources were indexed beyond the localized Moodle page text.")

    document.add_heading("Suggested Lesson Flow", level=1)
    flow_rows = [
        ("Minds On / Launch", short_text(expectation_text or lesson_text or f"Introduce {lesson_title} and connect it to the unit focus.", 360)),
        ("Teach / Model", short_text(lesson_text or "Use the localized lesson page and any playable iSpring/video/H5P as the main teaching sequence.", 360)),
        ("Guided / Hands On", short_text(hands_text or "Use the Hands On activity or teacher-created prompt as formative practice.", 360)),
        ("Consolidation", short_text(consolidation_text or "Use exit slip, reflection, summary task, or H5P attempt to check understanding.", 360)),
        ("Homework / Follow Up", short_text(homework_text or "Use indexed homework files and Moodle instructions for independent completion.", 360)),
    ]
    add_table(document, flow_rows)

    document.add_heading("Teacher Preparation Checklist", level=1)
    add_bullets(
        document,
        [
            "Open the localized lesson HTML page before class and confirm the visible order matches Moodle.",
            "Confirm all H5P/video/iSpring entries are local and playable; use standalone cards only for these localized media types.",
            "Confirm DOC/DOCX/PDF/PPT attachments are shown in the Files/attachments area of the owning page.",
            "Have one quick formative question ready in case a playable resource needs a teacher-led fallback.",
        ],
    )

    document.add_heading("Differentiation and Accessibility", level=1)
    add_bullets(
        document,
        [
            "Offer students the option to respond orally, visually, in writing, or through a structured organizer when appropriate.",
            "Preview vocabulary and model one example before independent work.",
            "Use chunked instructions and checkpoints for long activities or multi-step submissions.",
        ],
    )

    document.add_heading("Upload and Display QA", level=1)
    add_bullets(
        document,
        [
            "After packaging, verify this lesson page in the portal and confirm the file display matches ENG3U conventions.",
            "If a resource appears only inside a page, verify whether it is intentionally inline or should also be a localized standalone H5P/video/iSpring card.",
            "If an ordinary file appears as a standalone card, move it back under the owning page attachment list.",
        ],
    )
    document.save(out_path)


def write_source_files(course_code, course_root, manifest, profile, results):
    ensure_dir(course_root / "texts")
    source_path = course_root / "texts" / "SOURCES.md"
    existing = source_path.read_text(encoding="utf-8", errors="ignore") if source_path.exists() else ""
    header = f"# {course_code} Sources and Teacher-Prep Notes\n\n"
    body = (
        f"- Official curriculum reference: {profile['curriculum_title']}\n"
        f"- Official curriculum URL: {profile['curriculum_url']}\n"
        f"- Local Moodle/source base: {profile['source_url']}\n"
        f"- Teacher-prep promotion generated: {GENERATED_AT}\n\n"
        f"{profile['source_note']}\n\n"
        "Planning notes:\n"
        "- Unit and lesson plans are locally authored from indexed course materials for teacher preparation.\n"
        "- H5P/video/iSpring should be treated as playable/localized resources; ordinary files remain attached to their owning page.\n"
        "- ENG3U display conventions remain the shared front-end target for localized pages.\n"
    )
    if "Teacher-prep promotion generated" not in existing:
        source_path.write_text(header + body if not existing.strip() else existing.rstrip() + "\n\n" + body, encoding="utf-8")
        results["sourceNotesWritten"].append(course_code)

    if profile.get("curriculum_copy_from"):
        src = profile["curriculum_copy_from"]
        dest_rel = profile["curriculum_dest"]
        dest = course_root / dest_rel
        ensure_dir(dest.parent)
        if not dest.exists():
            shutil.copy2(src, dest)
            results["curriculumCopied"].append({"course": course_code, "from": str(src), "to": str(dest)})
        curriculum_rel = dest_rel.as_posix()
        curriculum_type = dest.suffix.lower().lstrip(".")
    else:
        dest_rel = Path("texts/ontario-curriculum/GLC2O-2024-curriculum-reference.md")
        dest = course_root / dest_rel
        ensure_dir(dest.parent)
        if not dest.exists():
            dest.write_text(
                "# GLC2O Official Curriculum Reference\n\n"
                f"Title: {profile['curriculum_title']}\n\n"
                f"Official URL: {profile['curriculum_url']}\n\n"
                "The Ontario Digital Curriculum Platform is JavaScript-rendered, so this local reference records the official course page used for teacher preparation.\n",
                encoding="utf-8",
            )
            results["curriculumCopied"].append({"course": course_code, "from": profile["curriculum_url"], "to": str(dest)})
        curriculum_rel = dest_rel.as_posix()
        curriculum_type = "md"

    manifest.setdefault("texts", [])
    curriculum_item = course_item(
        profile["curriculum_title"],
        curriculum_rel,
        "official_curriculum",
        "official_curriculum",
        curriculum_type,
        profile["curriculum_url"],
    )
    source_item = course_item(
        f"{course_code} Sources and Teacher-Prep Notes",
        "texts/SOURCES.md",
        "source_notes",
        "source_audit",
        "md",
        "local source audit",
    )
    update_bytes(course_root, curriculum_item)
    update_bytes(course_root, source_item)
    units = [int(unit.get("unit")) for unit in (manifest.get("units") or []) if unit.get("unit") is not None]
    upsert_path_item(
        manifest["texts"],
        text_registry_item(course_code, profile["curriculum_title"], curriculum_item, units, "Official curriculum/reference file indexed for teacher preparation."),
    )
    upsert_path_item(
        manifest["texts"],
        text_registry_item(course_code, f"{course_code} Sources and Teacher-Prep Notes", source_item, units, profile["source_note"]),
    )


def ensure_asm_course_outline(course_root, manifest, profile, results):
    if manifest.get("course", {}).get("code") != "ASM3M":
        return
    outline_rel = Path("plans/course/ASM3M_Course_Outline.md")
    outline_path = course_root / outline_rel
    ensure_dir(outline_path.parent)
    if not outline_path.exists():
        lines = [
            "# ASM3M Course Outline",
            "",
            "Course: ASM3M - Media Arts, Grade 11, University/College Preparation",
            f"Curriculum: {profile['curriculum_title']}",
            "",
            "Course organization:",
        ]
        for unit in manifest.get("units") or []:
            lines.append(f"- Unit {unit.get('unit')}: {unit.get('title')} ({len(unit.get('lessons') or [])} lessons)")
        lines.extend(
            [
                "",
                "Teacher-prep note: The original Moodle shell did not expose a separate course outline file, so this local outline summarizes the indexed course structure for planning and upload QA.",
            ],
        )
        outline_path.write_text("\n".join(lines) + "\n", encoding="utf-8")
        results["courseOutlinesWritten"].append("ASM3M")
    manifest.setdefault("courseDownloads", [])
    outline_item = course_item(
        "ASM3M Course Outline",
        outline_rel.as_posix(),
        "course_outline",
        "course_outline",
        "md",
        "locally generated from indexed course structure",
    )
    update_bytes(course_root, outline_item)
    upsert_path_item(manifest["courseDownloads"], outline_item)


def promote_course(course_code):
    profile = COURSE_PROFILES[course_code]
    course_root = COURSEWARE_ROOT / course_code
    manifest_path = course_root / "course-manifest.json"
    manifest = json.loads(manifest_path.read_text(encoding="utf-8"))
    results = {
        "course": course_code,
        "unitPlansCreated": [],
        "lessonPlansCreated": [],
        "sourceNotesWritten": [],
        "curriculumCopied": [],
        "courseOutlinesWritten": [],
    }

    write_source_files(course_code, course_root, manifest, profile, results)
    ensure_asm_course_outline(course_root, manifest, profile, results)

    for unit in manifest.get("units") or []:
        unit_no = unit.get("unit")
        if not unit.get("unitPlan") or not item_exists(course_root, unit.get("unitPlan")):
            out_dir = course_root / "plans" / "generated" / "unit-plans"
            ensure_dir(out_dir)
            out_path = out_dir / f"{course_code}-U{int(unit_no):02d}-unit-plan.docx"
            create_unit_plan(course_code, course_root, manifest, unit, profile, out_path)
            item = local_plan_item(
                f"{course_code} Unit {unit_no} Teacher Unit Plan",
                rel(out_path, course_root),
                "unit_plan",
            )
            update_bytes(course_root, item)
            unit["unitPlan"] = item
            results["unitPlansCreated"].append(item["path"])

        for lesson in unit.get("lessons") or []:
            if lesson.get("planningStatus") == "unit_overview":
                continue
            if lesson.get("lessonPlan") and item_exists(course_root, lesson.get("lessonPlan")):
                continue
            lesson_no = lesson.get("lesson")
            out_dir = course_root / "plans" / "generated" / "lesson-plans" / f"Unit-{int(unit_no):02d}"
            ensure_dir(out_dir)
            out_path = out_dir / f"{course_code}-U{int(unit_no):02d}-L{int(lesson_no):02d}-{slug(lesson.get('title'), 48)}-lesson-plan.docx"
            create_lesson_plan(course_code, course_root, unit, lesson, profile, out_path)
            item = local_plan_item(
                f"{course_code} U{unit_no}L{lesson_no} Teacher Lesson Plan - {lesson.get('title')}",
                rel(out_path, course_root),
                "lesson_plan",
            )
            update_bytes(course_root, item)
            lesson["lessonPlan"] = item
            results["lessonPlansCreated"].append(item["path"])

    manifest.setdefault("sourceAudit", {})
    manifest["sourceAudit"]["teacherPrepPromotion"] = {
        "promotedAt": GENERATED_AT,
        "standard": "ICS3U teacher-prep enrichment; ENG3U display conventions",
        "course": course_code,
        "unitPlansCreated": len(results["unitPlansCreated"]),
        "lessonPlansCreated": len(results["lessonPlansCreated"]),
        "note": "Generated plans are local teacher-facing scaffolds derived from the indexed Moodle course materials and official curriculum reference.",
    }
    manifest["generatedAt"] = GENERATED_AT
    manifest_path.write_text(json.dumps(manifest, indent=2, ensure_ascii=False) + "\n", encoding="utf-8")
    return results


def main():
    all_results = [promote_course(code) for code in ("GLC2O", "HFC3M", "ASM3M")]
    REPORT_PATH.write_text(json.dumps({"generatedAt": GENERATED_AT, "results": all_results}, indent=2, ensure_ascii=False) + "\n", encoding="utf-8")
    print(json.dumps({"generatedAt": GENERATED_AT, "results": all_results}, indent=2, ensure_ascii=False))


if __name__ == "__main__":
    main()
