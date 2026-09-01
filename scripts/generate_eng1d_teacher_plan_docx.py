import json
import re
from datetime import datetime, timezone
from html import unescape
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


WORKSPACE_ROOT = Path("D:/工作文件/SUNNYBROOK")
REPO_ROOT = WORKSPACE_ROOT / "ossd-course-portal"
COURSE_ROOT = WORKSPACE_ROOT / "courseware" / "ENG1D"
MANIFEST_PATH = COURSE_ROOT / "course-manifest.json"
REPORT_PATH = REPO_ROOT / "deployment" / "ENG1D-docx-plans-report.json"
GENERATED_AT = datetime.now(timezone.utc).replace(microsecond=0).isoformat()
TEMPLATE_ROOT = REPO_ROOT / "templates" / "teacher-plans"
LESSON_PLAN_TEMPLATE = TEMPLATE_ROOT / "mdm4u-style-lesson-plan-template.docx"
UNIT_PLAN_TEMPLATE = TEMPLATE_ROOT / "mdm4u-style-unit-plan-template.docx"
MDM4U_LESSON_PLAN_REFERENCE = (
    WORKSPACE_ROOT / "courseware" / "MDM4U" / "plans" / "source" / "Unit 1" / "Lesson Plan Unit 1 lesson 1.docx"
)
MDM4U_UNIT_PLAN_REFERENCE = WORKSPACE_ROOT / "courseware" / "MDM4U" / "plans" / "source" / "Unit 1" / "Unit 1 Plan .docx"


UNITS = {
    1: {
        "strand": "Reading and Literature Studies / Writing",
        "focus": "short-story elements, narrative craft, characterization, irony, conflict, point of view, symbolism, theme, and original short-story writing.",
        "essential": [
            "How do authors use literary elements to shape meaning and reader response?",
            "How can students support interpretation with precise evidence from a text?",
            "How does understanding craft help students plan and revise their own short story?",
        ],
    },
    2: {
        "strand": "Reading and Literature Studies / Oral Communication",
        "focus": "Shakespearean drama, context, plot, character motivation, dramatic conflict, themes, symbols, motifs, and evidence-based discussion.",
        "essential": [
            "How do language, context, and staging shape our understanding of Romeo and Juliet?",
            "How do characters' choices and social pressures drive dramatic conflict?",
            "How can students use quotations and scene evidence to explain a theme or motif?",
        ],
    },
    3: {
        "strand": "Reading and Literature Studies / Media and Critical Literacy",
        "focus": "novel pre-reading, chapter comprehension, character development, theme, symbol, conflict, inference, and sustained literary response.",
        "essential": [
            "How does a novel develop character, conflict, and theme over time?",
            "How can readers track evidence across chapters to revise interpretations?",
            "How do symbols and recurring ideas reveal a novel's larger message?",
        ],
    },
    4: {
        "strand": "Reading and Literature Studies / Cultural Texts",
        "focus": "mythology, oral tradition, archetypes, cultural context, Greek and Norse myth, heroic patterns, and comparative interpretation.",
        "essential": [
            "Why do myths and archetypes remain powerful across cultures and time periods?",
            "How do mythological stories communicate values, fears, and explanations?",
            "How can students compare texts while respecting cultural and historical context?",
        ],
    },
}


def clean_text(value):
    text = str(value or "")
    text = re.sub(r"<script[\s\S]*?</script>", " ", text, flags=re.I)
    text = re.sub(r"<style[\s\S]*?</style>", " ", text, flags=re.I)
    text = re.sub(r"<[^>]+>", " ", text)
    text = unescape(text)
    return re.sub(r"\s+", " ", text).strip()


def clean_label(value):
    text = clean_text(value)
    text = re.sub(r"^\s*[-:]+\s*", "", text)
    text = re.sub(r"\s*/\s*", "/", text)
    return text or "Untitled"


def read_local_text(relative_path):
    if not relative_path:
        return ""
    path = COURSE_ROOT / relative_path
    if not path.exists():
        return ""
    return clean_text(path.read_text(encoding="utf-8", errors="ignore"))


def short_text(value, limit=340):
    text = clean_text(value)
    if len(text) <= limit:
        return text
    clipped = text[:limit]
    cut = max(clipped.rfind(". "), clipped.rfind("; "), clipped.rfind(", "))
    if cut > 150:
        clipped = clipped[: cut + 1]
    return clipped.strip() + "..."


def find_section(lesson, pattern):
    regex = re.compile(pattern, re.I)
    for section in lesson.get("bookSections") or []:
        label = section.get("sectionLabel") or section.get("label") or ""
        if regex.search(label):
            return section
    return None


def section_text(lesson, pattern):
    section = find_section(lesson, pattern)
    if not section:
        return ""
    return read_local_text(section.get("path")) or clean_text(section.get("textPreview"))


def split_sentences(text, limit):
    text = clean_text(text)
    if not text:
        return []
    parts = re.split(r"(?<=[.;:])\s+|(?=\b(?:R|W|M|L|C|OC)\d(?:\.\d+)?\b)", text)
    cleaned = [short_text(part, 270) for part in parts if len(clean_text(part)) > 12]
    return cleaned[:limit] or [short_text(text, 270)]


def expectation_codes(text):
    codes = re.findall(r"\b(?:R|W|M|L|C|OC)\d(?:\.\d+)?\b", text or "", flags=re.I)
    return sorted({code.upper() for code in codes})[:18]


def extract_expectation_sections(text):
    text = clean_text(text)
    if not text:
        return {
            "overall": ["Use the localized Lesson Expectations page for exact wording."],
            "specific": ["Use the localized Lesson Expectations page for exact wording."],
        }
    overall_marker = re.search(r"\bOverall Expectations?:", text, flags=re.I)
    if overall_marker:
        text = text[overall_marker.start():]
    lower = text.lower()
    overall = text
    specific = ""
    for marker in ["specific lesson expectations", "specific expectations", "specific expectation"]:
        pos = lower.find(marker)
        if pos >= 0:
            overall = text[:pos]
            specific = text[pos:]
            break
    for marker in ["learning goals", "success criteria", "lesson "]:
        pos = overall.lower().find(marker)
        if pos > 25:
            overall = overall[:pos]
    for marker in ["learning goals", "success criteria"]:
        pos = specific.lower().find(marker)
        if pos >= 0:
            specific = specific[:pos]
    return {
        "overall": split_sentences(overall, 6),
        "specific": split_sentences(specific or text, 10),
    }


def learning_goals(expectations_text, lesson_title, unit_focus):
    match = re.search(
        r"Learning Goals?:\s*(.+?)(Success Criteria|Hands On|Consolidation|Homework|$)",
        expectations_text or "",
        flags=re.I,
    )
    if match:
        raw = match.group(1)
        parts = split_sentences(raw, 6)
        if len(parts) <= 1:
            parts = re.split(
                r"(?=By the end of the lesson\s+(?:I will|students will)|I can|Identify|Explain|Analyze|Analyse|Describe|Compare|Use|Create|Demonstrate|Write|Read)",
                raw,
                flags=re.I,
            )
        goals = [short_text(part, 180) for part in parts if len(clean_text(part)) > 12]
        if goals:
            return goals[:4]
    return [
        f"Read and interpret {lesson_title} using the literary vocabulary and reading strategies named in this unit.",
        f"Use specific textual evidence to explain how an author or text develops {short_text(unit_focus, 110)}",
        "Apply the lesson concept through the Hands On activity, consolidation task, and homework response.",
    ]


def success_criteria(lesson_title, codes):
    code_text = ", ".join(codes[:4]) if codes else "the listed lesson expectations"
    return [
        f"I can explain the main idea or literary skill in {lesson_title} using accurate English terminology.",
        f"I can connect my interpretation to {code_text} and cite specific evidence from the text or lesson resource.",
        "I can complete the practice activity and use feedback to improve my written or oral response.",
        "I can prepare a homework submission that follows the task instructions and uses clear paragraph-level communication.",
    ]


def all_resources(lesson):
    pools = ["downloads", "textExports", "handsOn", "videos", "ispring"]
    items = []
    for key in pools:
        for item in lesson.get(key) or []:
            if item:
                items.append(item)
    return items


def resources_matching(lesson, pattern):
    regex = re.compile(pattern, re.I)
    found = []
    for item in all_resources(lesson):
        haystack = " ".join(str(item.get(key, "")) for key in ("label", "type", "category", "role", "path", "previewPath", "parentSection"))
        if regex.search(haystack):
            found.append(item)
    return found


def resource_labels(items):
    return [clean_label(item.get("label") or item.get("title") or "Untitled resource") for item in items]


def role_label(value):
    label = str(value or "course").replace("_", " ").strip()
    if label.lower() == "handson":
        label = "hands on"
    return " ".join(part.capitalize() for part in label.split())


def lesson_materials(lesson):
    media = []
    seen_media = set()

    def add_media(prefix, label):
        entry = f"{prefix}: {clean_label(label)}"
        key = re.sub(r"\W+", "", entry).lower()
        if key not in seen_media:
            seen_media.add(key)
            media.append(entry)

    for item in lesson.get("ispring") or []:
        add_media("iSpring lesson module", item.get("label") or "lesson presentation")
    for item in lesson.get("handsOn") or []:
        add_media("Hands On H5P", item.get("label") or "H5P activity")
    for item in resources_matching(lesson, r"h5p"):
        add_media(f"{role_label(item.get('role'))} H5P", item.get("label") or "H5P activity")
    for item in lesson.get("videos") or []:
        add_media(f"{role_label(item.get('role'))} video", item.get("label") or "video")
    for item in resources_matching(lesson, r"mp4|video"):
        add_media(f"{role_label(item.get('role'))} video", item.get("label") or "video")
    files = resource_labels(resources_matching(lesson, r"docx?|doc\b|pdf|pptx?|xlsx?|txt|png|jpe?g|gif|worksheet|homework|handout|kwl|rubric"))[:12]
    return media[:12], files


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_text(cell, text, bold=False, size=9.5):
    cell.text = ""
    lines = str(text or "").split("\n")
    for line_index, line in enumerate(lines):
        paragraph = cell.paragraphs[0] if line_index == 0 else cell.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(0)
        run = paragraph.add_run(line)
        run.bold = bold
        run.font.name = "Calibri"
        run.font.size = Pt(size)


def set_row_text(row, text, bold=False, size=9.5):
    set_text(row.cells[0], text, bold=bold, size=size)


def template_document(kind):
    if kind == "lesson":
        path = LESSON_PLAN_TEMPLATE if LESSON_PLAN_TEMPLATE.exists() else MDM4U_LESSON_PLAN_REFERENCE
    elif kind == "unit":
        path = UNIT_PLAN_TEMPLATE if UNIT_PLAN_TEMPLATE.exists() else MDM4U_UNIT_PLAN_REFERENCE
    else:
        raise ValueError(f"Unknown plan template kind: {kind}")
    return Document(str(path))


def style_document(doc):
    section = doc.sections[0]
    section.top_margin = Inches(0.9)
    section.bottom_margin = Inches(0.9)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.12
    for name, size, color in [("Heading 1", 16, "002B57"), ("Heading 2", 12.5, "002B57"), ("Heading 3", 10.5, "002B57")]:
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)


def unit_summary(unit, unit_info, lessons):
    topics = ", ".join(clean_label(lesson.get("title")) for lesson in lessons[:5])
    return (
        f"This unit supports Grade 9 English teachers in moving students through {unit['title']} with a clear cycle: "
        "lesson expectations, guided iSpring instruction, Hands On practice, consolidation evidence, and homework submission. "
        f"The instructional focus is {unit_info['focus']} Early lessons include {topics}."
    )


def lesson_plan_doc(unit, lesson):
    unit_info = UNITS.get(unit["unit"], {"strand": "English", "focus": unit["title"], "essential": []})
    lesson_title = clean_label(lesson.get("title"))
    expectations = section_text(lesson, r"expectations")
    lesson_page = section_text(lesson, r"^Lesson$")
    consolidation = section_text(lesson, r"Consolidation")
    homework = section_text(lesson, r"Homework")
    codes = expectation_codes(expectations)
    media, files = lesson_materials(lesson)
    expectation_sections = extract_expectation_sections(expectations)
    goals = learning_goals(expectations or lesson_page, lesson_title, unit_info["focus"])
    criteria = success_criteria(lesson_title, codes)

    doc = template_document("lesson")
    style_document(doc)
    overview = doc.tables[0]
    rows = overview.rows
    set_text(rows[0].cells[0], "Lesson Plan", bold=True)
    set_text(rows[0].cells[2], "Subject: ENG1D English, Grade 9 Academic", bold=True)
    set_row_text(rows[1], f"Lesson Name: {lesson_title}")
    set_row_text(rows[2], f"Unit of Study: {unit['title']}")
    prior = (
        f"Connect this lesson to the previous reading, KWL reflection, or consolidation response in {unit['title']}. "
        "Ask students to name one literary term, reading strategy, or text detail they already understand before they begin the iSpring lesson."
    )
    set_row_text(rows[3], "PRIOR KNOWLEDGE What do my students already know? What key questions or strategies will activate prior knowledge? What connections can I help students make?\n" + prior)
    set_row_text(rows[4], "\n".join(["CURRICULUM EXPECTATIONS", "OVERALL", *expectation_sections["overall"], "SPECIFIC", *expectation_sections["specific"]]))
    set_row_text(rows[5], "\n".join(["LEARNING GOALS", "LEARNING GOALS What do I want students to know and be able to do?", *goals]))
    set_row_text(rows[6], "\n".join(["SUCCESS CRITERIA(S) How will students know they attained the goal? What tools can they use to self-assess?", *criteria]))
    set_text(rows[7].cells[0], "Assessment as Learning", bold=True)
    set_text(rows[7].cells[1], "Assessment for Learning", bold=True)
    set_text(rows[7].cells[2], "Assessment of Learning", bold=True)
    set_text(rows[8].cells[0], "[x] KWL / reflection\n[x] Self-check during Hands On\n[x] Exit slip\n[x] Revision notes\n[ ] Peer conference")
    set_text(rows[8].cells[1], "[x] Hands On H5P / practice quiz\n[x] Strategic questioning\n[x] Annotation or paragraph draft\n[x] Homework check\n[x] Consolidation response")
    set_text(rows[8].cells[2], "[x] Homework submission\n[x] Unit quiz/test/assignment where listed\n[x] Short written response\n[ ] Presentation / conference where assigned")
    set_row_text(rows[9], "What will I do?\nConfer with students during practice.\nObserve reading, annotation, and response habits.\nGrade submitted homework or AOL tasks when listed.")
    set_row_text(
        rows[10],
        "Accommodations: How will you change the lesson to meet the needs of individual students?\n"
        "Pre-teach key vocabulary; chunk the iSpring lesson into pause points; provide sentence frames for evidence-based responses; allow oral rehearsal before writing; offer a model annotation or paragraph starter; and use Hands On/consolidation evidence to decide what needs reteaching.",
    )
    material_lines = (media + files)[:12] or ["Localized Lesson Expectations, Lesson, Hands On, Consolidation, and Homework pages."]
    set_row_text(
        rows[11],
        "Materials and Resources: What do I need to prepare in advance? Equipment? Student resources? Teacher resources?\n"
        + "\n".join(material_lines)
        + f"\nSource status: reconstructed_from_moodle_content. Generated from localized ENG1D Moodle resources on {GENERATED_AT}.",
    )

    lesson_table = doc.tables[1]
    set_text(lesson_table.rows[0].cells[0], "DELIVERING THE LESSON", bold=True)
    lesson_rows = lesson_table.rows
    set_text(lesson_rows[1].cells[0], "Timing\n5-8\nminutes")
    set_text(lesson_rows[1].cells[1], "Grouping")
    set_text(lesson_rows[1].cells[4], "Minds On!\nOpen with a KWL, quick-write, quotation, image, or retrieval prompt. Have students identify the reading purpose and the success criteria before moving into the lesson media.")
    set_text(lesson_rows[1].cells[5], "Materials/Resources\nKWL chart or notebook\nLesson Expectations page")
    phases = [
        (
            2,
            "Timing\n20-30\nminutes",
            "Action!\nUse the Lesson page and iSpring module to teach the core reading or writing skill. Pause for text evidence, vocabulary, and think-aloud modelling rather than letting students passively watch.",
            "Materials/Resources\nLocalized lesson page\niSpring lesson module",
        ),
        (
            3,
            "Timing\n15-25\nminutes",
            "Action!\nUse the Hands On activity as formative practice. Ask students to explain why an answer or response is supported by the text, not only whether it is correct.",
            "Materials/Resources\n" + ("\n".join(media[:5]) if media else "Hands On page and H5P activity"),
        ),
        (
            4,
            "Timing\n10-15\nminutes",
            "Consolidation!\nUse the summary video, exit slip, or consolidation prompt to check transfer. Before dismissal, restate the homework product, file requirements, and submission location.",
            "Materials/Resources\n" + "\n".join(([short_text(consolidation, 160)] if consolidation else []) + ([short_text(homework, 160)] if homework else []) + files[:5]),
        ),
    ]
    for row_index, timing, action, materials in phases:
        row = lesson_rows[row_index]
        set_text(row.cells[0], timing)
        set_text(row.cells[1], "W\n[ ]")
        set_text(row.cells[2], "S\n[ ]")
        set_text(row.cells[3], "I\n[ ]")
        set_text(row.cells[4], action)
        set_text(row.cells[5], materials)

    notes = doc.tables[2]
    set_text(
        notes.rows[0].cells[0],
        "Notes:\nSupport: reduce reading load through chunking, vocabulary previews, modelled annotation, sentence frames, and oral rehearsal.\n"
        "Extension: ask students to compare texts, justify alternate interpretations, or revise a paragraph for stronger evidence and voice.\n"
        "Teacher reflection: note which literary term, reading habit, or paragraph skill should launch the next lesson.",
    )
    return doc


def activity_titles(unit, pattern):
    regex = re.compile(pattern, re.I)
    titles = []
    for lesson in unit.get("lessons") or []:
        for item in all_resources(lesson):
            haystack = " ".join(str(item.get(key, "")) for key in ("label", "type", "category", "role", "section", "path", "parentSection"))
            if regex.search(haystack):
                titles.append(item.get("label") or item.get("title") or "Untitled")
    return titles


def unit_plan_doc(unit):
    unit_info = UNITS.get(unit["unit"], {"strand": "English", "focus": unit["title"], "essential": []})
    lessons = unit.get("lessons") or []
    all_codes = []
    outlines = []
    for lesson in lessons:
        lesson_title = clean_label(lesson.get("title"))
        expectations = section_text(lesson, r"expectations")
        all_codes.extend(expectation_codes(expectations))
        media, files = lesson_materials(lesson)
        task = "; ".join((media[:2] or files[:2] or ["Lesson page, Hands On, consolidation, and homework resources"]))
        focus = (learning_goals(expectations, lesson_title, unit_info["focus"]) or [])[0]
        outlines.append(
            f"{lesson['id']}: {lesson_title}\n"
            f"Focus: {focus}\n"
            "Assessment: Hands On/H5P checks, consolidation evidence, homework review, and any listed AOL/evaluation task.\n"
            f"Resources: {task}"
        )
    codes = sorted(set(all_codes))[:24]
    evaluations = resource_labels((unit.get("unitResources") or {}).get("evaluations") or [])

    doc = template_document("unit")
    style_document(doc)
    if doc.paragraphs:
        doc.paragraphs[0].text = f"UNIT {unit['unit']}"
        doc.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    table = doc.tables[0]
    rows = table.rows
    for row_index, label in {0: "Unit Author", 5: "Unit Overview", 14: "Unit Foundation", 19: "Assessment Plan", 22: "Unit Details", 25: "Materials and Resources"}.items():
        set_text(rows[row_index].cells[0], label, bold=True)
    for row_index, label in [(1, "Name:"), (2, "School District:"), (3, "School Name:"), (4, "School City, Province:")]:
        set_row_text(rows[row_index], label)
    values = [
        (6, "Unit Title Name"),
        (7, unit["title"]),
        (8, "Unit Summary"),
        (9, unit_summary(unit, unit_info, lessons)),
        (10, "Year Level"),
        (11, "Grade 9, Academic English"),
        (12, "Approximate Time Needed"),
        (13, f"{max(len(lessons) * 3, 8)} instructional hours, plus homework, reading, drafting, and evaluation time"),
        (15, "Targeted Curriculum Expectations"),
        (
            16,
            "\n".join(
                [
                    f"Primary strand focus: {unit_info['strand']}",
                    f"Expectation evidence found across lessons: {', '.join(codes) if codes else 'see each Lesson Expectations page.'}",
                    f"Unit focus: {unit_info['focus']}",
                ]
            ),
        ),
        (17, "Learning Goals"),
        (
            18,
            "\n".join(
                [
                    f"Students will read, interpret, discuss, and write about {unit['title']} using accurate literary vocabulary.",
                    "Students will use specific textual evidence to support interpretations in oral and written responses.",
                    "Students will complete Hands On, consolidation, homework, reflection, and evaluation tasks as evidence of learning.",
                    "Essential Questions",
                    *unit_info.get("essential", []),
                ]
            ),
        ),
    ]
    for row_index, value in values:
        set_row_text(rows[row_index], value)
    set_text(rows[20].cells[0], "Assessment as Learning (ASL)", bold=True)
    set_text(rows[20].cells[2], "Assessment for Learning (AFL)", bold=True)
    set_text(rows[20].cells[3], "Assessment of Learning (AoL)", bold=True)
    set_text(rows[21].cells[0], "[x] KWL chart\n[x] Reflection and Learning Log routines\n[x] Exit slips\n[x] Self-assessment of reading/writing habits")
    set_text(rows[21].cells[2], "[x] Hands On/H5P checks\n[x] Strategic questioning\n[x] Annotation and paragraph drafts\n[x] Homework review\n[x] Consolidation responses")
    set_text(rows[21].cells[3], "[x] Unit quiz/test/assignment where listed\n" + ("[x] " + "\n[x] ".join(evaluations[:10]) if evaluations else "[x] Unit evaluation or culminating task listed in the Evaluation section."))
    set_row_text(rows[23], "Lesson and Assessment Outlines", bold=True)
    set_row_text(rows[24], "\n\n".join(outlines))
    set_text(rows[26].cells[0], "Technology", bold=True)
    tech = "\n".join(sorted(set(activity_titles(unit, r"ispring|h5p|mp4|video")))[:28] or ["Localized lesson pages, iSpring presentations, H5P activities, and consolidation videos listed in the manifest."])
    set_text(rows[26].cells[1], tech)
    set_text(rows[27].cells[0], "Printed", bold=True)
    printed = "\n".join(sorted(set(activity_titles(unit, r"docx?|doc\b|pdf|worksheet|homework|handout|kwl|rubric|assignment")))[:28] or ["Homework, KWL, worksheet, and activity files listed in the manifest."])
    set_text(rows[27].cells[1], printed)
    set_text(rows[28].cells[0], "Other Resources", bold=True)
    set_text(rows[28].cells[1], "Anecdotal notes of reading and discussion behaviours\nExit slips and consolidation responses\nStudent self-assessment and reflection\nKWL chart and Learning Log routines where present")
    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    return doc


def save_doc(doc, relative_path):
    path = COURSE_ROOT / relative_path
    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(path)
    return path.stat().st_size


def file_record(label, relative_path, category, role):
    path = COURSE_ROOT / relative_path
    preview = f"previews-html/{relative_path}.html"
    return {
        "label": label,
        "type": "docx",
        "category": category,
        "role": role,
        "path": relative_path.replace("\\", "/"),
        "bytes": path.stat().st_size,
        "source": "reconstructed_from_moodle_content",
        "sourceStatus": "locally_authored_from_course_materials",
        "teacherUse": "Teacher planning aid reconstructed from localized ENG1D course content; not an original Moodle teacher packet.",
        "templateReference": "MDM4U/ENG3U teacher plan format and local mdm4u-style plan templates.",
        "previewPath": preview.replace("\\", "/"),
    }


def main():
    manifest = json.loads(MANIFEST_PATH.read_text(encoding="utf-8"))
    unit_count = 0
    lesson_count = 0
    written = []
    for unit in manifest.get("units") or []:
        unit_rel = f"plans/generated/unit-plans/U{unit['unit']:02d}-unit-plan.docx"
        save_doc(unit_plan_doc(unit), unit_rel)
        unit["unitPlan"] = file_record(f"Unit Plan - {unit['title']}", unit_rel, "unit_plan", "unit_plan")
        unit_count += 1
        written.append(unit_rel)
        for lesson in unit.get("lessons") or []:
            lesson_rel = f"plans/generated/lesson-plans/{lesson['id']}-lesson-plan.docx"
            save_doc(lesson_plan_doc(unit, lesson), lesson_rel)
            lesson["lessonPlan"] = file_record(
                f"Lesson Plan - Unit {unit['unit']} Lesson {lesson['lesson']}",
                lesson_rel,
                "lesson_plan",
                "lesson_plan",
            )
            lesson["resourceCounts"] = {**(lesson.get("resourceCounts") or {}), "lessonPlan": 1}
            lesson_count += 1
            written.append(lesson_rel)

    manifest["generatedAt"] = GENERATED_AT
    manifest["sourceAudit"] = {
        **(manifest.get("sourceAudit") or {}),
        "teacherPlanGeneration": {
            "generatedAt": GENERATED_AT,
            "format": "docx",
            "unitPlans": unit_count,
            "lessonPlans": lesson_count,
            "source": "reconstructed_from_moodle_content",
            "basis": [
                "localized Moodle Lesson Expectations pages",
                "localized Moodle Lesson, Hands On, Consolidation, and Homework sections",
                "localized iSpring, H5P, video, worksheet, exit slip, KWL/reflection, and homework submission records",
                "ENG1D unit sequence and literary focus visible in localized Moodle course books",
                "MDM4U/ENG3U teacher plan formatting conventions and local mdm4u-style plan templates",
            ],
            "boundary": "Teacher planning aids only; no answer keys, rubrics, copyrighted textbook excerpts, or Moodle-original teacher packet documents were created.",
        },
    }
    MANIFEST_PATH.write_text(json.dumps(manifest, indent=2, ensure_ascii=False) + "\n", encoding="utf-8")
    REPORT_PATH.parent.mkdir(parents=True, exist_ok=True)
    REPORT_PATH.write_text(
        json.dumps({"course": "ENG1D", "unitPlans": unit_count, "lessonPlans": lesson_count, "written": written, "generatedAt": GENERATED_AT}, indent=2, ensure_ascii=False) + "\n",
        encoding="utf-8",
    )
    print(json.dumps({"course": "ENG1D", "unitPlans": unit_count, "lessonPlans": lesson_count, "generatedAt": GENERATED_AT}, indent=2))


if __name__ == "__main__":
    main()
