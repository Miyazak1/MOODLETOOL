import json
import re
import shutil
from datetime import datetime, timezone
from html import unescape
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


WORKSPACE_ROOT = Path("D:/工作文件/SUNNYBROOK")
COURSEWARE_ROOT = WORKSPACE_ROOT / "courseware"
REPO_ROOT = WORKSPACE_ROOT / "ossd-course-portal"
REPORT_PATH = REPO_ROOT / "deployment" / "teacher-prep-promotion-HHS4U-LKBDU-BBI1O.json"
GENERATED_AT = datetime.now(timezone.utc).replace(microsecond=0).isoformat()


PROFILES = {
    "HHS4U": {
        "subject": "Families in Canada",
        "grade": "Grade 12, University Preparation",
        "curriculum_title": "The Ontario Curriculum: Social Sciences and Humanities, Grades 9 to 12, 2013 (Revised)",
        "curriculum_url": "https://www.edu.gov.on.ca/eng/curriculum/secondary/ssciences9to122013.pdf",
        "source_url": "https://www.esunnybrook.com/course/view.php?id=54",
        "course_lens": "sociological, psychological, and anthropological inquiry into families, intimate relationships, parenting, adulthood, aging, diversity, research, and social change.",
        "unit_focus": {
            1: "social science inquiry, definitions and functions of family, family change, historical perspectives, and introductory research habits.",
            2: "identity, diversity, gender, culture, family roles, equity, and individual development in a diverse society.",
            3: "intimate relationships, marriage, commitment, conflict, relationship change, and diverse family structures.",
            4: "parent-child relationships, parenting styles, child development, family responsibility, and social supports.",
            5: "midlife, aging, later adulthood, caregiving, intergenerational relationships, and family transitions.",
            6: "culminating inquiry project planning, research evidence, synthesis, presentation, and reflection.",
            7: "final examination review, synthesis of course concepts, and final assessment preparation.",
        },
        "essential_questions": [
            "How do social scientists explain the changing forms and functions of families in Canada?",
            "How do culture, identity, power, and social institutions shape family experiences?",
            "How can research evidence help students evaluate claims about relationships, parenting, adulthood, and family change?",
        ],
        "source_note": "HHS4U teacher-prep planning is based on the localized Moodle course, the indexed course outline/planning files, and the official Social Sciences and Humanities curriculum already included in the course package.",
    },
    "LKBDU": {
        "subject": "International Languages, Simplified Chinese",
        "grade": "Level 4, University Preparation",
        "curriculum_title": "The Ontario Curriculum, Grades 9 to 12: Classical Studies and International Languages, 2016 (Revised)",
        "curriculum_url": "https://www.dcp.edu.gov.on.ca/en/curriculum/classical-studies-intl-languages/grades/lbadu-ldydu/context/enduring-ideas",
        "source_url": "https://www.esunnybrook.com/course/view.php?id=45",
        "course_lens": "advanced Chinese communication, literary response, cultural inquiry, intercultural understanding, metacognition, and final performance evidence.",
        "unit_focus": {
            1: "Su Dongpo's poetry and biography, classical-cultural context, interpretation, and response.",
            2: "drama study through Thunderstorm, character, conflict, performance, and historical-cultural context.",
            3: "novel study through The Joy Luck Club, identity, diaspora, voice, cultural comparison, and literary response.",
            4: "prose appreciation and writing, style, rhetoric, audience, and refined written expression.",
            5: "final project, final exam, portfolio evidence, and synthesis of language-learning growth.",
        },
        "essential_questions": [
            "How can students use Simplified Chinese to communicate ideas with accuracy, nuance, and cultural awareness?",
            "How do literary and informational texts reveal cultural values, identity, history, and perspective?",
            "How can final performance evidence show growth across reading, writing, listening, speaking, and intercultural understanding?",
        ],
        "source_note": "LKBDU teacher-prep planning is based on the localized Moodle course, local U1-U4 daily lesson plans, and an indexed official Ontario Classical Studies and International Languages curriculum reference.",
        "official_reference_only": True,
    },
    "BBI1O": {
        "subject": "Introduction to Business",
        "grade": "Grade 9, Open",
        "curriculum_title": "The Ontario Curriculum, Grades 9 and 10: Business Studies, 2006 (Revised) - legacy BBI1O/BBI2O reference",
        "curriculum_url": "https://www.edu.gov.on.ca/eng/curriculum/secondary/business910currb.pdf",
        "source_url": "https://www.esunnybrook.com/course/view.php?id=32",
        "course_lens": "personal finance, business fundamentals, functions of business, entrepreneurship, international business, culminating application, and final assessment.",
        "unit_focus": {
            1: "personal finance, renting, home ownership, budgets, moving costs, event planning, grocery shopping, and future planning.",
            2: "business fundamentals, needs and wants, production, marketing, accounting, ICT, human resources, and business functions.",
            3: "entrepreneurship, opportunity recognition, venture planning, product/service ideas, risk, and business pitch evidence.",
            4: "international business, imports/exports, trade partners, globalization, ethics, and Canadian business connections.",
            5: "culminating project planning, applied business evidence, presentation, and reflection.",
            6: "final examination review, knowledge consolidation, and final assessment preparation.",
            7: "teacher evaluation comments, learning skills, work habits, and final reporting evidence.",
        },
        "essential_questions": [
            "How do individuals and businesses make responsible financial and operational decisions?",
            "How do business functions work together to meet needs and create value?",
            "How can students apply business thinking to entrepreneurship, international trade, and everyday financial choices?",
        ],
        "source_note": "BBI1O is treated as a local legacy course package. The 2006 Business Studies BBI1O/BBI2O curriculum and The World of Business text reference are indexed from the existing local BBI2O source set; current Ontario 2024 business curriculum has replaced the old BBI1O/BBI2O course codes.",
        "copy_resources": [
            {
                "from": COURSEWARE_ROOT / "BBI2O" / "texts" / "ontario-curriculum" / "business910currb.pdf",
                "to": Path("texts/ontario-curriculum/business910currb.pdf"),
                "label": "The Ontario Curriculum, Grades 9 and 10: Business Studies, 2006 (Revised) - legacy BBI1O/BBI2O reference",
                "role": "official_curriculum",
                "category": "official_curriculum",
                "type": "pdf",
                "source": "https://www.edu.gov.on.ca/eng/curriculum/secondary/business910currb.pdf",
            },
            {
                "from": COURSEWARE_ROOT / "BBI2O" / "texts" / "textbook" / "the-world-of-business-5th-edition.pdf",
                "to": Path("texts/textbook/the-world-of-business-5th-edition.pdf"),
                "label": "The World of Business, 5th Edition - local BBI1O/BBI2O reference text",
                "role": "textbook_reference",
                "category": "textbook",
                "type": "pdf",
                "source": "local BBI2O source set",
            },
        ],
    },
}


def clean_text(value):
    text = str(value or "")
    text = re.sub(r"<script[\s\S]*?</script>", " ", text, flags=re.I)
    text = re.sub(r"<style[\s\S]*?</style>", " ", text, flags=re.I)
    text = re.sub(r"<[^>]+>", " ", text)
    text = unescape(text)
    return re.sub(r"\s+", " ", text).strip()


def read_text(course_root, path):
    if not path:
        return ""
    file_path = course_root / path
    if not file_path.exists() or file_path.suffix.lower() not in {".html", ".htm", ".md", ".txt"}:
        return ""
    return clean_text(file_path.read_text(encoding="utf-8", errors="ignore"))


def short(value, limit=300):
    text = clean_text(value)
    if len(text) <= limit:
        return text
    clipped = text[:limit]
    cut = max(clipped.rfind(". "), clipped.rfind("; "), clipped.rfind(", "))
    if cut > 130:
        clipped = clipped[: cut + 1]
    return clipped.strip() + "..."


def slug(value, limit=56):
    text = re.sub(r"[^A-Za-z0-9]+", "-", str(value or "")).strip("-")
    return (text[:limit].strip("-") or "item")


def ensure_dir(path):
    path.mkdir(parents=True, exist_ok=True)


def rel(path, course_root):
    return path.relative_to(course_root).as_posix()


def set_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell(cell, text, bold=False):
    cell.text = ""
    para = cell.paragraphs[0]
    run = para.add_run(str(text or ""))
    run.bold = bold
    run.font.name = "Calibri"
    run.font.size = Pt(9.5)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP


def configure_doc(doc):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    styles = doc.styles
    styles["Normal"].font.name = "Calibri"
    styles["Normal"].font.size = Pt(11)
    for name, size, color in [("Title", 18, "0B2545"), ("Heading 1", 16, "2E74B5"), ("Heading 2", 13, "2E74B5")]:
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)


def add_table(doc, rows):
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    for cell, header in zip(table.rows[0].cells, ["Planning Area", "Teacher Notes"]):
        set_cell(cell, header, True)
        set_shading(cell, "F2F4F7")
    for label, value in rows:
        cells = table.add_row().cells
        set_cell(cells[0], label, True)
        if isinstance(value, list):
            set_cell(cells[1], "\n".join(f"- {item}" for item in value))
        else:
            set_cell(cells[1], value)
    return table


def add_bullets(doc, items):
    for item in items:
        para = doc.add_paragraph(style="List Bullet")
        para.add_run(str(item or ""))


def all_resources(lesson):
    out = []
    for key in ("bookSections", "ispring", "downloads", "handsOn", "textExports"):
        for item in lesson.get(key) or []:
            if isinstance(item, dict):
                out.append(item)
    return out


def lesson_resource_rows(course_root, lesson):
    rows = []
    for item in all_resources(lesson):
        label = item.get("label") or item.get("title") or item.get("path") or "Untitled resource"
        item_type = item.get("type") or item.get("category") or "resource"
        role = item.get("role") or item.get("sectionLabel") or item.get("category") or ""
        path = item.get("path") or item.get("previewPath") or item.get("downloadPath") or ""
        exists = "yes" if path and (course_root / path).exists() else ("n/a" if not path else "check")
        rows.append((short(label, 90), str(item_type), str(role), exists))
    return rows[:24]


def section_text(course_root, lesson, pattern):
    regex = re.compile(pattern, re.I)
    for section in lesson.get("bookSections") or []:
        label = section.get("sectionLabel") or section.get("label") or ""
        if regex.search(label):
            return read_text(course_root, section.get("path")) or clean_text(section.get("textPreview"))
    return ""


def infer_codes(text):
    return sorted(set(re.findall(r"\b[A-F][1-3](?:\.\d+)?\b", text or "")))[:10]


def item_exists(course_root, item):
    if not isinstance(item, dict):
        return False
    for field in ("path", "downloadPath", "previewPath", "packagePath"):
        value = item.get(field)
        if value and (course_root / value).exists():
            return True
    return False


def generated_plan_path(item, plan_kind):
    if not isinstance(item, dict):
        return False
    path = str(item.get("path") or item.get("downloadPath") or "").replace("\\", "/")
    return path.startswith(f"plans/generated/{plan_kind}/") and path.lower().endswith(".docx")


def resource_item(label, rel_path, role, category, item_type=None, source="locally generated"):
    path = Path(rel_path).as_posix()
    return {
        "label": label,
        "type": item_type or Path(rel_path).suffix.lower().lstrip(".") or "document",
        "role": role,
        "category": category,
        "path": path,
        "downloadPath": path,
        "bytes": 0,
        "source": source,
        "sourceStatus": "local",
        "generatedAt": GENERATED_AT,
    }


def plan_item(label, rel_path, role):
    item = resource_item(label, rel_path, role, role, "docx", "locally authored from indexed course materials")
    item["sourceStatus"] = "generated_from_local_course_materials"
    return item


def update_bytes(course_root, item):
    path = item.get("path") or item.get("downloadPath")
    if path and (course_root / path).exists():
        item["bytes"] = (course_root / path).stat().st_size


def upsert(items, item):
    if not isinstance(items, list):
        return "skipped"
    key = (item.get("path") or item.get("downloadPath") or "").lower()
    label = (item.get("label") or "").lower()
    for idx, existing in enumerate(items):
        if not isinstance(existing, dict):
            continue
        existing_key = (existing.get("path") or existing.get("downloadPath") or "").lower()
        existing_label = (existing.get("label") or "").lower()
        if key and key == existing_key:
            items[idx] = {**existing, **item}
            return "updated"
        if label and label == existing_label:
            items[idx] = {**existing, **item}
            return "updated"
    items.append(item)
    return "added"


def text_registry_item(course_code, label, material, all_unit_numbers, notes):
    item = dict(material)
    if item.get("path") and not item.get("previewPath") and item.get("type") in {"pdf", "html", "md"}:
        item["previewPath"] = item["path"]
    if item.get("path") and not item.get("downloadPath"):
        item["downloadPath"] = item["path"]
    category_role = f"{item.get('category', '')} {item.get('role', '')} {label}".lower()
    official = "curriculum" in category_role
    return {
        "id": f"{course_code.lower()}-{slug(label).lower()}",
        "title": label,
        "label": label,
        "author": "Ontario Ministry of Education" if official else "Course resource",
        "type": "curriculum" if official else item.get("type", "document"),
        "units": all_unit_numbers,
        "copyrightStatus": "official_public_document" if official else "local_teacher_prep_reference",
        "sourceStatus": item.get("sourceStatus") or "local",
        "notes": notes,
        "materials": [item],
        "path": item.get("path"),
        "previewPath": item.get("previewPath"),
        "downloadPath": item.get("downloadPath"),
        "bytes": item.get("bytes"),
        "category": item.get("category"),
        "role": item.get("role"),
        "source": item.get("source"),
    }


def create_unit_doc(course_code, course_root, unit, profile, out_path):
    doc = Document()
    configure_doc(doc)
    doc.add_heading(f"{course_code} Unit {unit.get('unit')} Teacher Unit Plan", 0)
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run(unit.get("title") or f"Unit {unit.get('unit')}").bold = True
    lessons = [lesson for lesson in unit.get("lessons") or [] if lesson.get("planningStatus") != "unit_overview"]
    focus = profile["unit_focus"].get(unit.get("unit"), profile["course_lens"])
    doc.add_heading("Unit Snapshot", level=1)
    add_table(
        doc,
        [
            ("Course / Curriculum Lens", f"{profile['subject']} ({profile['grade']}): {profile['course_lens']}"),
            ("Unit Focus", focus),
            ("Lesson Sequence", [f"Lesson {lesson.get('lesson')}: {lesson.get('title')}" for lesson in lessons]),
            ("Teacher Preparation Checklist", [
                "Open localized activity/page resources before teaching and confirm attached files are visible.",
                "Use playable H5P/video/iSpring only when localized; keep ordinary files attached to the owning page.",
                "Prepare discussion prompts and assessment checkpoints from the localized Moodle text and documents.",
                "Run teacher-prep and package QA before upload so missing local files or wrong display roles are caught early.",
            ]),
        ],
    )
    doc.add_heading("Big Ideas and Essential Questions", level=1)
    add_bullets(doc, profile["essential_questions"])
    doc.add_heading("Assessment and Feedback Arc", level=1)
    add_table(
        doc,
        [
            ("Diagnostic / Entry", "Use KWL, brief discussion, vocabulary check, or opening reflection to surface prior knowledge."),
            ("Formative Evidence", "Use Moodle activities, practice worksheets, discussion, teacher observation, or checkpoints for feedback."),
            ("Consolidation Evidence", "Use exit slips, reflections, short answers, practice quizzes, or submitted work to verify understanding."),
            ("Summative / AOL Evidence", "Use final submissions, projects, tests, rubrics, presentation products, or portfolio evidence indexed in the course."),
        ],
    )
    doc.add_heading("Unit Sequence", level=1)
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    for cell, header in zip(table.rows[0].cells, ["Lesson", "Title", "Teacher Move", "Evidence to Collect"]):
        set_cell(cell, header, True)
        set_shading(cell, "F2F4F7")
    for lesson in lessons:
        row = table.add_row().cells
        set_cell(row[0], f"U{unit.get('unit')}L{lesson.get('lesson')}")
        set_cell(row[1], lesson.get("title"))
        set_cell(row[2], "Connect the localized page, attachments, and activities into a teacher-led learning flow.")
        set_cell(row[3], "Student response, submitted file, activity result, reflection, or discussion evidence.")
    doc.add_heading("Upload and Display QA", level=1)
    add_bullets(
        doc,
        [
            "Confirm this unit has a manifest unitPlan entry and every lesson has a lessonPlan entry.",
            "Confirm ENG3U-style page shell/file rows are preserved after packaging.",
            "Confirm ordinary files do not become standalone cards, while localized H5P/video/iSpring may appear as playable resources.",
        ],
    )
    doc.save(out_path)


def create_lesson_doc(course_code, course_root, unit, lesson, profile, out_path):
    title = lesson.get("title") or f"Lesson {lesson.get('lesson')}"
    expectations = section_text(course_root, lesson, r"expectation")
    lesson_body = section_text(course_root, lesson, r"lesson|activity|assign|discussion")
    hands = section_text(course_root, lesson, r"hands|practice")
    consolidation = section_text(course_root, lesson, r"consolidation|exit|reflection")
    homework = section_text(course_root, lesson, r"homework|submission|dropbox")
    codes = infer_codes(expectations)
    resources = lesson_resource_rows(course_root, lesson)
    doc = Document()
    configure_doc(doc)
    doc.add_heading(f"{course_code} U{unit.get('unit')}L{lesson.get('lesson')} Teacher Lesson Plan", 0)
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run(title).bold = True
    doc.add_heading("Lesson Identity", level=1)
    add_table(
        doc,
        [
            ("Course", f"{course_code} - {profile['subject']} ({profile['grade']})"),
            ("Unit", f"Unit {unit.get('unit')}: {unit.get('title')}"),
            ("Lesson", f"Lesson {lesson.get('lesson')}: {title}"),
            ("Curriculum / Source Base", profile["curriculum_title"]),
            ("Planning Lens", profile["unit_focus"].get(unit.get("unit"), profile["course_lens"])),
        ],
    )
    doc.add_heading("Learning Goals", level=1)
    goals = [
        f"Explain the main concept or task in {title} using accurate course vocabulary.",
        "Use the localized Moodle text, activity instructions, and attached files as learning evidence.",
        "Complete the activity/submission sequence with a clear link to success criteria and feedback.",
    ]
    if codes:
        goals.append(f"Connect work to expectation code(s): {', '.join(codes)}.")
    add_bullets(doc, goals)
    doc.add_heading("Success Criteria", level=1)
    add_bullets(
        doc,
        [
            "I can identify the purpose of the activity and the evidence I need to produce.",
            "I can use the local course page and attachments without relying on unavailable external resources.",
            "I can explain my reasoning, creative decision, research evidence, or language/business/social-science connection clearly.",
            "I can use teacher feedback or a consolidation prompt to improve the next task.",
        ],
    )
    doc.add_heading("Indexed Lesson Resources", level=1)
    if resources:
        table = doc.add_table(rows=1, cols=4)
        table.style = "Table Grid"
        for cell, header in zip(table.rows[0].cells, ["Resource", "Type", "Role / Section", "Local Path"]):
            set_cell(cell, header, True)
            set_shading(cell, "F2F4F7")
        for label, item_type, role, exists in resources:
            cells = table.add_row().cells
            set_cell(cells[0], label)
            set_cell(cells[1], item_type)
            set_cell(cells[2], role)
            set_cell(cells[3], exists)
    else:
        doc.add_paragraph("No additional lesson-level resources were indexed beyond the localized Moodle activity text.")
    doc.add_heading("Suggested Lesson Flow", level=1)
    add_table(
        doc,
        [
            ("Minds On / Launch", short(expectations or lesson_body or f"Frame {title} in the unit sequence.", 360)),
            ("Teach / Model", short(lesson_body or "Use the localized Moodle page/activity and model the first step or response expectation.", 360)),
            ("Guided / Hands On", short(hands or "Use an activity, partner talk, short prompt, or worked example as formative practice.", 360)),
            ("Consolidation", short(consolidation or "Use an exit slip, reflection, debrief, discussion post, or short written response to check understanding.", 360)),
            ("Homework / Follow Up", short(homework or "Use indexed activity instructions and attached files for independent completion or submission.", 360)),
        ],
    )
    doc.add_heading("Teacher Preparation Checklist", level=1)
    add_bullets(
        doc,
        [
            "Open the localized page before class and compare visible instructions with the course card.",
            "Confirm attached DOC/DOCX/PDF/PPT files remain in the Files/attachments area.",
            "Confirm localized H5P/video/iSpring, if any, is playable and appears in the correct flow order.",
            "Prepare a backup teacher-led prompt if a media activity needs facilitation.",
        ],
    )
    doc.add_heading("Upload and Display QA", level=1)
    add_bullets(
        doc,
        [
            "After packaging, check this page for ENG3U-style card shell, file rows, spacing, and action labels.",
            "If a normal file appears as a standalone playable card, move it back under the owning page attachments.",
            "If a local H5P/video/iSpring is embedded in the page but missing as a standalone playable resource, flag it for course-resource repair.",
        ],
    )
    doc.save(out_path)


def source_and_curriculum(course_code, course_root, manifest, profile, results):
    ensure_dir(course_root / "texts")
    source_path = course_root / "texts" / "SOURCES.md"
    existing = source_path.read_text(encoding="utf-8", errors="ignore") if source_path.exists() else ""
    block = (
        f"\n\n## Teacher-Prep Promotion {GENERATED_AT}\n\n"
        f"- Official/reference curriculum: {profile['curriculum_title']}\n"
        f"- Curriculum URL/reference: {profile['curriculum_url']}\n"
        f"- Course/source base: {profile['source_url']}\n"
        f"- Note: {profile['source_note']}\n"
    )
    if "Teacher-Prep Promotion" not in existing:
        source_path.write_text((existing or f"# {course_code} Sources and Localization Notes\n").rstrip() + block + "\n", encoding="utf-8")
        results["sourceNotesWritten"].append(course_code)
    manifest["texts"] = [
        item
        for item in (manifest.get("texts") or [])
        if isinstance(item, dict) and isinstance(item.get("units"), list) and isinstance(item.get("materials"), list)
    ]
    units = [int(unit.get("unit")) for unit in (manifest.get("units") or []) if unit.get("unit") is not None]
    source_item = resource_item(f"{course_code} Sources and Teacher-Prep Notes", "texts/SOURCES.md", "source_notes", "source_audit", "md", "local source audit")
    update_bytes(course_root, source_item)
    upsert(manifest["texts"], text_registry_item(course_code, f"{course_code} Sources and Teacher-Prep Notes", source_item, units, profile["source_note"]))

    copied_any = False
    for spec in profile.get("copy_resources") or []:
        dest = course_root / spec["to"]
        ensure_dir(dest.parent)
        if not dest.exists() and spec["from"].exists():
            shutil.copy2(spec["from"], dest)
            copied_any = True
            results["copiedResources"].append({"course": course_code, "from": str(spec["from"]), "to": str(dest)})
        item = resource_item(spec["label"], spec["to"].as_posix(), spec["role"], spec["category"], spec["type"], spec["source"])
        update_bytes(course_root, item)
        upsert(manifest["texts"], text_registry_item(course_code, spec["label"], item, units, f"{spec['label']} indexed for teacher preparation."))

    if profile.get("official_reference_only"):
        ref_rel = Path("texts/ontario-curriculum/classical-studies-international-languages-2016-reference.md")
        ref = course_root / ref_rel
        ensure_dir(ref.parent)
        if not ref.exists():
            ref.write_text(
                f"# {profile['curriculum_title']}\n\n"
                f"Official DCP reference: {profile['curriculum_url']}\n\n"
                "This local reference indexes the official Ontario curriculum page for LKBDU teacher preparation.\n",
                encoding="utf-8",
            )
            results["copiedResources"].append({"course": course_code, "from": profile["curriculum_url"], "to": str(ref)})
        item = resource_item(profile["curriculum_title"], ref_rel.as_posix(), "official_curriculum", "official_curriculum", "md", profile["curriculum_url"])
        update_bytes(course_root, item)
        upsert(manifest["texts"], text_registry_item(course_code, profile["curriculum_title"], item, units, "Official Ontario curriculum/reference file indexed for teacher preparation."))

    if course_code == "HHS4U":
        # The HHS4U course already has the curriculum PDF in courseDownloads. Add it to texts as an explicit curriculum reference.
        for item in manifest.get("courseDownloads") or []:
            label = item.get("label") or ""
            path = item.get("path") or ""
            if "Social Science and Humanities Ontario Curriculum" in label and path:
                ref_item = resource_item(label, path, "official_curriculum", "official_curriculum", item.get("type") or "pdf", item.get("source") or profile["curriculum_url"])
                update_bytes(course_root, ref_item)
                upsert(manifest["texts"], text_registry_item(course_code, label, ref_item, units, "Official curriculum/reference file indexed for teacher preparation."))
                copied_any = True
                break
    return copied_any


def promote(course_code):
    profile = PROFILES[course_code]
    course_root = COURSEWARE_ROOT / course_code
    manifest_path = course_root / "course-manifest.json"
    manifest = json.loads(manifest_path.read_text(encoding="utf-8"))
    results = {
        "course": course_code,
        "unitPlansCreated": [],
        "lessonPlansCreated": [],
        "sourceNotesWritten": [],
        "copiedResources": [],
    }
    source_and_curriculum(course_code, course_root, manifest, profile, results)

    for unit in manifest.get("units") or []:
        unit_no = int(unit.get("unit"))
        if not unit.get("unitPlan") or not item_exists(course_root, unit.get("unitPlan")) or not generated_plan_path(unit.get("unitPlan"), "unit-plans"):
            out_dir = course_root / "plans" / "generated" / "unit-plans"
            ensure_dir(out_dir)
            out = out_dir / f"{course_code}-U{unit_no:02d}-unit-plan.docx"
            create_unit_doc(course_code, course_root, unit, profile, out)
            item = plan_item(f"{course_code} Unit {unit_no} Teacher Unit Plan", rel(out, course_root), "unit_plan")
            update_bytes(course_root, item)
            unit["unitPlan"] = item
            results["unitPlansCreated"].append(item["path"])
        for lesson in unit.get("lessons") or []:
            if lesson.get("planningStatus") == "unit_overview":
                continue
            if lesson.get("lessonPlan") and item_exists(course_root, lesson.get("lessonPlan")) and generated_plan_path(lesson.get("lessonPlan"), "lesson-plans"):
                continue
            lesson_no = int(lesson.get("lesson"))
            out_dir = course_root / "plans" / "generated" / "lesson-plans" / f"Unit-{unit_no:02d}"
            ensure_dir(out_dir)
            out = out_dir / f"{course_code}-U{unit_no:02d}-L{lesson_no:02d}-{slug(lesson.get('title'))}-lesson-plan.docx"
            create_lesson_doc(course_code, course_root, unit, lesson, profile, out)
            item = plan_item(f"{course_code} U{unit_no}L{lesson_no} Teacher Lesson Plan - {lesson.get('title')}", rel(out, course_root), "lesson_plan")
            update_bytes(course_root, item)
            lesson["lessonPlan"] = item
            results["lessonPlansCreated"].append(item["path"])

    manifest.setdefault("sourceAudit", {})
    manifest["sourceAudit"]["teacherPrepPromotion"] = {
        "promotedAt": GENERATED_AT,
        "standard": "ICS3U teacher-prep enrichment; ENG3U display conventions",
        "course": course_code,
        "unitPlansCreated": len(results["unitPlansCreated"]),
        "lessonPlansCreated": len(results["lessonPlansCreated"]),
    }
    manifest["generatedAt"] = GENERATED_AT
    manifest_path.write_text(json.dumps(manifest, indent=2, ensure_ascii=False) + "\n", encoding="utf-8")
    return results


def main():
    results = [promote(code) for code in ("HHS4U", "LKBDU", "BBI1O")]
    output = {"generatedAt": GENERATED_AT, "results": results}
    REPORT_PATH.write_text(json.dumps(output, indent=2, ensure_ascii=False) + "\n", encoding="utf-8")
    print(json.dumps(output, indent=2, ensure_ascii=False))


if __name__ == "__main__":
    main()
