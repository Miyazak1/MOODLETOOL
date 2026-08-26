import json
import re
from datetime import datetime, timezone
from html import unescape
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


WORKSPACE_ROOT = Path("D:/工作文件/SUNNYBROOK")
REPO_ROOT = WORKSPACE_ROOT / "ossd-course-portal"
COURSE_ROOT = WORKSPACE_ROOT / "courseware" / "BOH4M"
MANIFEST_PATH = COURSE_ROOT / "course-manifest.json"
REPORT_PATH = REPO_ROOT / "deployment" / "BOH4M-docx-plans-report.json"
TEMPLATE_ROOT = REPO_ROOT / "templates" / "teacher-plans"
LESSON_PLAN_TEMPLATE = TEMPLATE_ROOT / "mdm4u-style-lesson-plan-template.docx"
UNIT_PLAN_TEMPLATE = TEMPLATE_ROOT / "mdm4u-style-unit-plan-template.docx"
GENERATED_AT = datetime.now(timezone.utc).replace(microsecond=0).isoformat()


GENERATED_SOURCE = (
    "Locally authored from the St.Mary BOH4M V2.0 Moodle course, localized lesson expectation "
    "pages, localized course files/media, the course outline, and The Ontario Curriculum, "
    "Grades 11 and 12: Business Studies, 2006 (revised), BOH4M."
)


CURRICULUM = {
    1: {
        "strand": "Foundations of Management",
        "overall": [
            "Assess the role of management within an organization.",
            "Demonstrate the use of appropriate communication techniques related to business management.",
            "Evaluate the impact of issues related to ethics and social responsibility on the management of organizations.",
        ],
        "specific_focus": [
            "Characteristics of organizations; levels of management; management theories and practices.",
            "Business communication using ICT, presentations, vocabulary, reports, and correspondence.",
            "Ethical and social-responsibility issues, stakeholder analysis, and management decision making.",
        ],
    },
    2: {
        "strand": "Leading",
        "overall": [
            "Apply an understanding of human behaviour to explain how individuals and groups function in the workplace.",
            "Demonstrate an understanding of group dynamics.",
            "Demonstrate an understanding of proper leadership techniques in a variety of situations.",
        ],
        "specific_focus": [
            "Personality, attitudes, perceptions, behaviour, and trait-assessment tools in the workplace.",
            "Stages of group development, group types, team success factors, and teamwork skills.",
            "Leadership traits, contemporary leadership theories, and leadership styles.",
        ],
    },
    3: {
        "strand": "Management Challenges",
        "overall": [
            "Demonstrate an understanding of the communication process within the workplace.",
            "Evaluate the strategies used by individuals and organizations to manage stress and conflict.",
            "Compare theories of how to motivate individuals and teams in a productive work environment.",
        ],
        "specific_focus": [
            "Communication barriers, active listening, feedback, technology tools, and perception.",
            "Workplace stress, conflict resolution, mediation, negotiation, and management responses.",
            "Content, process, and reinforcement motivation theories applied to individuals and teams.",
        ],
    },
    4: {
        "strand": "Planning and Controlling",
        "overall": [
            "Analyse the importance of planning to the success of an organization.",
            "Demonstrate an understanding of appropriate planning tools and techniques in a variety of situations.",
            "Analyse the relationship between strategic planning and the success of an organization.",
            "Analyse how companies respond to internal and external pressures for change.",
            "Assess the importance of control in management.",
        ],
        "specific_focus": [
            "Planning types, mission and vision, strategic management, and corporate culture.",
            "SWOT, BCG, Porter's Five Forces, PEST, and other strategic planning tools.",
            "Change management, internal/external pressures, legal considerations, and control systems.",
        ],
    },
    5: {
        "strand": "Organizing",
        "overall": [
            "Demonstrate an understanding of the various organizational structures used to manage the workforce effectively.",
            "Assess the ways in which organizational structures have changed to adapt to the changing nature of work.",
            "Evaluate the role of human resources within an organization.",
        ],
        "specific_focus": [
            "Traditional and current organizational structures, productivity, competitive advantage, and culture.",
            "Meaning of work, psychological contract, quality of work life, job design, and work arrangements.",
            "Legal considerations, recruitment, selection, development, retention, and performance appraisal.",
        ],
    },
}


UNIT_INTENT = {
    1: {
        "big_idea": "Managers shape organizational purpose, communication, ethical decisions, and social responsibility.",
        "teacher_intent": (
            "Help students move from naming management roles to judging how managers influence people, "
            "communication, stakeholder relationships, and responsible decision making."
        ),
        "essential": [
            "What makes management different from simply being in charge?",
            "How do communication choices affect trust and performance in organizations?",
            "How should managers balance organizational goals with ethics and social responsibility?",
        ],
    },
    2: {
        "big_idea": "Effective leadership depends on understanding people, groups, team dynamics, and leadership style.",
        "teacher_intent": (
            "Connect human behaviour and leadership theory to realistic workplace situations so students can "
            "explain why teams succeed or fail and how leaders adapt."
        ),
        "essential": [
            "How do personality, perception, and attitude influence behaviour at work?",
            "Why do groups need structure, roles, norms, and leadership?",
            "How should leaders choose techniques for different situations and team needs?",
        ],
    },
    3: {
        "big_idea": "Managers must communicate clearly, respond to stress and conflict, and motivate people productively.",
        "teacher_intent": (
            "Teach management challenges as practical decision points: diagnose the problem, select a strategy, "
            "communicate the rationale, and monitor the effect on people."
        ),
        "essential": [
            "How can managers reduce communication breakdowns in the workplace?",
            "What strategies help individuals and organizations manage stress and conflict responsibly?",
            "How do motivation theories help managers support performance and morale?",
        ],
    },
    4: {
        "big_idea": "Planning and controlling help organizations set direction, respond to change, and evaluate performance.",
        "teacher_intent": (
            "Have students use planning tools as thinking tools, not vocabulary lists: connect mission, analysis, "
            "strategy, change, and control to business decisions."
        ),
        "essential": [
            "How do mission, vision, and strategy guide management decisions?",
            "Which planning tool best fits a particular business problem or opportunity?",
            "How do control systems and change-management strategies help organizations adapt?",
        ],
    },
    5: {
        "big_idea": "Organizational structure and human resources decisions influence culture, productivity, and work quality.",
        "teacher_intent": (
            "Guide students to compare structures and HR practices through the lens of people, fairness, legal "
            "responsibility, and organizational effectiveness."
        ),
        "essential": [
            "How does structure affect communication, authority, productivity, and culture?",
            "How is work changing, and what does that mean for managers and employees?",
            "How can human resources practices support a high-quality, legally responsible workforce?",
        ],
    },
}


TOPIC_DETAILS = [
    (
        ["introduction to managers", "managers & organizations"],
        {
            "application": "identify organizational characteristics and evaluate levels of management, roles, responsibilities, skills, and competencies",
            "question": "What does a manager actually do to help an organization achieve its purpose?",
        },
    ),
    (
        ["management theories"],
        {
            "application": "compare major management theories and judge how they influence management practice",
            "question": "Which management theory best explains how people should be managed in a modern workplace?",
        },
    ),
    (
        ["ethics", "social responsibilities"],
        {
            "application": "evaluate ethical and social-responsibility issues and explain how they affect management decisions",
            "question": "How should managers decide when profit, people, and social responsibility are in tension?",
        },
    ),
    (
        ["social responsibility strategies"],
        {
            "application": "analyse corporate social-responsibility strategies using stakeholder evidence",
            "question": "What makes a social-responsibility strategy authentic rather than symbolic?",
        },
    ),
    (
        ["ethical behavior"],
        {
            "application": "apply an ethical decision-making lens to realistic workplace and community situations",
            "question": "How can managers recognize and respond to an ethical dilemma before it becomes organizational harm?",
        },
    ),
    (
        ["proper emails"],
        {
            "application": "compose business emails that use appropriate tone, format, audience awareness, and professional vocabulary",
            "question": "How does a professional email change the way a message is received and acted on?",
        },
    ),
    (
        ["presentation skills"],
        {
            "application": "prepare and deliver business presentations using clear structure, visuals, audience contact, and appropriate language",
            "question": "What makes a business presentation persuasive, professional, and useful?",
        },
    ),
    (
        ["multiple intelligence"],
        {
            "application": "explain how multiple-intelligence theory can influence leadership, training, and teamwork",
            "question": "How should leaders account for different ways people learn and contribute?",
        },
    ),
    (
        ["components of personality"],
        {
            "application": "analyse elements that shape personality and explain their impact on workplace behaviour",
            "question": "How can understanding personality help managers lead without stereotyping people?",
        },
    ),
    (
        ["learning theory", "behavioural psychology"],
        {
            "application": "connect learning theories and behaviour psychology to workplace training, feedback, and performance",
            "question": "How do people learn workplace behaviour, and how can managers influence it responsibly?",
        },
    ),
    (
        ["roles within groups"],
        {
            "application": "analyse group roles, stages, norms, and responsibilities in effective teams",
            "question": "Why do teams need roles, and what happens when roles are unclear or unhealthy?",
        },
    ),
    (
        ["group success", "group failure"],
        {
            "application": "analyse factors such as norms, cohesiveness, cultural expectations, and social loafing in team success or failure",
            "question": "What conditions make a group productive rather than merely busy?",
        },
    ),
    (
        ["leadership traits"],
        {
            "application": "analyse characteristics of effective leaders and connect them to equity, diversity, vision, and integrity",
            "question": "Are effective leaders born, developed, or shaped by the situation?",
        },
    ),
    (
        ["leadership styles"],
        {
            "application": "compare leadership styles and select an appropriate style for different workplace situations",
            "question": "When should a leader be directive, collaborative, supportive, or hands-off?",
        },
    ),
    (
        ["communication in the workplace"],
        {
            "application": "explain workplace communication processes, barriers, feedback, and perception",
            "question": "Why do workplace messages break down, and how can managers improve clarity?",
        },
    ),
    (
        ["stress", "conflict"],
        {
            "application": "evaluate stress and conflict-management strategies for individuals and organizations",
            "question": "How should managers respond when stress or conflict begins to affect performance and trust?",
        },
    ),
    (
        ["motivation in the workplace"],
        {
            "application": "explain how motivation affects productivity, morale, and management decisions",
            "question": "What motivates people to do high-quality work beyond simply being paid?",
        },
    ),
    (
        ["content motivation"],
        {
            "application": "compare content motivation theories and apply them to workplace needs and incentives",
            "question": "Which human needs should managers consider when designing motivating work?",
        },
    ),
    (
        ["process and reinforcement"],
        {
            "application": "compare process and reinforcement theories and evaluate how feedback and rewards shape behaviour",
            "question": "How do expectations, fairness, reinforcement, and feedback influence motivation?",
        },
    ),
    (
        ["introduction to planning"],
        {
            "application": "analyse why planning supports organizational direction, coordination, and control",
            "question": "Why is planning a management responsibility rather than a paperwork exercise?",
        },
    ),
    (
        ["mission", "vision"],
        {
            "application": "distinguish mission and vision statements and judge how they guide decisions",
            "question": "How do mission and vision statements shape what an organization chooses to do?",
        },
    ),
    (
        ["strategic management"],
        {
            "application": "connect strategic management to long-term goals, competitive position, and organizational success",
            "question": "How do managers turn analysis into strategy?",
        },
    ),
    (
        ["strategic planning tools"],
        {
            "application": "select and use planning tools to analyse organizational situations",
            "question": "Which strategic planning tool best fits the decision a manager needs to make?",
        },
    ),
    (
        ["swot"],
        {
            "application": "construct and interpret a SWOT analysis to support a business decision",
            "question": "How can a SWOT analysis become evidence for action instead of a list?",
        },
    ),
    (
        ["bcg"],
        {
            "application": "use the BCG matrix to analyse products or business units and recommend resource allocation",
            "question": "How should managers decide where to invest limited resources?",
        },
    ),
    (
        ["porter", "pest"],
        {
            "application": "use Porter's Five Forces and PEST analysis to evaluate external pressures and opportunities",
            "question": "How do market forces and external trends change strategic choices?",
        },
    ),
    (
        ["control process", "discipline"],
        {
            "application": "describe the control process, discipline, and internal/external control measures",
            "question": "How can control systems support accountability without damaging morale?",
        },
    ),
    (
        ["change management"],
        {
            "application": "analyse causes of change, attitudes toward change, and strategies for managing acceptance",
            "question": "Why do people resist change, and how can managers lead change responsibly?",
        },
    ),
    (
        ["organizational structures"],
        {
            "application": "compare traditional and current organizational structures and their impact on productivity and culture",
            "question": "How does structure shape how people communicate, decide, and work?",
        },
    ),
    (
        ["trends in business management"],
        {
            "application": "analyse trends in organizational design and management practice",
            "question": "How are management structures changing to fit new work realities?",
        },
    ),
    (
        ["nature of work"],
        {
            "application": "assess changing work arrangements, job design, quality of work life, and job satisfaction",
            "question": "What does meaningful and productive work look like in a changing organization?",
        },
    ),
    (
        ["human resources"],
        {
            "application": "evaluate HR processes, legal responsibilities, workforce development, and performance appraisal",
            "question": "How can HR decisions support both organizational goals and employee dignity?",
        },
    ),
]


def clean_text(value):
    text = str(value or "")
    text = re.sub(r"<script[\s\S]*?</script>", " ", text, flags=re.I)
    text = re.sub(r"<style[\s\S]*?</style>", " ", text, flags=re.I)
    text = re.sub(r"<br\s*/?>", "\n", text, flags=re.I)
    text = re.sub(r"</(p|div|li|tr|h[1-6])>", "\n", text, flags=re.I)
    text = re.sub(r"<[^>]+>", " ", text)
    text = unescape(text)
    text = text.replace("\u00a0", " ")
    return re.sub(r"\s+", " ", text).strip()


def read_local_text(relative_path):
    if not relative_path:
        return ""
    path = COURSE_ROOT / relative_path
    if not path.exists():
        return ""
    return clean_text(path.read_text(encoding="utf-8", errors="ignore"))


def short_text(value, limit=340):
    text = clean_text(value)
    if len(text) <= limit:
        return text
    clipped = text[:limit]
    cut = max(clipped.rfind(". "), clipped.rfind("; "), clipped.rfind(", "))
    if cut > 140:
        clipped = clipped[: cut + 1]
    return clipped.strip() + "..."


def unique(items):
    out = []
    seen = set()
    for item in items:
        value = clean_text(item)
        key = value.lower()
        if value and key not in seen:
            seen.add(key)
            out.append(value)
    return out


def item_label(item):
    return clean_text(item.get("label") or item.get("title") or Path(str(item.get("path") or "")).name)


def find_section(lesson, pattern):
    regex = re.compile(pattern, re.I)
    for section in lesson.get("bookSections") or []:
        label = section.get("sectionLabel") or section.get("label") or ""
        if regex.search(label):
            return section
    return None


def section_text(lesson, pattern):
    section = find_section(lesson, pattern)
    if not section:
        return ""
    return read_local_text(section.get("path")) or clean_text(section.get("textPreview"))


def extract_between(text, start_labels, end_labels):
    source = clean_text(text)
    lower = source.lower()
    start = -1
    label_len = 0
    for label in start_labels:
        idx = lower.find(label.lower())
        if idx >= 0 and (start < 0 or idx < start):
            start = idx
            label_len = len(label)
    if start < 0:
        return ""
    start += label_len
    while start < len(source) and source[start] in ": -":
        start += 1
    end = len(source)
    for label in end_labels:
        idx = lower.find(label.lower(), start)
        if idx >= 0:
            end = min(end, idx)
    return clean_text(source[start:end])


def strip_expectation_heading(text, lesson):
    out = clean_text(text)
    patterns = [
        rf"^Lesson\s+{lesson.get('lesson')}\s*:\s*{re.escape(str(lesson.get('title') or ''))}",
        rf"^Lesson\s+{lesson.get('lesson')}\s*:",
        r"^Lesson Expectations",
    ]
    changed = True
    while changed:
        changed = False
        for pattern in patterns:
            new = re.sub(pattern, "", out, flags=re.I).strip()
            if new != out:
                out = clean_text(new)
                changed = True
    return out


def split_items(block, limit=8):
    block = clean_text(block)
    if not block:
        return []
    pieces = re.split(r"\s*;\s*|\s*(?<=\.)\s+(?=[A-Z])", block)
    return [short_text(piece.strip(" .;"), 280) for piece in pieces if len(clean_text(piece)) > 10][:limit]


def expectation_for(lesson):
    raw = strip_expectation_heading(section_text(lesson, r"expectations"), lesson)
    overall = extract_between(
        raw,
        ["Overall Expectations", "Overall Expectation"],
        ["Specific Lesson Expectations", "Specific Expectations", "Learning Goals", "Success Criteria"],
    )
    specific = extract_between(
        raw,
        ["Specific Lesson Expectations", "Specific Expectations", "Specific Lesson Expectation"],
        ["Learning Goals", "Success Criteria"],
    )
    goals = extract_between(raw, ["Learning Goals", "Learning Goal"], ["Success Criteria"])
    success = extract_between(raw, ["Success Criteria", "Success Criterion"], [])
    return {
        "raw": raw,
        "overall": split_items(overall, 6),
        "specific": split_items(specific, 8),
        "goals": split_items(goals, 6),
        "success": split_items(success, 6),
        "exists": bool(raw),
    }


def lesson_topic_detail(lesson):
    title = clean_text(lesson.get("title", "")).lower()
    for keys, detail in TOPIC_DETAILS:
        if any(key in title for key in keys):
            return detail
    return {
        "application": f"apply BOH4M management concepts to explain {lesson.get('title', 'the lesson topic')}",
        "question": f"How does {lesson.get('title', 'this topic')} help managers make better decisions?",
    }


def resource_items(lesson):
    items = []
    items.extend(lesson.get("downloads") or [])
    items.extend(lesson.get("textExports") or [])
    for section in lesson.get("bookSections") or []:
        items.extend(section.get("attachments") or [])
    return items


def lesson_materials(lesson):
    media = []
    files = []
    for item in lesson.get("ispring") or []:
        media.append(f"iSpring lesson presentation: {item.get('label') or 'localized presentation'}")
    for item in resource_items(lesson):
        label = item_label(item)
        haystack = " ".join(str(item.get(k, "")) for k in ("type", "category", "role", "path", "label"))
        if re.search(r"h5p", haystack, re.I):
            media.append(f"Localized H5P activity: {label}")
        elif re.search(r"mp4|video", haystack, re.I):
            media.append(f"Video resource: {label}")
        elif re.search(r"docx?|pdf|pptx?|xlsx?|worksheet|rubric|assignment|kwl|reflection|learning log|exit", haystack, re.I):
            files.append(label)
    return unique(media), unique(files)


def lesson_activity_summary(lesson):
    pieces = []
    for label, pattern, limit in [
        ("Lesson page", r"^Lesson$", 230),
        ("Hands On", r"Hands On", 180),
        ("Consolidation", r"Consolidation", 180),
        ("Homework", r"Homework", 210),
    ]:
        text = section_text(lesson, pattern)
        if text:
            pieces.append(f"{label}: {short_text(text, limit)}")
    return "\n".join(pieces) or "Use the localized Moodle lesson page and attached course materials."


def prior_knowledge(unit, lesson):
    unit_no = int(unit["unit"])
    lesson_no = int(lesson.get("lesson", 1))
    if lesson_no == 1:
        return (
            f"Students begin {unit['title']} by activating everyday experiences with managers, teams, "
            "organizations, communication, or workplace decision making. Start with a KWL prompt or "
            "short business scenario and ask students to separate opinion from evidence."
        )
    prev = next((item for item in unit.get("lessons") or [] if item.get("lesson") == lesson_no - 1), None)
    previous = f"Unit {unit_no} Lesson {prev['lesson']} ({prev['title']})" if prev else "the previous lesson"
    return (
        f"Students build from {previous}. Revisit one key management term or decision from the previous "
        f"lesson, then connect it to {lesson.get('title')} with a practical workplace question."
    )


def lesson_learning_goals(unit, lesson, exp):
    if exp["goals"]:
        return exp["goals"]
    detail = lesson_topic_detail(lesson)
    return unique(
        [
            f"Students will explain how {lesson['title']} connects to {unit['title']} and to the work of managers.",
            f"Students will learn to {detail['application']}.",
            "Students will use business vocabulary and evidence from the lesson, presentation, activity, and homework task.",
        ]
    )


def lesson_success_criteria(unit, lesson, exp):
    if exp["success"]:
        return exp["success"]
    detail = lesson_topic_detail(lesson)
    return unique(
        [
            "I can define and correctly use the key management vocabulary from the lesson.",
            "I can support a management judgement with evidence from a scenario, lesson page, presentation, or activity.",
            f"I can {detail['application']} in a written, oral, or interactive response.",
            "I can complete the assigned homework/submission evidence using the correct Moodle folder or file.",
        ]
    )


def assessment_lists(unit, lesson):
    labels = " ".join(item_label(item) for item in resource_items(lesson)).lower()
    asl = ["[x] Observation", "[x] Anecdotal notes", "[x] Student self-check against success criteria"]
    asl.append("[x] KWL/reflection/exit evidence" if re.search(r"kwl|reflection|learning log|exit", labels) else "[ ] KWL/reflection/exit evidence")
    afl = ["[x] Strategic questioning", "[x] Guided worksheet or activity check", "[x] Homework preparation check"]
    if lesson.get("handsOn") or section_text(lesson, r"Hands On"):
        afl.append("[x] Hands On/localized H5P practice")
    if lesson.get("ispring"):
        afl.append("[x] iSpring note-taking checkpoint")
    unit_evals = [item_label(item) for item in (unit.get("unitResources") or {}).get("evaluations") or []]
    aol = [f"[x] {label}" for label in unit_evals[:6]]
    if not aol:
        aol = ["[ ] No formal AoL in this lesson unless identified by the Moodle Evaluation area"]
    return "\n".join(asl), "\n".join(afl), "\n".join(aol)


def accommodations_for(lesson):
    title = clean_text(lesson.get("title", ""))
    return (
        "Chunk readings, iSpring presentations, and case prompts into short checkpoints; pre-teach management "
        "vocabulary; provide sentence frames for claim, evidence, and recommendation; allow oral rehearsal before "
        "written responses; provide graphic organizers for theories, stakeholder analysis, planning tools, or HR "
        "processes; allow extra time for note-taking/uploading; extend by asking students to compare "
        f"{title} with a current business example selected by the teacher."
    )


def materials_for(lesson):
    media, files = lesson_materials(lesson)
    lines = unique(media[:10] + files[:12])
    return "\n".join(lines) or "Localized Moodle lesson page, teacher-selected case scenario, and course outline resources."


def teacher_assessment_action(unit, lesson):
    return (
        "Confer with students during case analysis, note-taking, and worksheet completion.\n"
        "Observe whether students use management evidence rather than unsupported opinion.\n"
        "Give descriptive feedback tied to the learning goal and success criteria.\n"
        "Grade only when the task appears in the Evaluation/AoL area; otherwise use work as practice evidence."
    )


def unit_learning_goals(unit):
    intent = UNIT_INTENT[int(unit["unit"])]
    return [
        f"Students will explain the big idea of {unit['title']} using accurate BOH4M management vocabulary.",
        "Students will interpret management situations by considering people, structure, communication, ethics, planning, and evidence.",
        "Students will communicate business recommendations through discussion, interactive practice, homework, and written responses.",
        *intent["essential"],
    ]


def unit_summary(unit):
    lessons = unit.get("lessons") or []
    intent = UNIT_INTENT[int(unit["unit"])]
    sequence = ", ".join(lesson.get("title", "") for lesson in lessons)
    return (
        f"{intent['big_idea']}\n\n"
        f"Instructional sequence: {sequence}.\n\n"
        f"Teacher planning focus: {intent['teacher_intent']}\n\n"
        "Students work through localized Moodle lesson pages, iSpring presentations, Hands On/H5P practice, "
        "consolidation/exit evidence, homework submission files, reflection routines, and Moodle Evaluation tasks."
    )


def unit_outline(unit):
    chunks = []
    for lesson in unit.get("lessons") or []:
        exp = expectation_for(lesson)
        detail = lesson_topic_detail(lesson)
        target = "; ".join(exp["specific"][:2]) or "; ".join(exp["overall"][:2]) or "Use the localized Lesson Expectations page for exact wording."
        media, files = lesson_materials(lesson)
        evidence = []
        if media:
            evidence.append("media/practice: " + "; ".join(media[:4]))
        if files:
            evidence.append("files: " + "; ".join(files[:5]))
        evidence.append("assessment: Hands On/consolidation checks and homework submission where listed")
        chunks.append(
            f"{lesson['id']} - {lesson['title']}\n"
            f"Management question: {detail['question']}\n"
            f"Target: {target}\n"
            "Teacher plan: activate prior knowledge, model management vocabulary and evidence use, guide the Moodle lesson sequence, "
            "and close with a reflection or submission-based check.\n"
            + "\n".join(evidence)
        )
    for item in (unit.get("unitResources") or {}).get("evaluations") or []:
        chunks.append(
            f"Assessment of Learning - {item_label(item)}\n"
            "Students apply unit learning in a Moodle-identified assessed task. Feedback should address knowledge, thinking/inquiry, communication, and application."
        )
    return "\n\n".join(chunks)


def unit_technology(unit):
    lines = []
    for lesson in unit.get("lessons") or []:
        media, _files = lesson_materials(lesson)
        for item in media:
            lines.append(f"{lesson['id']} - {item}")
    return "\n".join(unique(lines)) or "Localized Moodle activity pages and course resources."


def unit_printed(unit):
    lines = []
    for lesson in unit.get("lessons") or []:
        _media, files = lesson_materials(lesson)
        for item in files:
            lines.append(f"{lesson['id']} - {item}")
    for key in ["evaluations", "lessonDropboxes", "reflectionAndLogs"]:
        for item in (unit.get("unitResources") or {}).get(key) or []:
            for attachment in item.get("attachments") or []:
                lines.append(f"{item_label(item)} - {item_label(attachment)}")
    return "\n".join(unique(lines)) or "Course outline, homework files, evaluation files, and any Moodle attachments listed in the course."


def template_document(kind):
    if kind == "lesson":
        return Document(str(LESSON_PLAN_TEMPLATE))
    if kind == "unit":
        return Document(str(UNIT_PLAN_TEMPLATE))
    raise ValueError(f"Unknown template kind: {kind}")


def style_document(doc):
    section = doc.sections[0]
    section.top_margin = Inches(0.9)
    section.bottom_margin = Inches(0.9)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.12
    for name, size, color in [("Heading 1", 16, "002B57"), ("Heading 2", 12.5, "002B57"), ("Heading 3", 10.5, "002B57")]:
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)


def set_text(cell, text, bold=False, size=9.5):
    cell.text = ""
    lines = str(text or "").split("\n")
    for line_index, line in enumerate(lines):
        paragraph = cell.paragraphs[0] if line_index == 0 else cell.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(0)
        run = paragraph.add_run(line)
        run.bold = bold
        run.font.name = "Calibri"
        run.font.size = Pt(size)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP


def set_row_text(row, text, bold=False, size=9.5):
    set_text(row.cells[0], text, bold=bold, size=size)


def lesson_plan_doc(unit, lesson):
    exp = expectation_for(lesson)
    curriculum = CURRICULUM[int(unit["unit"])]
    doc = template_document("lesson")
    style_document(doc)
    rows = doc.tables[0].rows
    set_text(rows[0].cells[0], "Lesson Plan", bold=True)
    set_text(rows[0].cells[2], "Subject: BOH4M Business Leadership: Management Fundamentals, Grade 12 University/College", bold=True)
    set_row_text(rows[1], f"Lesson Name: {lesson['title']}")
    set_row_text(rows[2], f"Unit of Study: {unit['title']}")
    set_row_text(
        rows[3],
        "PRIOR KNOWLEDGE What do my students already know? What key questions or strategies will activate prior knowledge? What connections can I help students make?\n"
        + prior_knowledge(unit, lesson),
    )
    local_overall = exp["overall"] or curriculum["overall"]
    local_specific = exp["specific"] or curriculum["specific_focus"]
    set_row_text(
        rows[4],
        "\n".join(
            [
                "CURRICULUM EXPECTATIONS",
                "OVERALL",
                *local_overall,
                "SPECIFIC",
                *local_specific,
                "Curriculum alignment: " + curriculum["strand"],
            ]
        ),
    )
    set_row_text(rows[5], "\n".join(["LEARNING GOALS What do I want students to know and be able to do?", *lesson_learning_goals(unit, lesson, exp)]))
    set_row_text(rows[6], "\n".join(["SUCCESS CRITERIA(S) How will students know they attained the goal? What tools can they use to self-assess?", *lesson_success_criteria(unit, lesson, exp)]))
    asl, afl, aol = assessment_lists(unit, lesson)
    set_text(rows[7].cells[0], "Assessment as Learning", bold=True)
    set_text(rows[7].cells[1], "Assessment for Learning", bold=True)
    set_text(rows[7].cells[2], "Assessment of Learning", bold=True)
    set_text(rows[8].cells[0], asl)
    set_text(rows[8].cells[1], afl)
    set_text(rows[8].cells[2], aol)
    set_row_text(rows[9], "What will I do?\n" + teacher_assessment_action(unit, lesson))
    set_row_text(rows[10], "Accommodations: How will you change the lesson to meet the needs of individual students?\n" + accommodations_for(lesson))
    set_row_text(
        rows[11],
        "Materials and Resources: What do I need to prepare in advance? Equipment? Student resources? Teacher resources?\n"
        + materials_for(lesson)
        + f"\nSource status: locally authored from BOH4M St.Mary Moodle content on {GENERATED_AT}.",
    )

    lesson_table = doc.tables[1]
    set_text(lesson_table.rows[0].cells[0], "DELIVERING THE LESSON", bold=True)
    detail = lesson_topic_detail(lesson)
    phases = [
        (
            1,
            "Timing\n5-8\nminutes",
            "Minds On!\nOpen with a management scenario, KWL prompt, vocabulary prediction, or quick-write. Name the management question and connect it to the lesson expectation before students begin the Moodle sequence.",
            "Materials/Resources\nLesson Expectations page\nKWL or scenario prompt",
        ),
        (
            2,
            "Timing\n20-30\nminutes",
            "Action!\nUse the Lesson page and iSpring presentation to model the management concept. Pause for vocabulary checks and ask students to connect the idea to an organization, manager, or workplace decision.",
            "Materials/Resources\nLocalized lesson page\niSpring presentation",
        ),
        (
            3,
            "Timing\n15-25\nminutes",
            f"Action!\nGuide students through the Hands On/H5P/practice task. Require a short response that explains how students can {detail['application']} with evidence.",
            "Materials/Resources\nHands On section\nLocalized H5P/practice resources\nTeacher observation notes",
        ),
        (
            4,
            "Timing\n10-15\nminutes",
            "Consolidation!\nUse the consolidation/exit activity or homework instructions to check understanding. Confirm submission location and have students identify one revision target before uploading work.",
            "Materials/Resources\nHomework/consolidation instructions\nAttached worksheet or guide",
        ),
    ]
    for row_index, timing, action, materials in phases:
        row = lesson_table.rows[row_index]
        set_text(row.cells[0], timing)
        set_text(row.cells[1], "W\n[ ]")
        set_text(row.cells[2], "S\n[ ]")
        set_text(row.cells[3], "I\n[ ]")
        set_text(row.cells[4], action + "\n\n" + short_text(lesson_activity_summary(lesson), 430))
        set_text(row.cells[5], materials)

    set_text(
        doc.tables[2].rows[0].cells[0],
        "Notes:\nGenerated teacher planning aid reconstructed from localized BOH4M St.Mary Moodle content. Keep Answer Keys in the teacher packet area; use Homework Submission Folder resources for student-facing submission workflow. Teacher reflection: note which management vocabulary, theory, or evidence habit needs reteaching next lesson.",
    )
    return doc


def unit_plan_doc(unit):
    unit_no = int(unit["unit"])
    curriculum = CURRICULUM[unit_no]
    intent = UNIT_INTENT[unit_no]
    doc = template_document("unit")
    style_document(doc)
    if doc.paragraphs:
        doc.paragraphs[0].text = f"UNIT {unit_no}"
        doc.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    rows = doc.tables[0].rows
    for row_index, label in {0: "Unit Author", 5: "Unit Overview", 14: "Unit Foundation", 19: "Assessment Plan", 22: "Unit Details", 25: "Materials and Resources"}.items():
        set_text(rows[row_index].cells[0], label, bold=True)
    for row_index, label in [(1, "Name:"), (2, "School District:"), (3, "School Name:"), (4, "School City, Province:")]:
        set_row_text(rows[row_index], label)

    values = [
        (6, "Unit Title Name"),
        (7, unit["title"]),
        (8, "Unit Summary"),
        (9, unit_summary(unit)),
        (10, "Year Level"),
        (11, "Grade 12, University/College Business Leadership"),
        (12, "Approximate Time Needed"),
        (13, f"{max(len(unit.get('lessons') or []) * 3, 6)} instructional hours, plus homework and evaluation time"),
        (15, "Targeted Curriculum Expectations"),
        (
            16,
            "\n".join(
                [
                    f"Ontario curriculum strand: {curriculum['strand']}",
                    "Overall expectations:",
                    *curriculum["overall"],
                    "Specific focus for this unit:",
                    *curriculum["specific_focus"],
                    "Local lesson expectation pages provide exact lesson-level wording and success criteria.",
                ]
            ),
        ),
        (17, "Learning Goals"),
        (18, "\n".join(unit_learning_goals(unit))),
    ]
    for row_index, value in values:
        set_row_text(rows[row_index], value)

    set_text(rows[20].cells[0], "Assessment as Learning (ASL)", bold=True)
    set_text(rows[20].cells[2], "Assessment for Learning (AFL)", bold=True)
    set_text(rows[20].cells[3], "Assessment of Learning (AoL)", bold=True)
    set_text(rows[21].cells[0], "[x] KWL/reflection routines\n[x] Exit slips\n[x] Learning log\n[x] Student next-step notes")
    set_text(rows[21].cells[2], "[x] Hands On/H5P checks\n[x] Strategic questioning\n[x] iSpring note checks\n[x] Homework review")
    evaluations = [item_label(item) for item in (unit.get("unitResources") or {}).get("evaluations") or []]
    set_text(rows[21].cells[3], "[x] " + "\n[x] ".join(evaluations[:10]) if evaluations else "[x] Moodle Evaluation tasks where listed")
    set_row_text(rows[23], "Lesson and Assessment Outlines", bold=True)
    set_row_text(rows[24], unit_outline(unit))
    set_text(rows[26].cells[0], "Technology", bold=True)
    set_text(rows[26].cells[1], unit_technology(unit))
    set_text(rows[27].cells[0], "Printed", bold=True)
    set_text(rows[27].cells[1], unit_printed(unit))
    set_text(rows[28].cells[0], "Other Resources", bold=True)
    set_text(
        rows[28].cells[1],
        "\n".join(
            [
                "Course outline and Learning Log",
                "Anecdotal observation notes",
                "Current business examples selected by the teacher",
                "Student self-assessment and reflection evidence",
                "Ontario Business Studies curriculum expectations for BOH4M",
                f"Essential questions: {'; '.join(intent['essential'])}",
            ]
        ),
    )
    for row in rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    return doc


def save_doc(doc, relative_path):
    path = COURSE_ROOT / relative_path
    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(path)
    return path.stat().st_size


def file_record(label, relative_path, category, role):
    path = COURSE_ROOT / relative_path
    preview = f"previews-html/{relative_path}.html"
    return {
        "label": label,
        "type": "docx",
        "category": category,
        "role": role,
        "path": relative_path.replace("\\", "/"),
        "bytes": path.stat().st_size,
        "source": "reconstructed_from_moodle_content",
        "sourceStatus": "locally_authored_from_course_materials",
        "teacherUse": "Teacher planning aid reconstructed from localized BOH4M St.Mary course content; not an original Moodle teacher packet.",
        "templateReference": "MDM4U teacher plan format and local mdm4u-style plan templates.",
        "previewPath": preview.replace("\\", "/"),
    }


def main():
    if not MANIFEST_PATH.exists():
        raise FileNotFoundError(MANIFEST_PATH)
    manifest = json.loads(MANIFEST_PATH.read_text(encoding="utf-8"))
    unit_count = 0
    lesson_count = 0
    written = []
    missing_expectations = []

    for unit in manifest.get("units") or []:
        unit_rel = f"plans/generated/unit-plans/U{int(unit['unit']):02d}-unit-plan.docx"
        save_doc(unit_plan_doc(unit), unit_rel)
        unit["unitPlan"] = file_record(f"Unit Plan - {unit['title']}", unit_rel, "unit_plan", "unit_plan")
        unit_count += 1
        written.append(unit_rel)

        for lesson in unit.get("lessons") or []:
            exp = expectation_for(lesson)
            if not exp["exists"]:
                missing_expectations.append(lesson.get("id"))
            lesson_rel = f"plans/generated/lesson-plans/{lesson['id']}-lesson-plan.docx"
            save_doc(lesson_plan_doc(unit, lesson), lesson_rel)
            lesson["lessonPlan"] = file_record(
                f"Lesson Plan - Unit {unit['unit']} Lesson {lesson['lesson']}",
                lesson_rel,
                "lesson_plan",
                "lesson_plan",
            )
            lesson["resourceCounts"] = {**(lesson.get("resourceCounts") or {}), "lessonPlan": 1}
            lesson_count += 1
            written.append(lesson_rel)

    manifest["generatedAt"] = GENERATED_AT
    manifest["sourceAudit"] = {
        **(manifest.get("sourceAudit") or {}),
        "teacherPlanGeneration": {
            "generatedAt": GENERATED_AT,
            "format": "docx",
            "unitPlans": unit_count,
            "lessonPlans": lesson_count,
            "source": "reconstructed_from_moodle_content",
            "basis": [
                "St.Mary BOH4M V2.0 Moodle course id 40",
                "localized Lesson Expectations pages",
                "localized Lesson, Hands On, Consolidation, and Homework sections",
                "localized iSpring, H5P, worksheet, homework, reflection, exit card, Evaluation, and Answer Key records",
                "BOH4M Course Outline and local MDM4U-style teacher plan templates",
                "The Ontario Curriculum, Grades 11 and 12: Business Studies, 2006 (revised), BOH4M",
            ],
            "boundary": "Teacher planning aids only; no answer keys, rubrics, textbook excerpts, or Moodle-original teacher packet documents were fabricated.",
            "missingExpectationPages": missing_expectations,
        },
    }
    MANIFEST_PATH.write_text(json.dumps(manifest, indent=2, ensure_ascii=False) + "\n", encoding="utf-8")
    REPORT_PATH.parent.mkdir(parents=True, exist_ok=True)
    REPORT_PATH.write_text(
        json.dumps(
            {
                "course": "BOH4M",
                "unitPlans": unit_count,
                "lessonPlans": lesson_count,
                "written": written,
                "missingExpectationPages": missing_expectations,
                "generatedAt": GENERATED_AT,
            },
            indent=2,
            ensure_ascii=False,
        )
        + "\n",
        encoding="utf-8",
    )
    print(json.dumps({"course": "BOH4M", "unitPlans": unit_count, "lessonPlans": lesson_count, "missingExpectationPages": missing_expectations, "generatedAt": GENERATED_AT}, indent=2))


if __name__ == "__main__":
    main()
